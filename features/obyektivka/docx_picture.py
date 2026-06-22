"""Obyektivka header foto: VML placeholder (namuna) + wp:anchor (yuklangan rasm)."""

from __future__ import annotations

import html
from copy import deepcopy

from docx.image.image import Image
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Emu, Length, Pt
from docx.text.paragraph import Paragraph

from features.obyektivka.layout import (
    PHOTO_VML_HEIGHT_PT,
    PHOTO_VML_MARGIN_LEFT_PT,
    PHOTO_VML_MARGIN_TOP_PT,
    PHOTO_VML_WIDTH_PT,
    PHOTO_VML_Z_INDEX,
)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
V_NS = "urn:schemas-microsoft-com:vml"
O_NS = "urn:schemas-microsoft-com:office:office"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

_spid = 1029


def _emu(length: Length | int) -> int:
    return int(length)


def _next_spid() -> str:
    global _spid
    _spid += 1
    return f"_x0000_s{_spid}"


def _vml_rect_style() -> str:
    return (
        f"position:absolute;left:0pt;margin-left:{PHOTO_VML_MARGIN_LEFT_PT}pt;"
        f"margin-top:{PHOTO_VML_MARGIN_TOP_PT}pt;height:{PHOTO_VML_HEIGHT_PT}pt;"
        f"width:{PHOTO_VML_WIDTH_PT}pt;z-index:{PHOTO_VML_Z_INDEX};"
        "mso-width-relative:page;mso-height-relative:page;"
    )


def _hint_textbox_xml(hint_text: str) -> str:
    safe = html.escape(hint_text, quote=False)
    return f"""
      <v:textbox>
        <w:txbxContent>
          <w:p>
            <w:pPr>
              <w:ind w:left="-142" w:right="-119"/>
              <w:jc w:val="center"/>
              <w:rPr><w:sz w:val="22"/><w:szCs w:val="22"/></w:rPr>
            </w:pPr>
            <w:r>
              <w:rPr><w:sz w:val="22"/><w:szCs w:val="22"/></w:rPr>
              <w:t xml:space="preserve">{safe}</w:t>
            </w:r>
          </w:p>
        </w:txbxContent>
      </v:textbox>
    """


def add_vml_photo_placeholder(paragraph: Paragraph, *, hint_text: str) -> None:
    """Namuna p1: v:rect + matn (foto yuklanmaganda)."""
    run = paragraph.add_run()
    spid = _next_spid()
    style = html.escape(_vml_rect_style(), quote=True)
    hint_box = _hint_textbox_xml(hint_text or " ")
    pict_xml = f"""
    <w:pict
      xmlns:w="{W_NS}"
      xmlns:v="{V_NS}"
      xmlns:o="{O_NS}"
      xmlns:r="{R_NS}">
      <v:rect id="{spid}" o:spid="{spid}" o:spt="1"
        style="{style}" coordsize="21600,21600">
        <v:path/>
        <v:fill focussize="0,0"/>
        <v:stroke/>
        <v:imagedata o:title=""/>
        <o:lock v:ext="edit"/>
        {hint_box}
      </v:rect>
    </w:pict>
    """
    run._r.append(parse_xml(pict_xml))


def add_floating_picture(
    paragraph: Paragraph,
    image_path: str,
    *,
    width: Length,
    height: Length,
    pos_x: Length | None = None,
    pos_y: Length | None = None,
    relative_from_page: bool = False,
) -> None:
    """Floating image — namuna VML koordinatalari (page-relative) bilan."""
    run = paragraph.add_run()
    run.add_picture(image_path, width=width, height=height)
    inlines = run._r.xpath(".//wp:inline")
    if not inlines:
        return
    inline = inlines[0]
    anchor = OxmlElement("wp:anchor")
    anchor.set(qn("wp:distT"), "0")
    anchor.set(qn("wp:distB"), "0")
    anchor.set(qn("wp:distL"), "0")
    anchor.set(qn("wp:distR"), "0")
    anchor.set(qn("wp:simplePos"), "0")
    anchor.set(qn("wp:relativeHeight"), str(PHOTO_VML_Z_INDEX))
    anchor.set(qn("wp:behindDoc"), "0")
    anchor.set(qn("wp:locked"), "0")
    anchor.set(qn("wp:layoutInCell"), "1")
    anchor.set(qn("wp:allowOverlap"), "1")

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")
    anchor.append(simple_pos)

    pos_h = OxmlElement("wp:positionH")
    pos_h.set(qn("wp:relativeFrom"), "page" if relative_from_page else "column")
    if pos_x is not None:
        off = OxmlElement("wp:posOffset")
        off.text = str(_emu(pos_x))
        pos_h.append(off)
    else:
        align = OxmlElement("wp:align")
        align.text = "right"
        pos_h.append(align)
    anchor.append(pos_h)

    pos_v = OxmlElement("wp:positionV")
    pos_v.set(qn("wp:relativeFrom"), "page" if relative_from_page else "paragraph")
    off_v = OxmlElement("wp:posOffset")
    off_v.text = str(_emu(pos_y or Emu(0)))
    pos_v.append(off_v)
    anchor.append(pos_v)

    extent = inline.find(qn("wp:extent"))
    if extent is not None:
        anchor.append(deepcopy(extent))
    effect = inline.find(qn("wp:effectExtent"))
    if effect is not None:
        anchor.append(deepcopy(effect))

    doc_pr = inline.find(qn("wp:docPr"))
    if doc_pr is not None:
        anchor.append(deepcopy(doc_pr))

    c_nv = inline.find(qn("wp:cNvGraphicFramePr"))
    if c_nv is not None:
        anchor.append(deepcopy(c_nv))

    graphic = inline.find(qn("a:graphic"))
    if graphic is not None:
        anchor.append(deepcopy(graphic))

    inline.getparent().replace(inline, anchor)


def clear_photo_hint_from_paragraph(paragraph: Paragraph) -> None:
    """Foto yuklanganda VML matn ramkasi va dublikat izohni olib tashlash."""
    p_el = paragraph._p
    w_pict = f"{{{W_NS}}}pict"
    for pict in list(p_el.iter(w_pict)):
        parent = pict.getparent()
        if parent is not None:
            parent.remove(pict)
    for run in list(paragraph.runs):
        text = run.text or ""
        low = text.lower()
        if (
            "{{photo}}" in text
            or "fotosurat" in low
            or "фотосурат" in low
            or "3x4" in low
            or "3х4" in low
            or "rasmiy kiyimda" in low
            or "расмий кийимда" in low
        ):
            run.text = ""


def add_reference_photo(paragraph: Paragraph, image_path: str) -> None:
    """3×4 ramka koordinatalarida rasm (VML izoh o'rniga)."""
    clear_photo_hint_from_paragraph(paragraph)
    add_floating_picture(
        paragraph,
        image_path,
        width=Pt(PHOTO_VML_WIDTH_PT),
        height=Pt(PHOTO_VML_HEIGHT_PT),
        pos_x=Pt(PHOTO_VML_MARGIN_LEFT_PT),
        pos_y=Pt(PHOTO_VML_MARGIN_TOP_PT),
        relative_from_page=True,
    )


def find_photo_paragraph(doc) -> Paragraph | None:
    """VML placeholder yoki MA'LUMOTNOMA bandidagi paragraf."""
    for p in doc.paragraphs[:8]:
        xml = p._p.xml
        if "v:rect" in xml or "w:pict" in xml or "{{photo}}" in (p.text or ""):
            return p
    for p in doc.paragraphs[:6]:
        t = (p.text or "").upper()
        if "MA" in t and "LUMOT" in t:
            return p
    return doc.paragraphs[0] if doc.paragraphs else None
