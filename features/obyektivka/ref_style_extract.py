"""Extract full style profile from reference Obyektivka DOCX."""

from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Emu

sys.stdout.reconfigure(encoding="utf-8")


def emu_to_pt(v) -> float | None:
    if v is None:
        return None
    return round(int(v) / 12700, 2)


def emu_to_mm(v) -> float | None:
    if v is None:
        return None
    return round(int(v) / 36000, 2)


def run_style(r) -> dict:
    u = r.underline
    color = None
    if r.font.color and r.font.color.rgb:
        color = str(r.font.color.rgb)
    return {
        "text": r.text[:120],
        "bold": r.bold,
        "italic": r.italic,
        "underline": str(u) if u is not None else None,
        "size_pt": emu_to_pt(r.font.size),
        "name": r.font.name,
        "color": color,
    }


def para_style(p) -> dict:
    pf = p.paragraph_format
    tabs = []
    if pf.tab_stops:
        for ts in pf.tab_stops:
            tabs.append({"pos_emu": ts.position, "align": str(ts.alignment)})
    runs = [run_style(r) for r in p.runs if r.text or r._r.xpath(".//w:drawing") or r._r.xpath(".//w:pict")]
    return {
        "style": p.style.name if p.style else None,
        "align": str(p.alignment),
        "space_before_pt": emu_to_pt(pf.space_before),
        "space_after_pt": emu_to_pt(pf.space_after),
        "line_spacing": str(pf.line_spacing),
        "line_rule": str(pf.line_spacing_rule),
        "left_indent_mm": emu_to_mm(pf.left_indent),
        "right_indent_mm": emu_to_mm(pf.right_indent),
        "first_line_indent_mm": emu_to_mm(pf.first_line_indent),
        "tabs": tabs,
        "text": p.text[:200],
        "runs": runs,
    }


def cell_borders(cell) -> dict:
    tc = cell._tc
    tc_pr = tc.tcPr
    out = {"width_dxa": None, "borders": {}}
    if tc_pr is not None:
        tcw = tc_pr.find(qn("w:tcW"))
        if tcw is not None:
            out["width_dxa"] = tcw.get(qn("w:w"))
        tc_b = tc_pr.find(qn("w:tcBorders"))
        if tc_b is not None:
            for side in ("top", "left", "bottom", "right"):
                b = tc_b.find(qn(f"w:{side}"))
                if b is not None:
                    out["borders"][side] = dict(b.attrib)
        tc_mar = tc_pr.find(qn("w:tcMar"))
        if tc_mar is not None:
            out["margins"] = {}
            for side in ("top", "left", "bottom", "right"):
                m = tc_mar.find(qn(f"w:{side}"))
                if m is not None:
                    out["margins"][side] = m.get(qn("w:w"))
    return out


def extract(path: str) -> dict:
    doc = Document(path)
    sec = doc.sections[0]
    profile = {
        "source": path,
        "section": {
            "page_w_mm": emu_to_mm(sec.page_width),
            "page_h_mm": emu_to_mm(sec.page_height),
            "margin_top_mm": emu_to_mm(sec.top_margin),
            "margin_right_mm": emu_to_mm(sec.right_margin),
            "margin_bottom_mm": emu_to_mm(sec.bottom_margin),
            "margin_left_mm": emu_to_mm(sec.left_margin),
        },
        "styles": {},
        "paragraphs": [],
        "tables": [],
    }
    for sname in ("Normal", "Heading 1", "Heading 2", "Heading 6"):
        try:
            s = doc.styles[sname]
            entry = {
                "font_name": s.font.name,
                "font_size_pt": emu_to_pt(s.font.size),
                "bold": s.font.bold,
                "italic": s.font.italic,
            }
            if hasattr(s, "paragraph_format"):
                entry["align"] = str(s.paragraph_format.alignment)
            profile["styles"][sname] = entry
        except KeyError:
            pass

    for i, p in enumerate(doc.paragraphs):
        ps = para_style(p)
        if ps["text"].strip() or ps["runs"]:
            ps["idx"] = i
            profile["paragraphs"].append(ps)

    for ti, t in enumerate(doc.tables):
        tbl = {
            "idx": ti,
            "rows": len(t.rows),
            "cols": len(t.columns),
            "grid_dxa": [],
            "cells_sample": [],
        }
        grid = t._tbl.find(qn("w:tblGrid"))
        if grid is not None:
            tbl["grid_dxa"] = [c.get(qn("w:w")) for c in grid.findall(qn("w:gridCol"))]
        for ri in range(min(len(t.rows), 3)):
            row = []
            for ci, cell in enumerate(t.rows[ri].cells):
                p = cell.paragraphs[0]
                row.append(
                    {
                        "r": ri,
                        "c": ci,
                        "text": cell.text[:80],
                        "para": para_style(p),
                        "cell": cell_borders(cell),
                    }
                )
            tbl["cells_sample"].append(row)
        profile["tables"].append(tbl)
    return profile


if __name__ == "__main__":
    ref = Path("temp/ref_converted.docx")
    if not ref.exists():
        print("Missing", ref)
        sys.exit(1)
    data = extract(str(ref))
    out = Path("temp/ref_style_profile.json")
    out.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    print("written", out, "paras", len(data["paragraphs"]), "tables", len(data["tables"]))
