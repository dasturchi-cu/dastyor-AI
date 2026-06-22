"""
Official Obyektivka (Ma'lumotnoma) DOCX generator.
Layout mirrors «Намуна Объективка (18).doc» (tab-based fields, exact margins).
"""

import base64
import json
import logging
import os
import re
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, Twips

from backend.services.document_render.photo import process_passport_photo
from features.obyektivka.docx_picture import add_floating_picture, add_vml_photo_placeholder
from features.obyektivka.layout import (
    FONT_BODY_PT,
    FONT_FAMILY,
    FONT_REL_TITLE_PT,
    FONT_TITLE_PT,
    IND_DEP_STACK_HANGING_TWIPS,
    IND_DEP_STACK_LEFT_TWIPS,
    IND_HDR_RIGHT_TWIPS,
    IND_INLINE_HANGING_TWIPS,
    IND_INLINE_LEFT_TWIPS,
    IND_JOB_RIGHT_TWIPS,
    IND_REL_COL0_TWIPS,
    IND_STACK_HANGING_TWIPS,
    IND_STACK_LEFT_TWIPS,
    IND_VALUE_HANGING_TWIPS,
    IND_VALUE_LEFT_TWIPS,
    IND_VALUE_RIGHT_TWIPS,
    IND_WORK_HANGING_TWIPS,
    LINE_HEIGHT,
    PAGE_MARGIN_BOTTOM_MM,
    PAGE_MARGIN_LEFT_MM,
    PAGE_MARGIN_RIGHT_MM,
    PAGE_MARGIN_TOP_MM,
    PAGE_HEIGHT_MM,
    PAGE_WIDTH_MM,
    PHOTO_HEIGHT_MM,
    PHOTO_WIDTH_MM,
    REL_COL_DXA,
    TAB_COL_POS,
    TAB_NAME_CENTER_POS,
    TAB_PHOTO_POS,
    TAB_WORK_TITLE_POS,
    TAB_YEAR_POS,
    labels_for,
)

logger = logging.getLogger(__name__)


def _to_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()


def _parse_list(value: Any) -> list[dict[str, Any]]:
    if isinstance(value, list):
        return value
    if isinstance(value, str) and value.strip():
        try:
            parsed = json.loads(value)
            return parsed if isinstance(parsed, list) else []
        except Exception:
            return []
    return []


def _set_run_font(
    run,
    *,
    size: int = FONT_BODY_PT,
    bold: bool | None = None,
    underline: bool = False,
    highlight: bool = False,
) -> None:
    if bold is not None:
        run.bold = bold
    run.font.size = Pt(size)
    run.font.name = FONT_FAMILY
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), FONT_FAMILY)
    if underline:
        run.underline = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def _apply_hdr_right(p) -> None:
    p.paragraph_format.right_indent = Twips(IND_HDR_RIGHT_TWIPS)


def _apply_hanging(p, left_twips: int, hanging_twips: int, right_twips: int | None = None) -> None:
    pf = p.paragraph_format
    pf.left_indent = Twips(left_twips)
    pf.first_line_indent = Twips(-hanging_twips)
    if right_twips is not None:
        pf.right_indent = Twips(right_twips)


def _apply_rel_col0_indent(p) -> None:
    p.paragraph_format.left_indent = Twips(IND_REL_COL0_TWIPS)
    p.paragraph_format.right_indent = Twips(IND_REL_COL0_TWIPS)


def _set_cell_no_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for side in ("top", "left", "bottom", "right"):
        border = tc_borders.find(qn(f"w:{side}"))
        if border is None:
            border = OxmlElement(f"w:{side}")
            tc_borders.append(border)
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")


def _set_table_no_borders_strict(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_style = tbl_pr.find(qn("w:tblStyle"))
    if tbl_style is not None:
        tbl_pr.remove(tbl_style)

    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = tbl_borders.find(qn(f"w:{side}"))
        if border is None:
            border = OxmlElement(f"w:{side}")
            tbl_borders.append(border)
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")

    tbl_look = tbl_pr.find(qn("w:tblLook"))
    if tbl_look is None:
        tbl_look = OxmlElement("w:tblLook")
        tbl_pr.append(tbl_look)
    tbl_look.set(qn("w:firstRow"), "0")
    tbl_look.set(qn("w:lastRow"), "0")
    tbl_look.set(qn("w:firstColumn"), "0")
    tbl_look.set(qn("w:lastColumn"), "0")
    tbl_look.set(qn("w:noHBand"), "1")
    tbl_look.set(qn("w:noVBand"), "1")

    for row in table.rows:
        for cell in row.cells:
            _set_cell_no_borders(cell)


def _set_table_borders(table, size_pt: float = 1.0) -> None:
    size_eighth_pt = str(int(size_pt * 8))
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = tbl_borders.find(qn(f"w:{side}"))
        if border is None:
            border = OxmlElement(f"w:{side}")
            tbl_borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size_eighth_pt)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def _set_col_widths_dxa(table, widths: tuple[int, ...]) -> None:
    tbl = table._tbl
    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for old in list(grid.findall(qn("w:gridCol"))):
        grid.remove(old)
    for w in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        grid.append(gc)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = widths[i]


def _para_tab(p, before: float = 0, after: float = 0) -> None:
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = LINE_HEIGHT
    p.paragraph_format.tab_stops.add_tab_stop(TAB_COL_POS)


def _make_photo_placeholder(hint: str, out_path: str) -> str:
    """3x4 sm ramkali foto joyi (namuna VML pict o'rniga)."""
    from PIL import Image, ImageDraw, ImageFont

    w, h = 354, 472
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, w - 1, h - 1], outline="black", width=2)
    try:
        font = ImageFont.truetype("arial.ttf", 14)
    except OSError:
        font = ImageFont.load_default()
    words = hint.split()
    lines: list[str] = []
    line = ""
    for word in words:
        test = f"{line} {word}".strip()
        if draw.textlength(test, font=font) <= w - 20:
            line = test
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    y = 24
    for ln in lines[:8]:
        tw = draw.textlength(ln, font=font)
        draw.text(((w - tw) / 2, y), ln, fill="black", font=font)
        y += 18
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    img.save(out_path)
    return out_path


def _resolve_photo_path(photo_path: str | None, hint: str) -> tuple[str | None, bool]:
    """Return (path, is_placeholder). Placeholder = VML matn, PNG kerak emas."""
    del hint
    if photo_path and os.path.exists(photo_path):
        return photo_path, False
    return None, True


def _add_header_block(doc: Document, *, title: str, full_name: str, photo_path: str | None, photo_hint: str) -> str | None:
    """Namuna: p0 sarlavha (+ foto agar bor), p1 bo'sh (+ VML hint), p2 FISH."""
    temp_photo: str | None = None
    resolved, is_temp = _resolve_photo_path(photo_path, photo_hint)

    photo_w = Cm(PHOTO_WIDTH_MM / 10)
    photo_h = Cm(PHOTO_HEIGHT_MM / 10)

    try:
        p0 = doc.add_paragraph(style="Heading 6")
    except KeyError:
        p0 = doc.add_paragraph()
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_hdr_right(p0)
    p0.paragraph_format.tab_stops.add_tab_stop(TAB_PHOTO_POS)
    p0.paragraph_format.line_spacing = 1.0
    if not is_temp:
        add_floating_picture(p0, resolved, width=photo_w, height=photo_h)
    tr = p0.add_run(title)
    _set_run_font(tr, size=FONT_TITLE_PT)

    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_hdr_right(p1)
    p1.paragraph_format.tab_stops.add_tab_stop(TAB_PHOTO_POS)
    if is_temp:
        add_vml_photo_placeholder(p1, hint_text=photo_hint)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_hdr_right(p2)
    p2.paragraph_format.tab_stops.add_tab_stop(TAB_NAME_CENTER_POS, WD_TAB_ALIGNMENT.CENTER)
    p2.paragraph_format.tab_stops.add_tab_stop(TAB_PHOTO_POS)
    nr = p2.add_run(full_name)
    _set_run_font(nr, size=FONT_TITLE_PT, bold=True)

    doc.add_paragraph()
    return temp_photo


def _add_two_col_pair(
    doc: Document,
    left_label: str,
    left_value: str,
    right_label: str,
    right_value: str,
    *,
    none: str,
    first_row: bool = False,
) -> None:
    p = doc.add_paragraph()
    _para_tab(p, before=8)
    lr = p.add_run(f"{left_label}:")
    _set_run_font(lr, bold=True)
    p.add_run("\t")
    rr = p.add_run(f"{right_label}:")
    _set_run_font(rr, bold=True)

    p2 = doc.add_paragraph()
    _para_tab(p2)
    if first_row:
        _apply_hanging(p2, IND_VALUE_LEFT_TWIPS, IND_VALUE_HANGING_TWIPS, IND_VALUE_RIGHT_TWIPS)
    else:
        _apply_hanging(p2, IND_VALUE_LEFT_TWIPS, IND_VALUE_HANGING_TWIPS)
    v1 = p2.add_run(left_value or none)
    _set_run_font(v1)
    p2.add_run("\t")
    v2 = p2.add_run(right_value or none)
    _set_run_font(v2)


def _add_inline_field(doc: Document, label: str, value: str, *, none: str) -> None:
    p = doc.add_paragraph()
    _para_tab(p, before=8)
    lr = p.add_run(f"{label}:")
    _set_run_font(lr, bold=True)
    p.add_run("\t")
    _apply_hanging(p, IND_INLINE_LEFT_TWIPS, IND_INLINE_HANGING_TWIPS)
    vr = p.add_run(value or none)
    _set_run_font(vr)


def _add_stacked_field(doc: Document, label: str, value: str, *, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    _para_tab(p, before=8)
    lr = p.add_run(f"{label}:")
    _set_run_font(lr, bold=True)

    if value:
        p2 = doc.add_paragraph()
        _apply_hanging(p2, IND_STACK_LEFT_TWIPS, IND_STACK_HANGING_TWIPS)
        vr = p2.add_run(value)
        _set_run_font(vr, highlight=highlight)


def _add_split_label_field(doc: Document, line1: str, line2: str, value: str) -> None:
    p1 = doc.add_paragraph()
    _para_tab(p1, before=8)
    r1 = p1.add_run(line1)
    _set_run_font(r1, bold=True)

    p2 = doc.add_paragraph()
    _para_tab(p2)
    r2 = p2.add_run(f"{line2}:")
    _set_run_font(r2, bold=True)

    if value:
        p3 = doc.add_paragraph()
        _apply_hanging(p3, IND_DEP_STACK_LEFT_TWIPS, IND_DEP_STACK_HANGING_TWIPS)
        vr = p3.add_run(value)
        _set_run_font(vr)


def generate_obyektivka_docx(
    user_data: dict[str, Any] | None = None,
    photo_path: str | None = None,
    output_filepath: str | None = None,
    **kwargs: Any,
) -> str:
    data = user_data or kwargs.get("data") or {}
    temp_photo_from_data = None
    if not photo_path or not os.path.exists(photo_path):
        photo_data = _to_text(data.get("photo_data"))
        try:
            if photo_data.startswith("data:image/") and "," in photo_data:
                photo_data = process_passport_photo(photo_data)
                header, b64 = photo_data.split(",", 1)
                mime = header.split(";")[0].split(":")[1].lower()
                ext = {
                    "image/png": "png",
                    "image/jpeg": "jpg",
                    "image/jpg": "jpg",
                    "image/webp": "webp",
                }.get(mime, "jpg")
                os.makedirs("temp", exist_ok=True)
                temp_photo_from_data = os.path.join("temp", f"oby_local_photo_{os.getpid()}.{ext}")
                with open(temp_photo_from_data, "wb") as f:
                    f.write(base64.b64decode(b64))
                photo_path = temp_photo_from_data
        except Exception as exc:
            logger.warning("Failed to decode photo_data in generator: %s", exc)
            photo_path = None

    if not isinstance(data, dict):
        raise ValueError("user_data must be a dictionary")

    if not output_filepath:
        os.makedirs("temp", exist_ok=True)
        safe_name = (_to_text(data.get("fullname")) or "Obyektivka").replace(" ", "_").replace("/", "_")
        output_filepath = os.path.join("temp", f"obyektivka_{safe_name}_@DastyorAiBot.docx")
    else:
        os.makedirs(os.path.dirname(output_filepath) or ".", exist_ok=True)

    lang = _to_text(data.get("lang")) or "uz_lat"
    L = labels_for(lang)
    none = L["none"]

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_FAMILY
    style.font.size = Pt(FONT_BODY_PT)
    style.paragraph_format.line_spacing = LINE_HEIGHT
    style.paragraph_format.space_after = Pt(0)
    for sname in ("Heading 2", "Heading 6"):
        try:
            hs = doc.styles[sname]
            hs.font.name = FONT_FAMILY
        except KeyError:
            pass
    for section in doc.sections:
        section.page_width = Cm(PAGE_WIDTH_MM / 10)
        section.page_height = Cm(PAGE_HEIGHT_MM / 10)
        section.top_margin = Cm(PAGE_MARGIN_TOP_MM / 10)
        section.bottom_margin = Cm(PAGE_MARGIN_BOTTOM_MM / 10)
        section.left_margin = Cm(PAGE_MARGIN_LEFT_MM / 10)
        section.right_margin = Cm(PAGE_MARGIN_RIGHT_MM / 10)

    full_name = _to_text(data.get("fullname")) or "FAMILIYA ISM SHARIF"
    work_items = _parse_list(data.get("work_experience"))

    current_job = _to_text(data.get("current_job"))
    current_job_year = _to_text(data.get("current_job_year"))
    if not current_job:
        for idx, item in enumerate(work_items):
            year_raw = _to_text(item.get("year") or item.get("from"))
            year_norm = re.sub(r"[\s.\-_/]", "", year_raw.lower())
            is_current = any(
                key in year_norm
                for key in ("hv", "hvgacha", "hozirgacha", "ҳв", "ҳвгача", "ҳозиргача")
            )
            position_raw = _to_text(item.get("position") or item.get("description") or item.get("job"))
            if is_current and position_raw:
                current_job = position_raw
                if not current_job_year:
                    from_raw = _to_text(item.get("from"))
                    if from_raw:
                        current_job_year = from_raw
                    else:
                        match = re.search(r"(19|20)\d{2}", year_raw)
                        if match:
                            current_job_year = match.group(0)
                work_items.pop(idx)
                break

    section = doc.sections[0]
    temp_header_photo: str | None = None

    temp_header_photo = _add_header_block(
        doc,
        title=L["title"],
        full_name=full_name,
        photo_path=photo_path,
        photo_hint=L["photo_hint"],
    )

    if current_job:
        if current_job_year:
            yr_p = doc.add_paragraph()
            yr_p.paragraph_format.tab_stops.add_tab_stop(TAB_YEAR_POS)
            yr_run = yr_p.add_run(f"{current_job_year.rstrip('.')}:")
            _set_run_font(yr_run)
        try:
            job_p = doc.add_paragraph(style="Heading 2")
        except KeyError:
            job_p = doc.add_paragraph()
        job_p.paragraph_format.right_indent = Twips(IND_JOB_RIGHT_TWIPS)
        job_p.paragraph_format.space_after = Pt(4)
        pos_run = job_p.add_run(current_job)
        _set_run_font(pos_run, bold=True)

    _add_two_col_pair(
        doc,
        L["r1l"],
        _to_text(data.get("birthdate")),
        L["r1r"],
        _to_text(data.get("birthplace")),
        none=none,
        first_row=True,
    )
    _add_two_col_pair(
        doc,
        L["r2l"],
        _to_text(data.get("nation")),
        L["r2r"],
        _to_text(data.get("party")),
        none=none,
    )
    _add_two_col_pair(
        doc,
        L["r3l"],
        _to_text(data.get("education")),
        L["r3r"],
        _to_text(data.get("graduated")),
        none=none,
    )
    _add_inline_field(doc, L["rSpec"], _to_text(data.get("specialty")), none=none)
    _add_two_col_pair(
        doc,
        L["r4l"],
        _to_text(data.get("degree")),
        L["r4r"],
        _to_text(data.get("scientific_title")),
        none=none,
    )
    _add_two_col_pair(
        doc,
        L["r5l"],
        _to_text(data.get("languages")),
        L["r5r"],
        _to_text(data.get("military_rank")),
        none=none,
    )
    _add_stacked_field(doc, L["rAw"], _to_text(data.get("awards")))
    _add_stacked_field(doc, L["rIdo"], _to_text(data.get("departmental_awards")), highlight=True)
    _add_split_label_field(doc, L["rDep1"], L["rDep2"], _to_text(data.get("deputy")))

    work_title = doc.add_paragraph()
    work_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    work_title.paragraph_format.tab_stops.add_tab_stop(TAB_WORK_TITLE_POS)
    work_title.paragraph_format.space_before = Pt(0)
    work_title.paragraph_format.space_after = Pt(4)
    _set_run_font(work_title.add_run(), size=FONT_BODY_PT)
    wr = work_title.add_run(L["work"])
    _set_run_font(wr, size=FONT_TITLE_PT, bold=True)

    yy_suffix = "йй." if lang == "uz_cyr" else "yy."
    if work_items:
        for item in work_items:
            year = _to_text(item.get("year")) or _to_text(item.get("from"))
            end = _to_text(item.get("to"))
            if year and end and yy_suffix.rstrip(".") not in year.lower():
                year = f"{year}-{end} {yy_suffix}"
            year = year.rstrip(".") if year else year
            position = _to_text(item.get("position") or item.get("description") or item.get("job"))
            if not (year or position):
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            _apply_hanging(p, IND_WORK_HANGING_TWIPS, IND_WORK_HANGING_TWIPS)
            if year and position:
                yshow = year if (yy_suffix.rstrip(".") in year.lower()) else f"{year} {yy_suffix}"
                line = f"{yshow} - {position}"
            else:
                line = year or position
            wrun = p.add_run(line)
            _set_run_font(wrun)
    else:
        p = doc.add_paragraph("-")
        _set_run_font(p.runs[0] if p.runs else p.add_run("-"))

    relatives = _parse_list(data.get("relatives"))
    doc.add_page_break()

    try:
        t1 = doc.add_paragraph(style="Heading 1")
    except KeyError:
        t1 = doc.add_paragraph()
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = t1.add_run(f"{full_name}{L['rel_line1_suffix']}")
    _set_run_font(r1, size=FONT_REL_TITLE_PT, bold=True)
    t1.paragraph_format.space_after = Pt(0)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run(L["rel_line2"])
    _set_run_font(r2, size=FONT_REL_TITLE_PT, bold=True)
    t2.paragraph_format.space_after = Pt(6)

    rel_tbl = doc.add_table(rows=1, cols=5)
    rel_tbl.autofit = False
    _set_col_widths_dxa(rel_tbl, REL_COL_DXA)
    _set_table_borders(rel_tbl, size_pt=1.0)

    header_cells = rel_tbl.rows[0].cells
    # col0
    c0 = header_cells[0]
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_rel_col0_indent(p0)
    if lang == "uz_cyr":
        for part in ("Қариндош", "-", "лиги"):
            _set_run_font(p0.add_run(part), bold=True)
    else:
        _set_run_font(p0.add_run(L["qar"]), bold=True)
    # col1-4 headers (namuna: qator bo'lib, underline yo'q)
    hdr_texts = [
        None,
        "Familiyasi, ismi \nva otasining ismi " if lang != "uz_cyr" else "Фамилияси, исми \nва отасининг исми ",
        "Tug'ilgan yili \nva joyi" if lang != "uz_cyr" else "Туғилган йили \nва жойи",
        L["ish"],
        L["tur"],
    ]
    for i in range(1, 5):
        cell = header_cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(p.add_run(hdr_texts[i]), bold=True)

    if relatives:
        for rel in relatives:
            row = rel_tbl.add_row().cells
            vals = [
                _to_text(rel.get("degree") or rel.get("type")),
                _to_text(rel.get("fullname") or rel.get("name")),
                _to_text(rel.get("birth_year_place") or rel.get("birth")),
                _to_text(rel.get("work_place") or rel.get("job")),
                _to_text(rel.get("address") or rel.get("addr")),
            ]
            for i, val in enumerate(vals):
                cell = row[i]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if i == 0:
                    _apply_rel_col0_indent(p)
                rr = p.add_run(val or none)
                _set_run_font(rr, bold=(i == 0))
    else:
        row = rel_tbl.add_row().cells
        merged = row[0].merge(row[4])
        p = merged.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(L["no_rel"])
        _set_run_font(run)

    doc.save(output_filepath)
    for tmp in (temp_photo_from_data, temp_header_photo):
        if tmp and os.path.exists(tmp):
            try:
                os.remove(tmp)
            except Exception:
                pass
    return output_filepath
