import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout.reconfigure(encoding="utf-8")

doc = Document("temp/ref_converted.docx")
t = doc.tables[0]
print("=== REL TABLE HEADER ===")
for i, cell in enumerate(t.rows[0].cells):
    p = cell.paragraphs[0]
    print(f"col{i}: align={p.alignment} text={repr(cell.text[:50])}")
    for r in p.runs:
        print(f"  run u={r.underline} b={r.bold} sz={r.font.size}")

print("\n=== REL TABLE BODY row1 ===")
for i, cell in enumerate(t.rows[1].cells):
    p = cell.paragraphs[0]
    print(f"col{i}: align={p.alignment} text={repr(cell.text[:40])}")

print("\n=== CURRENT JOB P4-P5 ===")
for i in [4, 5]:
    p = doc.paragraphs[i]
    print(f"P{i} style={p.style.name} align={p.alignment}")
    for r in p.runs:
        print(f"  bold={r.bold} u={r.underline} sz={r.font.size} text={repr(r.text[:50])}")
