import sys
from docx import Document

sys.stdout.reconfigure(encoding="utf-8")

ref = Document("temp/ref_converted.docx")
gen = Document("temp/test_v2.docx")

print("REF tables:", len(ref.tables), "| GEN tables:", len(gen.tables))
print("\nHEADER STRUCTURE:")
for label, doc in [("REF", ref), ("GEN", gen)]:
    print(f"--- {label} ---")
    for i in range(6):
        p = doc.paragraphs[i]
        bold = [r.bold for r in p.runs if r.text.strip()]
        print(f"  P{i}: {repr(p.text[:55])} align={p.alignment} bold={bold}")

print("\nREL COLS:")
from docx.oxml.ns import qn

for label, doc in [("REF", ref), ("GEN", gen)]:
    t = doc.tables[0]
    grid = t._tbl.find(qn("w:tblGrid"))
    cols = [c.get(qn("w:w")) for c in grid.findall(qn("w:gridCol"))]
    print(label, cols)
