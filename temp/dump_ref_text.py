import sys
from docx import Document

sys.stdout.reconfigure(encoding="utf-8")

d = Document("temp/ref_converted.docx")
for i, p in enumerate(d.paragraphs):
    t = p.text.replace("\n", " | ")
    if t.strip():
        print(f"P{i:02d} [{p.style.name if p.style else ''}]: {t[:120]}")

print("---TABLE---")
t = d.tables[0]
for ri, row in enumerate(t.rows):
    cells = [c.text.replace("\n", " ")[:40] for c in row.cells]
    print(f"R{ri:02d}", cells)
