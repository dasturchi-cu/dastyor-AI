"""Audit reference vs generated obyektivka header structure."""
import json
import sys
import zipfile
from pathlib import Path

from docx import Document

sys.stdout.reconfigure(encoding="utf-8")

REF = Path("temp/ref_converted.docx")
GEN = Path("temp/audit_gen.docx")


def para_info(doc, limit=12):
    out = []
    for i, p in enumerate(doc.paragraphs[:limit]):
        has_pict = bool(p._p.xpath(".//w:pict"))
        has_drawing = bool(p._p.xpath(".//w:drawing"))
        out.append(
            {
                "i": i,
                "style": p.style.name if p.style else None,
                "align": str(p.alignment),
                "text": p.text[:80],
                "pict": has_pict,
                "drawing": has_drawing,
                "right_mm": round(int(p.paragraph_format.right_indent or 0) / 36000, 2)
                if p.paragraph_format.right_indent
                else 0,
            }
        )
    return out


def main():
    if not GEN.exists():
        from features.obyektivka.docx_official import generate_obyektivka_docx

        generate_obyektivka_docx(
            {
                "fullname": "Aliyev Vali Valiyevich",
                "lang": "uz_lat",
                "birthdate": "01.01.1990",
                "birthplace": "Toshkent",
                "nation": "O'zbek",
                "party": "yo'q",
                "education": "Oliy",
                "graduated": "TATU",
                "specialty": "IT",
                "degree": "yo'q",
                "scientific_title": "yo'q",
                "languages": "yo'q",
                "military_rank": "yo'q",
                "work_experience": [{"year": "2020", "position": "Dasturchi"}],
            },
            output_filepath=str(GEN),
        )

    ref = Document(str(REF))
    gen = Document(str(GEN))
    print("=== REF paragraphs ===")
    print(json.dumps(para_info(ref), ensure_ascii=False, indent=2))
    print("=== GEN paragraphs ===")
    print(json.dumps(para_info(gen), ensure_ascii=False, indent=2))

    ref_xml = zipfile.ZipFile(REF).read("word/document.xml").decode()
    gen_xml = zipfile.ZipFile(GEN).read("word/document.xml").decode()
    print("ref pict", ref_xml.count("w:pict"), "gen pict", gen_xml.count("w:pict"))
    print("ref anchor", ref_xml.count("wp:anchor"), "gen anchor", gen_xml.count("wp:anchor"))


if __name__ == "__main__":
    main()
