import sys
from docx import Document
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding="utf-8")

doc = Document("temp/ref_converted.docx")
body = doc.element.body

print("=== HEADER PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs[:10]):
    pf = p.paragraph_format
    tabs = [(ts.position, str(ts.alignment)) for ts in pf.tab_stops] if pf.tab_stops else []
    print(f"P{i}: align={p.alignment} sb={pf.space_before} sa={pf.space_after}")
    print(f"  text: {repr(p.text[:80])}")
    print(f"  tabs: {tabs}")
    for ri, r in enumerate(p.runs):
        has_pic = bool(r._r.xpath(".//w:drawing") or r._r.xpath(".//w:pict"))
        if has_pic or r.text.strip():
            print(f"  run{ri}: {repr(r.text[:40])} pic={has_pic} bold={r.bold} u={r.underline}")

print("\n=== DRAWING/POSITION ===")
for el in body.iter():
    tag = el.tag.split("}")[-1]
    if tag in ("anchor", "inline", "pict"):
        print("TAG", tag)
        for c in el.iter():
            t = c.tag.split("}")[-1]
            if t in ("posOffset", "align", "relativeFrom", "cx", "cy", "simplePos"):
                print(f"  {t}: {dict(c.attrib)}")

print("\n=== TABLES", len(doc.tables))
