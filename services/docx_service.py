from __future__ import annotations

import os
from typing import Iterable


def lines_to_docx(lines: Iterable[str], output_path: str) -> str:
    """
    Create a DOCX where each input line becomes a new paragraph.
    Returns the saved path.
    """
    from docx import Document

    out_path = str(output_path)
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)

    doc = Document()
    for raw in lines:
        line = str(raw or "").strip()
        if not line:
            # Preserve spacing intentionally: blank line => blank paragraph
            doc.add_paragraph("")
            continue
        doc.add_paragraph(line)
    doc.save(out_path)
    return out_path

