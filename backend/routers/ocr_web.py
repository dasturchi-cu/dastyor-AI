"""OCR and PDF merge endpoints for the WebApp (Paddle + Gemini paths)."""
from __future__ import annotations

import asyncio
import io
import logging
import os
import time
from typing import List, Optional

from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, UploadFile
from fastapi.responses import JSONResponse, StreamingResponse
from telegram import InputFile

from backend.dependencies import get_ptb_application
from backend.services.html_text_normalize import html_ocr_to_plain
from backend.services.paddle_ocr_runtime import (
    docx_image_then_text,
    get_ocr_thread_pool,
    paddle_extract_plain_text,
)
from backend.services.paddle_ocr_runtime import resize_for_ocr as paddle_resize
from backend.services.temp_files import safe_remove, temp_image_path
from backend.services.upload_io import EmptyUploadError, UploadTooLargeError, read_upload_limited
from backend.services.user_resolve import resolve_telegram_uid, safe_filename_part
from backend.services.web_quota import web_quota_after, web_quota_before
from backend.settings import get_settings
from bot.utils.delivery import send_docx_with_confirmation

logger = logging.getLogger(__name__)

router = APIRouter(tags=["web-ocr"])


@router.post("/api/ocr_extract")
async def api_ocr_extract(file: UploadFile = File(...)):
    try:
        raw = await read_upload_limited(file)
    except EmptyUploadError:
        raise HTTPException(status_code=400, detail="Fayl bo'sh")
    except UploadTooLargeError:
        raise HTTPException(status_code=400, detail="Fayl juda katta")

    with temp_image_path(suffix=safe_filename_part(file.filename or "", ".jpg")) as img_path:
        try:
            with open(img_path, "wb") as f:
                f.write(raw)
        except OSError as e:
            raise HTTPException(status_code=500, detail=f"Fayl saqlashda xato: {e}")

        try:
            from bot.services.ocr_service import extract_text_from_image

            html_text = await extract_text_from_image(img_path)
        except Exception as e:
            logger.error("api_ocr_extract OCR xatosi: %s", e, exc_info=True)
            raise HTTPException(status_code=502, detail=f"OCR xatosi: {str(e)[:200]}")

    if not html_text or not html_text.strip():
        raise HTTPException(status_code=422, detail="Rasmdan matn ajratib bo'lmadi")

    plain = html_ocr_to_plain(html_text)
    return {"ok": True, "text": plain, "html": html_text}


@router.post("/api/ocr_extract_docx")
async def api_ocr_extract_docx(file: UploadFile = File(...)):
    """
    Bot bilan bir xil: Gemini → HTML → python-docx (1:1 layout).
    """
    try:
        raw = await read_upload_limited(file)
    except EmptyUploadError:
        raise HTTPException(status_code=400, detail="Fayl bo'sh")
    except UploadTooLargeError:
        raise HTTPException(status_code=400, detail="Fayl juda katta")

    with temp_image_path(suffix=safe_filename_part(file.filename or "", ".jpg")) as img_path:
        try:
            with open(img_path, "wb") as f:
                f.write(raw)
        except OSError as e:
            raise HTTPException(status_code=500, detail=f"Fayl saqlashda xato: {e}")

        try:
            from bot.services.ocr_service import extract_text_from_image

            html_text = await extract_text_from_image(img_path)
        except Exception as e:
            logger.error("api_ocr_extract_docx OCR xatosi: %s", e, exc_info=True)
            raise HTTPException(status_code=502, detail=f"OCR xatosi: {str(e)[:200]}")

    if not html_text or not html_text.strip():
        raise HTTPException(status_code=422, detail="Rasmdan matn ajratib bo'lmadi")

    try:
        from docx import Document

        from bot.handlers.ocr_to_word import add_html_to_docx

        loop = asyncio.get_running_loop()

        def _build():
            doc = Document()
            add_html_to_docx(doc, html_text)
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf.read()

        docx_bytes = await loop.run_in_executor(None, _build)
    except Exception as e:
        logger.error("api_ocr_extract_docx DOCX xatosi: %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail=f"Word yaratishda xato: {str(e)[:200]}")

    ts = int(time.time())
    fname = f"OCR_1to1_{ts}.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


@router.post("/ocr")
async def ocr_paddle(
    file: UploadFile = File(...),
    layout: Optional[str] = Query(None, description="1/true=yes — bbox va qatorlar bilan JSON"),
):
    try:
        from backend.services.paddle_ocr_runtime import (
            bytes_to_bgr,
            ensure_paddle_imports,
            paddle_extract_structured,
        )

        ensure_paddle_imports()
    except ImportError:
        raise HTTPException(status_code=500, detail="PaddleOCR deps missing on server")

    try:
        raw = await read_upload_limited(file)
    except EmptyUploadError:
        raise HTTPException(status_code=400, detail="Empty file")
    except UploadTooLargeError:
        raise HTTPException(status_code=400, detail="File too large")

    from backend.services.paddle_ocr_runtime import get_paddle_engine

    get_paddle_engine()
    img = bytes_to_bgr(raw)
    if img is None:
        raise HTTPException(status_code=400, detail="Invalid image")

    img = paddle_resize(img)
    loop = asyncio.get_running_loop()
    pool = get_ocr_thread_pool()
    want_layout = (layout or "").strip().lower() in ("1", "true", "yes", "full")

    def _run():
        if want_layout:
            return paddle_extract_structured(img)
        return {"text": paddle_extract_plain_text(img), "lines": [], "preprocess": ""}

    data = await loop.run_in_executor(pool, _run)
    text = (data.get("text") or "").strip()
    out: dict = {"text": text}
    if want_layout:
        out["lines"] = data.get("lines") or []
        out["preprocess"] = data.get("preprocess") or "none"
    return out


@router.post("/ocr-word")
async def ocr_word_paddle(file: UploadFile = File(...)):
    try:
        from backend.services.paddle_ocr_runtime import bytes_to_bgr, ensure_paddle_imports

        ensure_paddle_imports()
    except ImportError:
        raise HTTPException(status_code=500, detail="PaddleOCR deps missing on server")

    try:
        raw = await read_upload_limited(file)
    except EmptyUploadError:
        raise HTTPException(status_code=400, detail="Empty file")
    except UploadTooLargeError:
        raise HTTPException(status_code=400, detail="File too large")

    from backend.services.paddle_ocr_runtime import get_paddle_engine

    get_paddle_engine()
    img = bytes_to_bgr(raw)
    if img is None:
        raise HTTPException(status_code=400, detail="Invalid image")

    img = paddle_resize(img)
    loop = asyncio.get_running_loop()
    pool = get_ocr_thread_pool()
    try:
        text = await loop.run_in_executor(pool, paddle_extract_plain_text, img)
    except Exception as ocr_err:
        logger.warning("Paddle OCR failed, returning image-only DOCX: %s", ocr_err)
        text = ""
    docx_bytes = await loop.run_in_executor(pool, docx_image_then_text, raw, text)
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="ocr_result.docx"'},
    )


@router.post("/api/ocr_direct")
async def api_ocr_direct(
    file: UploadFile = File(...),
    telegram_id: Optional[str] = Form(None),
    mode: Optional[str] = Form("docx"),
    ptb=Depends(get_ptb_application),
):
    ts = int(time.time())
    safe_upload_name = safe_filename_part(file.filename or "", "img.jpg")

    try:
        raw = await read_upload_limited(file)
    except EmptyUploadError:
        raise HTTPException(status_code=400, detail="Fayl bo'sh")
    except UploadTooLargeError:
        raise HTTPException(status_code=400, detail="Fayl juda katta")

    with temp_image_path(suffix=os.path.splitext(safe_upload_name)[1] or ".jpg") as img_path:
        try:
            with open(img_path, "wb") as f:
                f.write(raw)
        except OSError as e:
            raise HTTPException(status_code=500, detail=f"Fayl saqlashda xato: {e}")

        if (mode or "").strip().lower() in ("image_docx", "image", "imgdocx", "scan_image"):
            try:
                from docx import Document
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.shared import Inches

                def build_image_docx() -> bytes:
                    doc = Document()
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(6.5))
                    buf = io.BytesIO()
                    doc.save(buf)
                    buf.seek(0)
                    return buf.read()

                loop = asyncio.get_running_loop()
                docx_bytes = await loop.run_in_executor(None, build_image_docx)
            except Exception as e:
                logger.error("Image DOCX build error: %s", e, exc_info=True)
                raise HTTPException(status_code=500, detail=f"Word hujjat yaratishda xato: {str(e)[:200]}")

            if telegram_id and telegram_id.strip().isdigit():
                chat_id = int(telegram_id)

                async def send_to_telegram():
                    try:
                        buf = io.BytesIO(docx_bytes)
                        buf.name = f"OCR_Image_{ts}.docx"
                        await send_docx_with_confirmation(
                            ptb.bot,
                            chat_id,
                            buf,
                            filename=buf.name,
                            caption="✅ Rasm Word faylga joylandi (1:1 ko'rinish).",
                        )
                    except Exception as tg_err:
                        logger.warning("Telegram send failed (non-fatal): %s", tg_err)

                asyncio.create_task(send_to_telegram())

            filename = f"DASTYOR_IMAGE_{ts}_@DastyorAiBot.docx"
            return StreamingResponse(
                io.BytesIO(docx_bytes),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{filename}"'},
            )

        ocr_start = time.perf_counter()
        try:
            from bot.services.ocr_service import extract_text_from_image

            html_text = await extract_text_from_image(img_path)
            logger.info(
                "api_ocr_direct OCR done in %.1fs path=%s",
                time.perf_counter() - ocr_start,
                img_path,
            )
        except Exception as e:
            logger.error(
                "api_ocr_direct OCR error after %.1fs: %s",
                time.perf_counter() - ocr_start,
                e,
                exc_info=True,
            )
            raise HTTPException(status_code=502, detail=f"OCR xatosi: {str(e)[:200]}")

    if not html_text or not html_text.strip():
        raise HTTPException(status_code=422, detail="Rasmdan matn ajratib bo'lmadi. Aniqroq rasm yuboring.")

    if (mode or "").strip().lower() in ("text", "txt", "preview"):
        plain = html_ocr_to_plain(html_text)
        return {"ok": True, "text": plain, "html": html_text}

    try:
        from bot.handlers.ocr_to_word import add_html_to_docx
        from docx import Document

        def build_docx() -> bytes:
            doc = Document()
            doc.add_heading("OCR Natijasi", 0)
            try:
                add_html_to_docx(doc, html_text)
            except Exception as parse_err:
                logger.warning("HTML parse fallback: %s", parse_err)
                doc.add_paragraph(html_text)
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf.read()

        loop = asyncio.get_running_loop()
        docx_bytes = await loop.run_in_executor(None, build_docx)
    except Exception as e:
        logger.error("DOCX build error: %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail=f"Word hujjat yaratishda xato: {str(e)[:200]}")

    if telegram_id and telegram_id.strip().isdigit():
        chat_id = int(telegram_id)

        async def send_to_telegram():
            try:
                buf = io.BytesIO(docx_bytes)
                buf.name = f"OCR_Natija_{ts}.docx"
                await send_docx_with_confirmation(
                    ptb.bot,
                    chat_id,
                    buf,
                    filename=buf.name,
                    caption="✅ Rasm Word ga aylantirildi!\n📎 Fayl tayyor.",
                )
            except Exception as tg_err:
                logger.warning("Telegram send failed (non-fatal): %s", tg_err)

        asyncio.create_task(send_to_telegram())

    filename = f"DASTYOR_OCR_{ts}_@DastyorAiBot.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/api/upload_ocr")
async def api_upload_ocr(
    telegram_id: str = Form(...),
    files: List[UploadFile] = File(...),
    ptb=Depends(get_ptb_application),
):
    try:
        chat_id = int(telegram_id)
        legacy_ts = int(time.time())
        os.makedirs("temp", exist_ok=True)
        img_paths: list[str] = []
        settings = get_settings()
        for f in files:
            raw = await f.read()
            if len(raw) > settings.max_upload_bytes:
                raise HTTPException(status_code=400, detail="Fayl juda katta")
            path = f"temp/legacy_ocr_{legacy_ts}_{safe_filename_part(f.filename or 'img', 'img.jpg')}"
            with open(path, "wb") as fh:
                fh.write(raw)
            img_paths.append(path)

        async def run_legacy_ocr():
            from bot.handlers.ocr_to_word import add_html_to_docx
            from bot.services.ocr_service import extract_text_from_image
            from docx import Document

            try:
                msg = await ptb.bot.send_message(
                    chat_id=chat_id,
                    text="⏳ Rasm qabul qilindi. OCR qilinmoqda...",
                )
                for img_path in img_paths:
                    html_text = await extract_text_from_image(img_path)
                    if not html_text:
                        continue
                    ts2 = int(time.time())
                    docx_path = f"temp/legacy_{chat_id}_{ts2}.docx"
                    doc = Document()
                    doc.add_heading("OCR Natijasi", 0)
                    add_html_to_docx(doc, html_text)
                    doc.save(docx_path)
                    with open(docx_path, "rb") as df:
                        await send_docx_with_confirmation(
                            ptb.bot,
                            chat_id,
                            df,
                            filename=f"OCR_{ts2}.docx",
                            caption="✅ Word fayl tayyor!",
                        )
                    safe_remove(docx_path)
                await ptb.bot.delete_message(chat_id=chat_id, message_id=msg.message_id)
            except Exception as ex:
                logger.error("Legacy OCR error: %s", ex, exc_info=True)
            finally:
                for p in img_paths:
                    safe_remove(p)

        asyncio.create_task(run_legacy_ocr())
        return {"status": "ok"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Upload OCR API error: %s", e, exc_info=True)
        return {"error": str(e)}


@router.post("/api/pdf_direct")
async def api_pdf_direct(
    files: List[UploadFile] = File(...),
    telegram_id: Optional[str] = Form(None),
    token: Optional[str] = Form(None),
    send_only: Optional[str] = Form(None),
    ptb=Depends(get_ptb_application),
):
    try:
        logger.info("[PDF API] Boshlandi. telegram_id=%s, files_count=%s", telegram_id, len(files))
        uid_pdf = resolve_telegram_uid(telegram_id, token)
        uid_int = int(uid_pdf) if uid_pdf else None
        if uid_int:
            from bot.services.plan_limits import CAT_IMAGE_PDF

            web_quota_before(uid_int, CAT_IMAGE_PDF)
        os.makedirs("temp", exist_ok=True)
        img_paths: list[str] = []
        ts = int(time.time())
        settings = get_settings()
        n_files = len(files)
        want_send_only = str(send_only or "").lower() in ("1", "true", "yes", "on")
        try:

            async def _save_one(i: int, file: UploadFile) -> str:
                if not file.filename:
                    raise ValueError("Noto'g'ri fayl nomi")
                ext = os.path.splitext(file.filename)[1] or ".jpg"
                path = f"temp/pdf_req_{ts}_{i}{ext}"
                content = await file.read()
                if not content:
                    raise ValueError(f"Bo'sh fayl: {file.filename}")
                if len(content) > settings.max_upload_bytes:
                    raise ValueError(f"Fayl juda katta: {file.filename}")
                with open(path, "wb") as f:
                    f.write(content)
                return path

            img_paths = await asyncio.gather(*[_save_one(i, f) for i, f in enumerate(files)])
        except ValueError as ve:
            raise HTTPException(status_code=400, detail=str(ve)) from ve
        except HTTPException:
            raise
        except Exception as e:
            logger.error("[PDF API] Rasm yuklashda xatolik: %s", e, exc_info=True)
            raise HTTPException(status_code=500, detail=f"Rasm yuklash xatosi: {e}")

        if not img_paths:
            raise HTTPException(status_code=400, detail="Fayl yuklanmadi")

        pdf_path = f"temp/merged_{ts}.pdf"
        try:
            from bot.services.pdf_service import images_to_pdf

            loop = asyncio.get_running_loop()
            await loop.run_in_executor(None, images_to_pdf, img_paths, pdf_path)
        except Exception as build_err:
            logger.error("[PDF API] PDF yaralishda xato: %s", build_err, exc_info=True)
            safe_remove(*img_paths)
            raise HTTPException(status_code=500, detail=f"PDF yaratishda xato: {build_err}")

        if not os.path.exists(pdf_path):
            safe_remove(*img_paths)
            raise HTTPException(status_code=500, detail="PDF fayl yaratilmadi")

        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()

        safe_remove(pdf_path, *img_paths)

        if uid_int:
            from bot.services.plan_limits import CAT_IMAGE_PDF

            web_quota_after(uid_int, CAT_IMAGE_PDF, "Web Rasm→PDF")

        tid_ok = telegram_id and telegram_id.strip().isdigit()
        if tid_ok:
            chat_id = int(telegram_id)
            n_merge = n_files

            async def send_pdf_to_telegram():
                try:
                    buf = io.BytesIO(pdf_bytes)
                    buf.name = f"DASTYOR_AI_Rasmlar_{ts}_@DastyorAiBot.pdf"
                    await ptb.bot.send_document(
                        chat_id=chat_id,
                        document=InputFile(buf, filename=buf.name),
                        caption=(
                            "✅ PDF tayyor va botga yuborildi\n"
                            f"📄 {n_merge} ta rasm birlashtirildi."
                        ),
                    )
                    try:
                        from bot.services.supabase_db import db_insert_action_log

                        db_insert_action_log(chat_id, "pdf", buf.name or "merged.pdf")
                    except Exception:
                        pass
                except Exception as tg_err:
                    logger.error("[PDF API Background] Telegram xatosi: %s", tg_err, exc_info=True)

            asyncio.create_task(send_pdf_to_telegram())

        if want_send_only and tid_ok:
            return JSONResponse(
                content={
                    "ok": True,
                    "message": "PDF tayyor va botga yuborildi.",
                    "files": n_files,
                }
            )

        filename = f"DASTYOR_AI_Rasmlar_{ts}_@DastyorAiBot.pdf"
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error("[PDF API] Umumiy xatolik: %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))
