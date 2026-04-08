"""
Shared synchronous PaddleOCR pipeline for:
- FastAPI handlers (via asyncio.run_in_executor + shared pool)
- Celery workers (same code path = no drift)

GPU: set PADDLE_OCR_USE_GPU=1 if Paddle build supports CUDA (optional).
"""
from __future__ import annotations

import base64
import io
import inspect
import logging
import os
import threading
from concurrent.futures import ThreadPoolExecutor
from typing import Any

os.environ.setdefault("FLAGS_enable_pir_api", "0")
os.environ.setdefault("FLAGS_use_new_executor", "0")
os.environ.setdefault("FLAGS_enable_onednn", "0")
os.environ.setdefault("FLAGS_use_mkldnn", "0")
os.environ.setdefault("FLAGS_enable_pir_in_executor", "0")

logger = logging.getLogger("dastyor.paddle_ocr")

_OCR_POOL = ThreadPoolExecutor(max_workers=int(os.getenv("OCR_WORKERS", "4")), thread_name_prefix="paddle")
_OCR_ENGINE = None
_OCR_LOCK = threading.Lock()
_OCR_PROFILE = "default"
_WARMUP_FUTURE = None
_WARMUP_LOCK = threading.Lock()


def get_ocr_thread_pool() -> ThreadPoolExecutor:
    return _OCR_POOL


def _runtime_executor_error(exc: Exception) -> bool:
    msg = str(exc)
    return "ConvertPirAttribute2RuntimeAttribute" in msg or "onednn_instruction.cc" in msg


def _filter_paddle_kwargs(kwargs: dict[str, Any]) -> dict[str, Any]:
    try:
        from paddleocr import PaddleOCR

        sig = inspect.signature(PaddleOCR)
        # PaddleOCR v3+ often exposes a long explicit signature and forwards extra options via **kwargs.
        # If **kwargs is present, we must NOT drop unknown keys (e.g. enable_mkldnn, run_mode).
        has_varkw = any(p.kind == inspect.Parameter.VAR_KEYWORD for p in sig.parameters.values())
        if not has_varkw:
            return {k: v for k, v in kwargs.items() if k in sig.parameters}

        # When **kwargs exists, only pass:
        # - explicitly supported args (in the signature), plus
        # - a small allow-list of runtime options consumed by PaddleOCR/PaddleX.
        allow_extra = {
            "enable_mkldnn",
            "run_mode",
            "device",
            "precision",
            "cpu_threads",
        }
        return {k: v for k, v in kwargs.items() if k in sig.parameters or k in allow_extra}
    except Exception:
        return kwargs


def _build_paddle_kwargs(profile: str) -> dict[str, Any]:
    lang = os.getenv("PADDLE_OCR_LANG", "en")
    det_limit = int(os.getenv("PADDLE_OCR_DET_LIMIT_SIDE_LEN", "1800"))
    kwargs: dict[str, Any] = {
        "lang": lang,
        "use_angle_cls": True,
        # PaddleOCR v3 uses explicit text_det_* args; older versions silently ignore unknown keys.
        "text_det_limit_side_len": det_limit,
        # Critical: disable MKLDNN/oneDNN to avoid PIR→oneDNN runtime crash on some PaddlePaddle builds.
        # (ConvertPirAttribute2RuntimeAttribute not support ... onednn_instruction.cc)
        "enable_mkldnn": False,
    }
    if profile == "fallback_v4":
        kwargs["ocr_version"] = os.getenv("PADDLE_OCR_FALLBACK_VERSION", "PP-OCRv4")
    else:
        kwargs["ocr_version"] = os.getenv("PADDLE_OCR_VERSION", "PP-OCRv4")
    if os.getenv("PADDLE_OCR_USE_GPU", "").lower() in ("1", "true", "yes"):
        kwargs["use_gpu"] = True
    return _filter_paddle_kwargs(kwargs)


def _init_paddle_engine(profile: str):
    global _OCR_ENGINE, _OCR_PROFILE
    from paddleocr import PaddleOCR

    try:
        import paddle

        paddle.set_flags(
            {
                "FLAGS_enable_pir_api": False,
                "FLAGS_use_new_executor": False,
                "FLAGS_enable_onednn": False,
                "FLAGS_use_mkldnn": False,
                "FLAGS_enable_pir_in_executor": False,
            }
        )
    except Exception:
        pass

    filtered = _build_paddle_kwargs(profile)
    try:
        _OCR_ENGINE = PaddleOCR(**filtered)
    except TypeError as e:
        if "unexpected keyword argument" in str(e):
            _OCR_ENGINE = PaddleOCR(
                lang=os.getenv("PADDLE_OCR_LANG", "en"),
                use_angle_cls=True,
                show_log=False,
                # Ensure MKLDNN stays disabled even on older PaddleOCR signatures.
                enable_mkldnn=False,
            )
        else:
            raise
    _OCR_PROFILE = profile
    logger.info("PaddleOCR initialized profile=%s", profile)


def get_paddle_engine():
    """Lazy-init PaddleOCR (thread-safe); tries default then PP-OCRv4 fallback."""
    global _OCR_ENGINE
    if _OCR_ENGINE is not None:
        return _OCR_ENGINE
    with _OCR_LOCK:
        if _OCR_ENGINE is not None:
            return _OCR_ENGINE
        last_err: Exception | None = None
        for profile in ("default", "fallback_v4"):
            try:
                _init_paddle_engine(profile)
                return _OCR_ENGINE
            except Exception as e:
                last_err = e
                _OCR_ENGINE = None
                continue
        raise RuntimeError(f"PaddleOCR init failed: {last_err}")


def warmup_paddle_engine_async() -> bool:
    """
    Kick off PaddleOCR initialization in background (non-blocking).
    Returns True if warmup was scheduled or already done; False if deps missing.
    """
    global _WARMUP_FUTURE
    try:
        import paddle  # noqa: F401
        from paddleocr import PaddleOCR  # noqa: F401
    except Exception:
        return False
    with _WARMUP_LOCK:
        if _OCR_ENGINE is not None:
            return True
        if _WARMUP_FUTURE is not None:
            return True
        try:
            _WARMUP_FUTURE = _OCR_POOL.submit(get_paddle_engine)
            logger.info("PaddleOCR warmup scheduled")
            return True
        except Exception:
            _WARMUP_FUTURE = None
            return False


def is_paddle_engine_ready() -> bool:
    return _OCR_ENGINE is not None


def is_paddle_warmup_done() -> bool:
    f = _WARMUP_FUTURE
    if f is None:
        return _OCR_ENGINE is not None
    try:
        return bool(f.done())
    except Exception:
        return False


def reinit_paddle_engine_fallback():
    """After runtime/oneDNN crashes, force PP-OCRv4 profile."""
    global _OCR_ENGINE
    with _OCR_LOCK:
        _OCR_ENGINE = None
        _init_paddle_engine("fallback_v4")


def ensure_paddle_imports():
    """Raise ImportError if numpy/cv2/paddleocr missing."""
    import numpy as np  # noqa: F401
    import cv2  # noqa: F401
    from paddleocr import PaddleOCR  # noqa: F401


def bytes_to_bgr(image_bytes: bytes):
    import numpy as np
    import cv2

    arr = np.frombuffer(image_bytes, dtype=np.uint8)
    return cv2.imdecode(arr, cv2.IMREAD_COLOR)


def resize_for_ocr(bgr, max_side: int | None = None):
    import cv2

    ms = max_side if max_side is not None else int(os.getenv("PADDLE_OCR_MAX_SIDE", "1800"))
    try:
        h, w = bgr.shape[:2]
        m = max(h, w)
        if m <= ms:
            return bgr
        scale = ms / float(m)
        nh = max(1, int(round(h * scale)))
        nw = max(1, int(round(w * scale)))
        return cv2.resize(bgr, (nw, nh), interpolation=cv2.INTER_AREA)
    except Exception:
        return bgr


def cv_preprocess(bgr):
    """Aggressive binarization — use when CLAHE/light path fails (faded scans)."""
    import cv2
    import numpy as np

    gray = cv2.cvtColor(bgr, cv2.COLOR_BGR2GRAY)
    equalized = cv2.equalizeHist(gray)
    denoised = cv2.GaussianBlur(equalized, (3, 3), 0)
    binary = cv2.adaptiveThreshold(
        denoised,
        255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY,
        31,
        15,
    )
    kernel = np.array([[0, -1, 0], [-1, 5, -1], [0, -1, 0]], dtype=np.float32)
    return cv2.filter2D(binary, -1, kernel)


def cv_preprocess_light(bgr):
    """CLAHE + unsharp on L channel — preserves photos/screenshots better than binary-only."""
    import cv2

    try:
        lab = cv2.cvtColor(bgr, cv2.COLOR_BGR2LAB)
        l_ch, a_ch, b_ch = cv2.split(lab)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        cl = clahe.apply(l_ch)
        merged = cv2.merge((cl, a_ch, b_ch))
        enhanced = cv2.cvtColor(merged, cv2.COLOR_LAB2BGR)
        blurred = cv2.GaussianBlur(enhanced, (0, 0), sigmaX=3)
        sharp = cv2.addWeighted(enhanced, 1.45, blurred, -0.45, 0)
        return sharp
    except Exception:
        return bgr


def _paddle_ocr_raw(img_bgr):
    try:
        engine = get_paddle_engine()
        return engine.ocr(img_bgr) or []
    except Exception as e:
        if _runtime_executor_error(e):
            logger.warning("Paddle runtime executor issue; reinitializing with fallback profile")
            reinit_paddle_engine_fallback()
            return get_paddle_engine().ocr(img_bgr) or []
        raise


def _reading_order_lines(ocr_result) -> list[tuple[list, str, float]]:
    """Sort detection boxes top-to-bottom, left-to-right; return (box, text, conf)."""
    items: list[tuple[list, str, float]] = []
    for block in ocr_result or []:
        for item in block or []:
            try:
                box = item[0]
                text = str(item[1][0] or "").strip()
                conf = float(item[1][1]) if len(item[1]) > 1 else 1.0
            except Exception:
                continue
            if text:
                items.append((box, text, conf))

    def sort_key(entry):
        box = entry[0]
        try:
            ys = [float(p[1]) for p in box]
            xs = [float(p[0]) for p in box]
            return (sum(ys) / max(len(ys), 1), sum(xs) / max(len(xs), 1))
        except Exception:
            return (0.0, 0.0)

    items.sort(key=sort_key)
    return items


def paddle_extract_structured(bgr_color, *, include_html_layout: bool = True) -> dict[str, Any]:
    """
    Run OCR with fallbacks: color BGR → light preprocess → aggressive binary.
    Returns plain text + line boxes for layout-aware clients.
    If include_html_layout=False, skips absolute-position HTML (faster for table-grid-only paths).
    """
    import cv2

    base = resize_for_ocr(bgr_color)
    variants: list[tuple[str, Any]] = [
        ("color", base),
        ("light", cv_preprocess_light(base)),
        ("binary", cv_preprocess(base)),
    ]
    last_items: list[tuple[list, str, float]] = []
    for name, im in variants:
        try:
            to_run = im
            if len(to_run.shape) == 2:
                to_run = cv2.cvtColor(to_run, cv2.COLOR_GRAY2BGR)
            result = _paddle_ocr_raw(to_run)
            items = _reading_order_lines(result)
            if items:
                last_items = items
                lines_text = "\n".join(t for _, t, _ in items).strip()
                if lines_text:
                    logger.debug("Paddle OCR path=%s lines=%s", name, len(items))
                    line_dicts = [
                        {"text": t, "bbox": box, "confidence": conf}
                        for box, t, conf in items
                    ]
                    ih, iw = int(to_run.shape[0]), int(to_run.shape[1])
                    html_layout = ""
                    if include_html_layout:
                        try:
                            from backend.services.ocr_layout_reconstruct import build_absolute_layout_html

                            html_layout = build_absolute_layout_html(to_run, line_dicts, iw, ih)
                        except Exception as le:
                            logger.debug("Layout HTML skipped: %s", le)
                    return {
                        "text": lines_text,
                        "lines": line_dicts,
                        "preprocess": name,
                        "html_layout": html_layout,
                        "width": iw,
                        "height": ih,
                    }
        except Exception as e:
            logger.warning("Paddle variant %s failed: %s", name, e)
            continue

    return {"text": "", "lines": [], "preprocess": "none"}


def paddle_extract_plain_text(bgr_color) -> str:
    """Backward-compatible: structured extract, plain text only."""
    return (paddle_extract_structured(bgr_color, include_html_layout=False) or {}).get("text") or ""


def docx_image_then_text(image_bytes: bytes, text: str) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    doc = Document()
    try:
        sec = doc.sections[0]
        sec.top_margin = Inches(0.4)
        sec.bottom_margin = Inches(0.4)
        sec.left_margin = Inches(0.4)
        sec.right_margin = Inches(0.4)
    except Exception:
        pass

    bio = io.BytesIO(image_bytes)
    bio.name = "upload.jpg"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(bio, width=Inches(7.2))
    doc.add_page_break()
    for line in (text or "").splitlines():
        doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def ocr_extract_text_from_bytes(image_bytes: bytes) -> dict[str, Any]:
    img = bytes_to_bgr(image_bytes)
    if img is None:
        raise ValueError("Invalid image")
    img = resize_for_ocr(img)
    structured = paddle_extract_structured(img)
    text = (structured.get("text") or "").strip()
    out: dict[str, Any] = {"text": text}
    lines = structured.get("lines") or []
    if lines:
        out["lines"] = lines
    if structured.get("preprocess"):
        out["preprocess"] = structured["preprocess"]
    hl = (structured.get("html_layout") or "").strip()
    if hl:
        out["html_layout"] = hl
    if structured.get("width") is not None:
        out["width"] = structured["width"]
    if structured.get("height") is not None:
        out["height"] = structured["height"]
    return out


def ocr_image_to_docx_from_bytes(image_bytes: bytes) -> dict[str, Any]:
    img = bytes_to_bgr(image_bytes)
    if img is None:
        raise ValueError("Invalid image")
    img = resize_for_ocr(img)
    try:
        text = paddle_extract_plain_text(img)
    except Exception as e:
        logger.warning("Paddle OCR failed, image-only DOCX: %s", e)
        text = ""
    docx_bytes = docx_image_then_text(image_bytes, text)
    return {"docx_b64": base64.b64encode(docx_bytes).decode("ascii"), "text": text}
