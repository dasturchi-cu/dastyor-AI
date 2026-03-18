from __future__ import annotations

import base64
import io
import os
import threading
from typing import Any

from backend.celery_app import celery_app


# Paddle runtime flags for stability (must be set before import/initialization)
os.environ.setdefault("FLAGS_enable_pir_api", "0")
os.environ.setdefault("FLAGS_use_new_executor", "0")
os.environ.setdefault("FLAGS_enable_onednn", "0")
os.environ.setdefault("FLAGS_use_mkldnn", "0")


_OCR_LOCK = threading.Lock()
_OCR_ENGINE = None


def _get_ocr_engine():
    """
    Initialize PaddleOCR once per worker process and reuse it.
    """
    global _OCR_ENGINE
    if _OCR_ENGINE is not None:
        return _OCR_ENGINE
    with _OCR_LOCK:
        if _OCR_ENGINE is not None:
            return _OCR_ENGINE
        from paddleocr import PaddleOCR

        lang = os.getenv("PADDLE_OCR_LANG", "en")
        det_limit = int(os.getenv("PADDLE_OCR_DET_LIMIT_SIDE_LEN", "1800"))
        _OCR_ENGINE = PaddleOCR(
            lang=lang,
            use_angle_cls=True,
            show_log=False,
            det_limit_side_len=det_limit,
            enable_mkldnn=False,
        )
        return _OCR_ENGINE


def _bytes_to_bgr(image_bytes: bytes):
    import numpy as np
    import cv2

    arr = np.frombuffer(image_bytes, dtype=np.uint8)
    img = cv2.imdecode(arr, cv2.IMREAD_COLOR)
    return img


def _resize_for_ocr(bgr, max_side: int = 1800):
    import cv2

    try:
        h, w = bgr.shape[:2]
        m = max(h, w)
        if m <= max_side:
            return bgr
        scale = max_side / float(m)
        nh = max(1, int(round(h * scale)))
        nw = max(1, int(round(w * scale)))
        return cv2.resize(bgr, (nw, nh), interpolation=cv2.INTER_AREA)
    except Exception:
        return bgr


def _cv_preprocess(bgr):
    import cv2
    import numpy as np

    gray = cv2.cvtColor(bgr, cv2.COLOR_BGR2GRAY)
    equalized = cv2.equalizeHist(gray)
    denoised = cv2.GaussianBlur(equalized, (3, 3), 0)
    binary = cv2.adaptiveThreshold(
        denoised,
        255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY,
        31,
        15,
    )
    kernel = np.array([[0, -1, 0], [-1, 5, -1], [0, -1, 0]], dtype=np.float32)
    sharpened = cv2.filter2D(binary, -1, kernel)
    return sharpened


def _paddle_extract_text(processed) -> str:
    import cv2

    engine = _get_ocr_engine()
    img = cv2.cvtColor(processed, cv2.COLOR_GRAY2BGR) if len(processed.shape) == 2 else processed
    result = engine.ocr(img) or []
    lines: list[str] = []
    for block in result:
        for item in block or []:
            try:
                t = item[1][0]
            except Exception:
                t = ""
            if t:
                lines.append(t)
    return "\n".join(lines).strip()


def _docx_image_then_text(image_bytes: bytes, text: str) -> bytes:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    try:
        sec = doc.sections[0]
        sec.top_margin = Inches(0.4)
        sec.bottom_margin = Inches(0.4)
        sec.left_margin = Inches(0.4)
        sec.right_margin = Inches(0.4)
    except Exception:
        pass

    bio = io.BytesIO(image_bytes)
    bio.name = "upload.jpg"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(bio, width=Inches(7.2))
    doc.add_page_break()
    for line in (text or "").splitlines():
        doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@celery_app.task(name="ocr.extract_text")
def ocr_extract_text(image_bytes: bytes) -> dict[str, Any]:
    img = _bytes_to_bgr(image_bytes)
    if img is None:
        raise ValueError("Invalid image")
    img = _resize_for_ocr(img, max_side=int(os.getenv("PADDLE_OCR_MAX_SIDE", "1800")))
    processed = _cv_preprocess(img)
    text = _paddle_extract_text(processed)
    return {"text": text}


@celery_app.task(name="ocr.image_to_docx")
def ocr_image_to_docx(image_bytes: bytes) -> dict[str, Any]:
    img = _bytes_to_bgr(image_bytes)
    if img is None:
        raise ValueError("Invalid image")
    img = _resize_for_ocr(img, max_side=int(os.getenv("PADDLE_OCR_MAX_SIDE", "1800")))
    processed = _cv_preprocess(img)
    text = _paddle_extract_text(processed)
    docx_bytes = _docx_image_then_text(image_bytes, text)
    return {"docx_b64": base64.b64encode(docx_bytes).decode("ascii"), "text": text}

