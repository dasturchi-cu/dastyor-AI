"""
Build templates/obyektivka_master.docx from reference.
Run: python scripts/build_obyektivka_master.py
"""
from __future__ import annotations

import re
import shutil
import sys
import zipfile
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
REF_DOCX = ROOT / "temp" / "ref_converted.docx"
REF_DOC = ROOT / "Намуна Объективка (18).doc"
OUT = ROOT / "templates" / "obyektivka_master.docx"

TEXT_REPLACEMENTS: list[tuple[str, str]] = [
    ("Эшматов Ботир Баҳодирович", "{{fish}}"),
    ("2007 йил 5 октябрдан:", "{{hozirgi_yil}}"),
    ("25.10.1960", "{{tugilgan_sana}}"),
    ("Тошкент вилояти, Қибрай тумани ", "{{tugilgan_joy}}"),
    ("ўзбек", "{{millati}}"),
    ("олий", "{{malumoti}}"),
    ("1982 й. Тошкент давлат университети (кундузги)", "{{tamomlagan}}"),
    ("иқтисодчи", "{{mutaxassisligi}}"),
    ("иқтисод фанлари доктори", "{{ilmiy_darajasi}}"),
    ("профессор", "{{ilmiy_unvoni}}"),
    ("рус, инглиз тиллари ", "{{chet_tillari}}"),
    (
        "2005 й. “Шуҳрати” медал, 2010 й. “Меҳнат шуҳрати” ордени",
        "{{mukofotlari}}",
    ),
    (
        "2008 й. «Халқ таълим аълочиси» кўкрак нишони, 2017 й. «Ўзбекистон Конституциясига 25 йил» эсдалик нишони, 2025 й. «Ғалабанинг 80 йиллиги» юбилей нишони, 2026 й. «Фидокорона меҳнати учун» кўкрак нишони",
        "{{idoriy_mukofotlari}}",
    ),
    (
        "2024 й.  - ҳ.в. - Халқ депутатлари Тошкент вилояти Кенгаши депутати, Ўзбекистон Республикаси Олий Мажлиси Сенати аъзоси ",
        "{{deputatligi}}",
    ),
    ("1977-1982 йй. - Тошкент давлат университети талабаси", "{{mehnat_faoliyati}}"),
    ("1982-1988 йй. - Тошкент давлат университети иқтисодиёт факультети кичик илмий ходими", "{{mehnat_faoliyati_2}}"),
    ("1988-1991 йй. - Тошкент давлат университети иқтисодиёт факультети аспиранти", "{{mehnat_faoliyati_3}}"),
    (
        "1991-1995 йй. - Тошкент давлат иқтисодиёт университети иқтисодиёт факультети катта илмий ходими",
        "{{mehnat_faoliyati_4}}",
    ),
    (
        "1995-1998 йй. - Тошкент давлат иқтисодиёт университети иқтисодиёт факультети докторанти",
        "{{mehnat_faoliyati_5}}",
    ),
    (
        "1998-2004 йй. - Ўзбекистон Республикаси Иқтисодиёт вазирлиги таълимни ривожлантириш бўлими мутахассиси, етакчи мутахассиси, бош мутахассиси",
        "{{mehnat_faoliyati_6}}",
    ),
    (
        "2004-2007 йй. - Тошкент давлат иқтисодиёт университети микроиқтисодиёт факультети декани",
        "{{mehnat_faoliyati_7}}",
    ),
    (
        "2007 й. -  ҳ.в.  - Андижон вилояти Андижон тумани “Makon Mirzo” масъулияти чекланган жамияти раҳбари",
        "{{mehnat_faoliyati_8}}",
    ),
    (
        "Андижон вилояти Андижон тумани “Makon Mirzo” масъулияти чекланган жамияти раҳбари",
        "{{hozirgi_ish}}",
    ),
]

REL_ROW_PREFIXES: list[tuple[str, str]] = [
    ("Отаси", "ota"),
    ("Онаси", "ona"),
    ("Опаси", "opa"),
    ("Синглиси", "singil"),
    ("Укаси", "uka"),
    ("Укаси", "aka"),
    ("Турмуш ўртоғи", "turmush_ortogi"),
    ("Ўғли", "farzandlar"),
    ("Қизи", "farzandlar_2"),
    ("Ўғли", "farzandlar_3"),
    ("Ўғли", "farzandlar_4"),
    ("Қизи", "farzandlar_5"),
    ("Ўғли", "farzandlar_6"),
    ("Қайнотаси", "qaynota"),
    ("Қайнонаси", "qaynona"),
]


def _set_cell_placeholder(cell, placeholder: str) -> None:
    if not cell.paragraphs:
        cell.add_paragraph()
    p = cell.paragraphs[0]
    for run in list(p.runs):
        p._p.remove(run._r)
    p.add_run(placeholder)
    for extra in cell.paragraphs[1:]:
        for run in list(extra.runs):
            extra._p.remove(run._r)


def _replace_relatives_table(doc: Document) -> None:
    if not doc.tables:
        return
    table = doc.tables[0]
    uka_seen = 0
    prefix_queue: list[str] = []
    for lbl, pfx in REL_ROW_PREFIXES:
        if lbl == "Укаси":
            prefix_queue.append("uka" if uka_seen == 0 else "aka")
            uka_seen += 1
        else:
            prefix_queue.append(pfx)

    for row, prefix in zip(table.rows[1:], prefix_queue):
        for i, suffix in enumerate(("", "_yil", "_ish", "_tur"), start=1):
            if i < len(row.cells):
                _set_cell_placeholder(row.cells[i], f"{{{{{prefix}{suffix}}}}}")


def _replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    full = paragraph.text
    if old not in full:
        return False
    new_full = full.replace(old, new, 1)
    if not paragraph.runs:
        paragraph.add_run(new_full)
        return True
    paragraph.runs[0].text = new_full
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def _replace_all(doc: Document, old: str, new: str) -> int:
    n = 0
    for p in doc.paragraphs:
        if _replace_in_paragraph(p, old, new):
            n += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if _replace_in_paragraph(p, old, new):
                        n += 1
    return n


def _fix_party_military(doc: Document) -> None:
    for p in doc.paragraphs:
        text = p.text
        if text.startswith("{{millati}}\t"):
            _replace_in_paragraph(p, "йўқ", "{{partiyaviyligi}}")
        if text.startswith("{{chet_tillari}}\t"):
            _replace_in_paragraph(p, "йўқ", "{{harbiy_unvoni}}")


def _fix_relatives_title(doc: Document) -> None:
    for p in doc.paragraphs:
        if "яқин қариндошлари ҳақида" in p.text and "{{fish}}" not in p.text:
            for run in p.runs:
                if "Эшматов" in run.text:
                    run.text = run.text.replace("Эшматов Ботир Баҳодирович", "{{fish}}")
                    return


def _mark_photo_placeholder(path: Path) -> None:
    with zipfile.ZipFile(path, "r") as zin:
        parts = {item.filename: zin.read(item.filename) for item in zin.infolist()}
    xml = parts["word/document.xml"].decode("utf-8")
    xml = re.sub(
        r"<w:t[^>]*>Оқ фондаги[^<]*</w:t>",
        "<w:t>{{photo}}</w:t>",
        xml,
        count=1,
    )
    parts["word/document.xml"] = xml.encode("utf-8")
    with zipfile.ZipFile(path, "w") as zout:
        for name, data in parts.items():
            zout.writestr(name, data)


def ensure_ref_docx() -> Path:
    if REF_DOCX.is_file():
        return REF_DOCX
    REF_DOCX.parent.mkdir(parents=True, exist_ok=True)
    if REF_DOC.is_file():
        try:
            import win32com.client  # type: ignore

            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(str(REF_DOC.resolve()))
            doc.SaveAs(str(REF_DOCX.resolve()), FileFormat=16)
            doc.Close()
            word.Quit()
            return REF_DOCX
        except Exception as exc:
            print("Word COM conversion failed:", exc, file=sys.stderr)
    raise FileNotFoundError(f"Reference DOCX missing: {REF_DOCX}")


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    src = ensure_ref_docx()
    OUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, OUT)
    doc = Document(str(OUT))

    for old, new in TEXT_REPLACEMENTS:
        _replace_all(doc, old, new)
    _replace_relatives_table(doc)
    _fix_party_military(doc)
    _fix_relatives_title(doc)
    doc.save(OUT)
    _mark_photo_placeholder(OUT)
    print("written", OUT)


if __name__ == "__main__":
    main()
