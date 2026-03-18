"""
Main Bot Entry Point (with Ban Check Middleware)
"""
import os
import io
import asyncio
import logging
import threading
import inspect
from concurrent.futures import ThreadPoolExecutor

# ── Paddle / PaddleOCR stability flags (Windows/oneDNN/PIR issues) ─────────────
# Set before PaddleOCR initializes models.
# Force-disable flags (do not use setdefault) because environment may already
# provide conflicting values on some Windows setups.
os.environ["FLAGS_enable_pir_api"] = "0"
os.environ["FLAGS_use_new_executor"] = "0"
os.environ["FLAGS_enable_onednn"] = "0"
os.environ["FLAGS_use_mkldnn"] = "0"
os.environ["FLAGS_enable_pir_in_executor"] = "0"

from dotenv import load_dotenv
from telegram import Update
from telegram.ext import (
    ApplicationBuilder, CommandHandler, MessageHandler, 
    filters, ContextTypes, TypeHandler, CallbackQueryHandler, ChatMemberHandler
)

try:
    load_dotenv()
except: pass

from config import BOT_TOKEN, logger

# ──────────────────────────────────────────────────────────────────────────────
# DASTYOR AI — OCR Backend (FastAPI + PaddleOCR + OpenCV + python-docx)
#
# Endpoints:
#   POST /ocr      → JSON: {"text": "..."}
#   POST /ocr-word → DOCX download with detected text
#
# Note:
# - PaddleOCR is initialized once at server startup.
# - OCR runs in a thread pool (CPU-bound).
#
# Run API mode:
#   RUN_MODE=api uvicorn main:app --host 0.0.0.0 --port 8000
# Run bot mode (default):
#   python main.py
# ──────────────────────────────────────────────────────────────────────────────

try:
    import numpy as np
    import cv2
    from paddleocr import PaddleOCR
    from docx import Document
    from fastapi import FastAPI, UploadFile, File, HTTPException
    from fastapi.responses import StreamingResponse
    from fastapi.middleware.cors import CORSMiddleware
except Exception:
    # Allow bot-only mode even if OCR deps not installed.
    np = None
    cv2 = None
    PaddleOCR = None
    Document = None
    FastAPI = None
    UploadFile = None
    File = None
    HTTPException = Exception
    StreamingResponse = None
    CORSMiddleware = None

_ocr_logger = logging.getLogger("dastyor.ocr")
_OCR_ENGINE = None
_OCR_ENGINE_PROFILE = "default"
_OCR_LOCK = threading.Lock()
_OCR_POOL = ThreadPoolExecutor(max_workers=int(os.getenv("OCR_WORKERS", "2")))

app = FastAPI(title="Dastyor AI OCR API") if FastAPI else None

if app and CORSMiddleware:
    app.add_middleware(
        CORSMiddleware,
        allow_origins=["*"],
        allow_credentials=True,
        allow_methods=["*"],
        allow_headers=["*"],
    )


def _ensure_ocr_deps():
    if any(x is None for x in (np, cv2, PaddleOCR, Document, StreamingResponse)) or app is None:
        raise RuntimeError(
            "OCR dependencies are missing. Install: paddleocr, opencv-python, numpy, python-docx, fastapi, uvicorn"
        )


def _runtime_executor_error(exc: Exception) -> bool:
    msg = str(exc)
    return "ConvertPirAttribute2RuntimeAttribute" in msg or "onednn_instruction.cc" in msg


def _build_ocr_engine(lang: str, profile: str):
    """
    Build PaddleOCR engine with profile-based kwargs.
    """
    def _safe_paddleocr_init(kwargs):
        try:
            sig = inspect.signature(PaddleOCR)
            allowed = set(sig.parameters.keys())
            filtered = {k: v for k, v in kwargs.items() if k in allowed}
        except Exception:
            filtered = kwargs
        return PaddleOCR(**filtered)

    if profile == "fallback_v4":
        return _safe_paddleocr_init(
            {
                "lang": lang,
                "use_angle_cls": True,
                "ocr_version": os.getenv("PADDLE_OCR_FALLBACK_VERSION", "PP-OCRv4"),
                "enable_mkldnn": False,
                "show_log": False,
            }
        )
    return _safe_paddleocr_init(
        {
            "lang": lang,
            "use_angle_cls": True,
            "ocr_version": os.getenv("PADDLE_OCR_VERSION", "PP-OCRv4"),
            "enable_mkldnn": False,
            "show_log": False,
        }
    )


def _init_ocr_engine(lang: str, force_fallback: bool = False):
    global _OCR_ENGINE, _OCR_ENGINE_PROFILE
    with _OCR_LOCK:
        profiles = ["fallback_v4"] if force_fallback else ["default", "fallback_v4"]
        last_err = None
        for profile in profiles:
            try:
                _ocr_logger.info("Initializing PaddleOCR (lang=%s, profile=%s)...", lang, profile)
                _OCR_ENGINE = _build_ocr_engine(lang, profile)
                _OCR_ENGINE_PROFILE = profile
                _ocr_logger.info("PaddleOCR initialized (profile=%s).", profile)
                return
            except TypeError as e:
                # Some PaddleOCR versions may not support specific kwargs.
                last_err = e
                if "unexpected keyword argument" in str(e):
                    try:
                        _OCR_ENGINE = PaddleOCR(lang=lang, use_angle_cls=True)
                        _OCR_ENGINE_PROFILE = f"{profile}_compat"
                        _ocr_logger.info("PaddleOCR initialized (profile=%s_compat).", profile)
                        return
                    except Exception as compat_err:
                        last_err = compat_err
                        continue
            except Exception as e:
                last_err = e
                continue
        raise RuntimeError(f"PaddleOCR init failed: {last_err}")


def preprocess_image_opencv(image_bgr):
    """
    OpenCV preprocessing:
    - grayscale
    - threshold (Otsu)
    Returns a processed image suitable for OCR.
    """
    gray = cv2.cvtColor(image_bgr, cv2.COLOR_BGR2GRAY)
    # Otsu thresholding improves contrast for text
    _, th = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    return th


def extract_layout_blocks_with_paddle(processed_img):
    """
    Run PaddleOCR and return normalized blocks with bbox information.
    """
    global _OCR_ENGINE
    if _OCR_ENGINE is None:
        raise RuntimeError("OCR engine is not initialized")

    # PaddleOCR works well with 3-channel images; convert if needed.
    if len(processed_img.shape) == 2:
        img_for_ocr = cv2.cvtColor(processed_img, cv2.COLOR_GRAY2BGR)
    else:
        img_for_ocr = processed_img

    try:
        result = _OCR_ENGINE.ocr(img_for_ocr) or []
    except Exception as e:
        if _runtime_executor_error(e):
            _ocr_logger.warning(
                "Paddle runtime executor issue detected. Reinitializing OCR engine with fallback profile."
            )
            lang = os.getenv("PADDLE_OCR_LANG", "en")
            _init_ocr_engine(lang=lang, force_fallback=True)
            result = _OCR_ENGINE.ocr(img_for_ocr) or []
        else:
            raise
    blocks = []
    # result: list[ list[ [box], (text, score) ] ]
    for block in result:
        for item in block or []:
            try:
                points = item[0]
                text = str(item[1][0]).strip()
                confidence = float(item[1][1])
                xs = [float(p[0]) for p in points]
                ys = [float(p[1]) for p in points]
            except Exception:
                continue
            if text:
                left = min(xs)
                top = min(ys)
                right = max(xs)
                bottom = max(ys)
                blocks.append(
                    {
                        "text": text,
                        "confidence": confidence,
                        "bbox": [left, top, right, bottom],
                        "left": left,
                        "top": top,
                        "right": right,
                        "bottom": bottom,
                        "center_y": (top + bottom) / 2.0,
                        "height": max(1.0, bottom - top),
                        "width": max(1.0, right - left),
                    }
                )
    return sorted(blocks, key=lambda b: (b["top"], b["left"]))


def group_blocks_by_line(blocks):
    if not blocks:
        return []

    heights = [b["height"] for b in blocks]
    median_h = float(np.median(heights)) if heights else 18.0
    y_tolerance = max(8.0, median_h * 0.65)
    sorted_blocks = sorted(blocks, key=lambda b: (b["center_y"], b["left"]))
    lines = []

    for block in sorted_blocks:
        placed = False
        for line in reversed(lines):
            if abs(block["center_y"] - line["avg_center_y"]) <= y_tolerance:
                line["blocks"].append(block)
                line["avg_center_y"] = float(
                    np.mean([x["center_y"] for x in line["blocks"]])
                )
                line["top"] = min(line["top"], block["top"])
                line["bottom"] = max(line["bottom"], block["bottom"])
                placed = True
                break
        if not placed:
            lines.append(
                {
                    "blocks": [block],
                    "avg_center_y": block["center_y"],
                    "top": block["top"],
                    "bottom": block["bottom"],
                }
            )

    for line in lines:
        line["blocks"] = sorted(line["blocks"], key=lambda b: b["left"])
        line["left"] = min(b["left"] for b in line["blocks"])
        line["right"] = max(b["right"] for b in line["blocks"])
        line["height"] = max(1.0, line["bottom"] - line["top"])

    return sorted(lines, key=lambda l: l["top"])


def render_layout_text(lines):
    rendered_lines = []
    for line in lines:
        blocks = line["blocks"]
        if not blocks:
            rendered_lines.append("")
            continue

        avg_char_w = max(
            6.0,
            float(
                np.median(
                    [
                        max(1.0, b["width"]) / max(1, len(b["text"]))
                        for b in blocks
                    ]
                )
            ),
        )
        line_text = blocks[0]["text"]
        for idx in range(1, len(blocks)):
            prev = blocks[idx - 1]
            curr = blocks[idx]
            gap = max(0.0, curr["left"] - prev["right"])
            spaces = max(1, min(24, int(round(gap / avg_char_w))))
            line_text += (" " * spaces) + curr["text"]
        rendered_lines.append(line_text.strip())

    return "\n".join(rendered_lines).strip()


def build_docx_with_layout(lines) -> bytes:
    """
    Create a DOCX with approximate original text layout.
    """
    from docx.shared import Inches

    doc = Document()
    if not lines:
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    all_left = min(line["left"] for line in lines)
    all_right = max(line["right"] for line in lines)
    doc_width_px = max(1.0, all_right - all_left)
    target_text_width_in = 6.2
    vertical_unit = max(10.0, float(np.median([line["height"] for line in lines])))

    prev_bottom = None
    for line in lines:
        if prev_bottom is not None:
            v_gap = max(0.0, line["top"] - prev_bottom)
            extra_blank = max(0, min(4, int(v_gap / vertical_unit) - 1))
            for _ in range(extra_blank):
                doc.add_paragraph("")
        prev_bottom = line["bottom"]

        p = doc.add_paragraph()
        line_indent_in = ((line["left"] - all_left) / doc_width_px) * target_text_width_in
        p.paragraph_format.left_indent = Inches(max(0.0, min(target_text_width_in, line_indent_in)))

        blocks = line["blocks"]
        avg_char_w = max(
            6.0,
            float(
                np.median(
                    [
                        max(1.0, b["width"]) / max(1, len(b["text"]))
                        for b in blocks
                    ]
                )
            ),
        )
        for idx, block in enumerate(blocks):
            p.add_run(block["text"])
            if idx < len(blocks) - 1:
                next_block = blocks[idx + 1]
                h_gap = max(0.0, next_block["left"] - block["right"])
                spaces = max(1, min(24, int(round(h_gap / avg_char_w))))
                p.add_run(" " * spaces)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


if app:
    @app.on_event("startup")
    async def _startup_init_ocr():
        _ensure_ocr_deps()
        global _OCR_ENGINE
        if _OCR_ENGINE is None:
            lang = os.getenv("PADDLE_OCR_LANG", "en")
            # Initialize once
            try:
                import paddle
                paddle.set_flags({
                    "FLAGS_enable_pir_api": False,
                    "FLAGS_use_new_executor": False,
                    "FLAGS_enable_onednn": False,
                    "FLAGS_use_mkldnn": False,
                    "FLAGS_enable_pir_in_executor": False,
                })
            except Exception:
                pass
            _init_ocr_engine(lang=lang, force_fallback=False)


    async def _read_image_as_bgr(file: UploadFile):
        raw = await file.read()
        if not raw:
            raise HTTPException(status_code=400, detail="Empty file")
        if len(raw) > 15 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="File too large (max 15MB)")
        arr = np.frombuffer(raw, dtype=np.uint8)
        img = cv2.imdecode(arr, cv2.IMREAD_COLOR)
        if img is None:
            raise HTTPException(status_code=400, detail="Invalid image")
        return img


    @app.post("/ocr")
    async def ocr_endpoint(file: UploadFile = File(...)):
        """
        Receive image → preprocess (gray + threshold) → PaddleOCR → return JSON text.
        """
        try:
            _ensure_ocr_deps()
            img = await _read_image_as_bgr(file)
            processed = preprocess_image_opencv(img)
            loop = asyncio.get_running_loop()
            blocks = await loop.run_in_executor(_OCR_POOL, extract_layout_blocks_with_paddle, processed)
            lines = group_blocks_by_line(blocks)
            text = render_layout_text(lines)
            return {
                "text": text,
                "line_count": len(lines),
                "block_count": len(blocks),
                "blocks": [
                    {
                        "text": b["text"],
                        "confidence": round(b["confidence"], 4),
                        "bbox": [round(v, 2) for v in b["bbox"]],
                    }
                    for b in blocks
                ],
            }
        except HTTPException:
            raise
        except Exception as e:
            _ocr_logger.error("OCR /ocr failed: %s", e, exc_info=True)
            raise HTTPException(status_code=500, detail="OCR failed")


    @app.post("/ocr-word")
    async def ocr_word_endpoint(file: UploadFile = File(...)):
        """
        Receive image → OCR → generate DOCX → return as download.
        """
        try:
            _ensure_ocr_deps()
            raw = await file.read()
            if not raw:
                raise HTTPException(status_code=400, detail="Empty file")
            if len(raw) > 15 * 1024 * 1024:
                raise HTTPException(status_code=400, detail="File too large (max 15MB)")
            arr = np.frombuffer(raw, dtype=np.uint8)
            img = cv2.imdecode(arr, cv2.IMREAD_COLOR)
            if img is None:
                raise HTTPException(status_code=400, detail="Invalid image")
            processed = preprocess_image_opencv(img)
            loop = asyncio.get_running_loop()
            blocks = await loop.run_in_executor(_OCR_POOL, extract_layout_blocks_with_paddle, processed)
            lines = group_blocks_by_line(blocks)
            docx_bytes = await loop.run_in_executor(_OCR_POOL, build_docx_with_layout, lines)
            filename = "ocr_result.docx"
            return StreamingResponse(
                io.BytesIO(docx_bytes),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{filename}"'},
            )
        except HTTPException:
            raise
        except Exception as e:
            _ocr_logger.error("OCR /ocr-word failed: %s", e, exc_info=True)
            raise HTTPException(status_code=500, detail="OCR-word failed")

# Handlers
from bot.handlers.admin import (
    admin_panel_command, stats_command, broadcast_command,
    handle_admin_text, add_channel_command, remove_channel_command,
    add_premium_command, remove_premium_command, set_limit_command,
    user_info_command, top_users_command, ban_user_command, unban_user_command,
    search_command, support_panel_callback, add_admin_command, remove_admin_command,
    approve_premium_command
)

from bot.handlers.admin_middleware import track_user
from bot.handlers.premium_callbacks import premium_callback_handler
from bot.handlers.premium import (
    premium_handler,
    premium_purchase_callback,
    handle_premium_screenshot,
    premium_payment_review_callback,
)
from bot.handlers.help import help_command
from bot.handlers.chat_member import chat_member_updated
from bot.handlers.common import balance_handler, help_button_handler
from bot.handlers.feedback import start_feedback, handle_feedback
from bot.handlers.support_group import support_group_router, SUPPORT_GROUP_ID


from bot.handlers.ocr_to_word import (
    ocr_to_word_handler as ocr_handler,
    handle_ocr_image as process_ocr_image,
    process_ocr_tayyor,
)
from bot.handlers.obyektivka import obyektivka_handler, handle_obyektivka_audio as process_obyektivka_audio
from bot.handlers.transliterate import transliterate_handler, process_transliteration as process_transliterate, krill_to_lotin_handler, lotin_to_krill_handler, translit_direction_callback
from bot.handlers.translate import translate_handler, process_translation as process_translate_doc, set_translation_direction
from bot.handlers.image_to_pdf import image_to_pdf_handler, collect_pdf_images as process_image_to_pdf
from bot.handlers.spell_check import spell_check_handler, process_spell_check
from bot.handlers.start import start_command, menu_command
from bot.keyboards.reply_keyboards import get_main_menu, get_back_button, get_more_menu
from bot.utils.i18n import get_regex_for_key, t
from bot.handlers.smart_logic import (
    handle_smart_photo, handle_smart_document, handle_smart_audio, smart_callback_handler
)
from bot.handlers.webapp_data import web_app_data_handler

# Services
from bot.services.settings_service import is_premium
from bot.services.user_service import increment_file_count, get_user_lang

async def back_to_main_menu(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not update.effective_chat or update.effective_chat.type != "private":
        return
    context.user_data.clear()
    lang = get_user_lang(update.effective_user.id) if update.effective_user else "uz_lat"
    await update.message.reply_text(t("or_menu", lang), reply_markup=get_main_menu(update.effective_user.id if update.effective_user else None, lang))

async def more_menu_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Show 'Boshqa xizmatlar' sub-menu"""
    if not update.effective_chat or update.effective_chat.type != "private":
        return
    uid = update.effective_user.id if update.effective_user else None
    lang = get_user_lang(uid)
    await update.message.reply_text(t("more_menu_title", lang), reply_markup=get_more_menu(lang))

async def cv_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Open CV Resume webapp page via WebApp inline button"""
    if not update.effective_chat or update.effective_chat.type != "private":
        return
    from bot.handlers.start import _ACTION_MAP, WEBAPP_BASE
    from telegram import WebAppInfo, InlineKeyboardMarkup, InlineKeyboardButton
    uid = update.effective_user.id if update.effective_user else 0
    lang = get_user_lang(uid)
    page_file, btn_label, desc = _ACTION_MAP["cv"]
    url = f"{WEBAPP_BASE}/{page_file}?telegram_id={uid}&lang={lang}"
    kb = InlineKeyboardMarkup([[InlineKeyboardButton(btn_label, web_app=WebAppInfo(url=url))]])
    await update.message.reply_text(f"🚀 <b>{desc}</b>", reply_markup=kb, parse_mode="HTML")

async def premium_info_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await premium_handler(update, context)

async def unified_router_check(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Central check for ban status"""
    if context.user_data.get('is_banned'):
        await update.message.reply_text("🚫 Siz botdan foydalanishdan bloklangansiz.")
        return False
    return True

from bot.handlers.admin import process_admin_state_input

async def handle_router_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not update.effective_chat or update.effective_chat.type != "private":
        return
    if not await unified_router_check(update, context): return
    if await process_admin_state_input(update, context): return
    
    state = context.user_data.get('waiting_for')
    text = (update.message.text or "").strip().lower()
    
    # 1. State-based routing
    if state == 'ocr_image' and text and 'tayyor' in text:
        if await process_ocr_tayyor(update, context):
            return
        await update.message.reply_text("❌ Hech qanday rasm yuklanmagan. Avval rasmlar yuboring.")
        return
    if state in ['transliterate_text', 'translit_content'] or context.user_data.get('transliterate_direction'):
         await process_transliterate(update, context)
         return
    elif state == 'translate_input' or context.user_data.get('translate_direction'):
         await process_translate_doc(update, context)
         return
    elif state == 'spell_check_doc' or state == 'spellcheck_file':
         await process_spell_check(update, context)
         return
    elif state == 'ocr_image' and context.user_data.get('ocr_images') and text and 'tayyor' in text:
         if await process_ocr_tayyor(update, context):
             return
    elif state == 'pdf_images':
         await process_image_to_pdf(update, context)
         return
    elif state == 'feedback':
         await handle_feedback(update, context)
         return

    # 2. NLP / Keyword Routing
    import re
    # Obyektivka (obyektivga, obyektovka, obyektvka, abyektiv)
    uid = update.effective_user.id
    lang = get_user_lang(uid)
    if re.search(r'(obyektiv|obyektov|abyektiv|obekt|resume|rezume|sivi|ma\'lumotnoma)', text):
        await update.message.reply_text(t("opening_service", lang, service="Obyektivka"))
        await obyektivka_handler(update, context)
        return

    # OCR / Word (docx, doc, dox, vord, ocr, textga)
    elif re.search(r'(ocr|word|vord|docx|doc|dox|matn|textga|oqib ber)', text) or (('rasm' in text or 'skan' in text) and ('o\'qi' in text or 'qil' in text)):
        await update.message.reply_text(t("opening_service", lang, service="Rasm -> Word"))
        await ocr_handler(update, context)
        return

    # Image 2 PDF (rasm... pdf)
    elif 'pdf' in text and ('rasm' in text or 'qo\'sh' in text or 'birlash' in text):
        await update.message.reply_text(t("opening_service", lang, service="Rasm -> PDF"))
        await image_to_pdf_handler(update, context)
        return

    # Translate (tarjima, pervod, perevod, translate)
    elif re.search(r'(tarjima|perevod|pervod|translate|tarjma|o\'gir)', text):
        await update.message.reply_text(t("opening_service", lang, service="Tarjima"))
        await translate_handler(update, context)
        return
    
    # Spell Check (imlo, xato, grammatika)
    elif re.search(r'(imlo|xato|tekshir|grammatika)', text):
        await update.message.reply_text(t("opening_service", lang, service="Imlo tekshirish"))
        await spell_check_handler(update, context)
        return

    # 3. Fallback
    await update.message.reply_text(t("unknown_cmd", lang), reply_markup=get_main_menu(uid, lang))

async def handle_router_doc(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not update.effective_chat or update.effective_chat.type != "private":
        return
    if not await unified_router_check(update, context): return
    if await handle_premium_screenshot(update, context):
        return
    if await process_admin_state_input(update, context): return
    
    state = context.user_data.get('waiting_for')
    transliterate_dir = context.user_data.get('transliterate_direction')
    translate_dir = context.user_data.get('translate_direction')  # e.g. 'ru_uz', 'uz_en'
    uid = update.effective_user.id
    
    # Transliterate mode
    if state == 'translit_content' or transliterate_dir:
        await process_transliterate(update, context)
        return
    # Translate mode — detected by presence of translate_direction key
    elif translate_dir or state == 'translate_input':
        await process_translate_doc(update, context)
        increment_file_count(uid, "Translate Doc")
    elif state == 'spell_check_doc' or state == 'spellcheck_file':
        await process_spell_check(update, context)
        increment_file_count(uid, "Spell Check")
    elif state == 'ocr_image' or state == 'ocr_image_doc':
        # Some users send images as documents
        await process_ocr_image(update, context)
        increment_file_count(uid, "OCR Doc-Image")
    elif state == 'feedback':
        await handle_feedback(update, context)
    else:
        # Smart Logic
        await handle_smart_document(update, context)

async def handle_router_photo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not update.effective_chat or update.effective_chat.type != "private":
        return
    if not await unified_router_check(update, context): return
    if await handle_premium_screenshot(update, context):
        return
    if await process_admin_state_input(update, context): return
    
    state = context.user_data.get('waiting_for')
    uid = update.effective_user.id
    
    if state == 'ocr_image':
        await process_ocr_image(update, context)
        increment_file_count(uid, "OCR Image")
    elif state == 'pdf_images':
        await process_image_to_pdf(update, context)
        increment_file_count(uid, "Image to PDF")
    elif state == 'feedback':
        await handle_feedback(update, context)
    else:
        # Smart Logic (Photo)
        await handle_smart_photo(update, context)

async def handle_router_audio(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not update.effective_chat or update.effective_chat.type != "private":
        return
    if not await unified_router_check(update, context): return
    if await process_admin_state_input(update, context): return
    
    state = context.user_data.get('waiting_for')
    uid = update.effective_user.id
    
    if state == 'obyektivka_audio':
        await process_obyektivka_audio(update, context)
        increment_file_count(uid, "Obyektivka Audio")
    elif state == 'feedback':
        await handle_feedback(update, context)
    else:
        # Smart logic handles unknown audio
        await handle_smart_audio(update, context)

async def button_callback_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not await unified_router_check(update, context): return
    
    query = update.callback_query
    await query.answer()
    
    if query.data == "check_subs":
        user_id = query.from_user.id
        
        if is_premium(user_id):
            await query.message.delete()
            await query.message.reply_text("✅ Premium hisob: Obuna shart emas!", reply_markup=get_main_menu())
            return
            
        # ... rest of logic ...
        await query.message.delete()
        await query.message.reply_text("✅ Rahmat!", reply_markup=get_main_menu())


async def error_handler(update: object, context: ContextTypes.DEFAULT_TYPE) -> None:
    logger.error(msg="Exception while handling an update:", exc_info=context.error)

async def _webapp_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE, action: str):
    """Generic handler: sends inline button opening the correct webapp page."""
    if not update.effective_chat or update.effective_chat.type != "private":
        return
    from bot.handlers.start import _ACTION_MAP, WEBAPP_BASE
    from bot.services.user_service import get_user_lang
    uid = update.effective_user.id if update.effective_user else 0
    lang = get_user_lang(uid)
    page_info = _ACTION_MAP.get(action)
    if not page_info:
        await update.message.reply_text("❌ Noma'lum buyruq.") # Or t("unknown_cmd", lang)
        return
    page_file, btn_label, desc = page_info
    from telegram import WebAppInfo, InlineKeyboardMarkup, InlineKeyboardButton
    url = f"{WEBAPP_BASE}/{page_file}?telegram_id={uid}&lang={lang}"
    kb = InlineKeyboardMarkup([[InlineKeyboardButton(btn_label, web_app=WebAppInfo(url=url))]])
    await update.message.reply_text(f"🚀 <b>{desc}</b>", reply_markup=kb, parse_mode="HTML")

async def cmd_cv(u, c):        await _webapp_cmd(u, c, "cv")
async def cmd_obyektivka(u,c): await _webapp_cmd(u, c, "obyektivka")
async def cmd_ocr(u, c):       await _webapp_cmd(u, c, "ocr")
async def cmd_pdf(u, c):       await _webapp_cmd(u, c, "pdf")
async def cmd_translit(u, c):  await _webapp_cmd(u, c, "translit")
async def cmd_translate(u, c): await _webapp_cmd(u, c, "translate")
async def cmd_premium(u, c):   await premium_handler(u, c)


def setup_application():
    if not BOT_TOKEN:
        logger.error("BOT_TOKEN is missing!")
        return None

    application = ApplicationBuilder().token(BOT_TOKEN).connection_pool_size(8).build()
    
    # 1. CRM Middleware (Tracks + Checks Ban)
    application.add_handler(TypeHandler(Update, track_user), group=-1)
    # 1.1 Support group strict router (ignore all other bot features there)
    application.add_handler(MessageHandler(filters.Chat(chat_id=SUPPORT_GROUP_ID), support_group_router), group=0)

    # 2. Core Commands
    application.add_handler(CommandHandler("start",       start_command))
    application.add_handler(CommandHandler("menu",        menu_command))
    application.add_handler(CommandHandler("help",        help_command))
    # ── Feature shortcut commands (open the matching webapp page directly) ──
    application.add_handler(CommandHandler("cv",          cmd_cv))
    application.add_handler(CommandHandler("obyektivka",  cmd_obyektivka))
    application.add_handler(CommandHandler("ocr",         cmd_ocr))
    application.add_handler(CommandHandler("pdf",         cmd_pdf))
    application.add_handler(CommandHandler("translit",    cmd_translit))
    application.add_handler(CommandHandler("translate",   cmd_translate))
    application.add_handler(CommandHandler("premium",     cmd_premium))

    # Track bot block/unblock
    application.add_handler(ChatMemberHandler(chat_member_updated, ChatMemberHandler.MY_CHAT_MEMBER))
    
    # Admin Commands
    application.add_handler(CommandHandler("admin", admin_panel_command))
    application.add_handler(CommandHandler("stats", stats_command))
    application.add_handler(CommandHandler("send", broadcast_command))
    application.add_handler(CommandHandler("user_info", user_info_command))
    application.add_handler(CommandHandler("users", user_info_command)) 
    application.add_handler(CommandHandler("top", top_users_command))
    application.add_handler(CommandHandler("search", search_command))
    application.add_handler(CommandHandler("ban", ban_user_command))
    application.add_handler(CommandHandler("unban", unban_user_command))
    
    application.add_handler(CommandHandler("add_channel", add_channel_command))
    application.add_handler(CommandHandler("remove_channel", remove_channel_command))
    application.add_handler(CommandHandler("add_premium", add_premium_command))
    application.add_handler(CommandHandler("remove_premium", remove_premium_command))
    application.add_handler(CommandHandler("approve", approve_premium_command))
    application.add_handler(CommandHandler("set_limit", set_limit_command))
    application.add_handler(CommandHandler("add_admin", add_admin_command))
    application.add_handler(CommandHandler("remove_admin", remove_admin_command))

    # 3. Callback Queries
    application.add_handler(CallbackQueryHandler(
        premium_callback_handler,
        pattern="^prem_"
    ))
    application.add_handler(CallbackQueryHandler(
        premium_purchase_callback,
        pattern="^buy_"
    ))
    application.add_handler(CallbackQueryHandler(
        premium_payment_review_callback,
        pattern=r"^prempay_(approve|reject)_\d+$"
    ))
    
    # Language callback handler removed — bot uses Uzbek by default
    
    application.add_handler(CallbackQueryHandler(smart_callback_handler, pattern="^smart_"))
    application.add_handler(CallbackQueryHandler(translit_direction_callback, pattern="^trl_"))
    application.add_handler(CallbackQueryHandler(support_panel_callback, pattern="^support_"))
    application.add_handler(CallbackQueryHandler(button_callback_handler))

    # 4. Text Menu Navigation — Asosiy tugmalar
    application.add_handler(MessageHandler(filters.Regex(get_regex_for_key("back_to_menu")), back_to_main_menu))
    application.add_handler(MessageHandler(filters.Regex("^(🔙 Orqaga|🔙 Назад|🔙 Back|🔙 Оркага)$"), back_to_main_menu))
    application.add_handler(MessageHandler(filters.Regex(get_regex_for_key("btn_more")) & filters.ChatType.PRIVATE, more_menu_handler))
    application.add_handler(MessageHandler(filters.Regex(get_regex_for_key("btn_cv")) & filters.ChatType.PRIVATE, cv_handler))
    
    admin_buttons = "^(📊 Statistika|📨 Xabar yuborish|📢 Kanallar|💎 Premium Boshqaruv|⚙️ Sozlamalar|👥 Foydalanuvchilar|➕ Admin qo'shish|❌ Admin o'chirish|🆘 Support so'rovlar|🚪 Panelni yopish)$"
    application.add_handler(MessageHandler(filters.Regex(admin_buttons), handle_admin_text))

    application.add_handler(MessageHandler(filters.Regex(get_regex_for_key("btn_ocr")), ocr_handler))
    application.add_handler(MessageHandler(filters.Regex(get_regex_for_key("btn_oby")) & filters.ChatType.PRIVATE, obyektivka_handler))
    application.add_handler(MessageHandler(filters.Regex(get_regex_for_key("btn_translit")), transliterate_handler))
    application.add_handler(MessageHandler(filters.Regex(get_regex_for_key("btn_translate")), translate_handler))
    application.add_handler(MessageHandler(filters.Regex(get_regex_for_key("btn_pdf")), image_to_pdf_handler))
    application.add_handler(MessageHandler(filters.Regex(get_regex_for_key("btn_spell")), spell_check_handler))
    application.add_handler(MessageHandler(filters.Regex(get_regex_for_key("btn_premium")) & filters.ChatType.PRIVATE, premium_info_handler))
    application.add_handler(MessageHandler(filters.Regex("^Premium sotib olish$") & filters.ChatType.PRIVATE, premium_info_handler))
    
    application.add_handler(MessageHandler(filters.Regex("^(🔡 Kirill → Lotin|🔡 Кирилл → Лотин)$"), krill_to_lotin_handler))
    application.add_handler(MessageHandler(filters.Regex("^(🔠 Lotin → Kirill|🔠 Лотин → Кирилл)$"), lotin_to_krill_handler))

    async def go_translate(update: Update, context: ContextTypes.DEFAULT_TYPE):
        text = update.message.text
        direction = "uz_en"
        if "O'zbek → Ingliz" in text: direction = "uz_en"
        elif "Ingliz → O'zbek" in text: direction = "en_uz"
        elif "Rus → O'zbek" in text: direction = "ru_uz"
        elif "O'zbek → Rus" in text: direction = "uz_ru"
        elif "Rus → Ingliz" in text: direction = "ru_en"
        # Clear any stale state before starting new translation session
        context.user_data.pop('waiting_for', None)
        await set_translation_direction(update, context, direction)

    application.add_handler(MessageHandler(
        filters.Regex("(O'zbek → Ingliz|Ingliz → O'zbek|Rus → O'zbek|O'zbek → Rus|Rus → Ingliz)"),
        go_translate
    ))
    
    application.add_handler(MessageHandler(filters.Regex(get_regex_for_key("btn_balance")) & filters.ChatType.PRIVATE, balance_handler))
    application.add_handler(MessageHandler(filters.Regex(get_regex_for_key("btn_contact")), start_feedback))
    application.add_handler(MessageHandler(filters.Regex(get_regex_for_key("btn_help")), help_button_handler))

    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_router_text))
    application.add_handler(MessageHandler(filters.Document.ALL, handle_router_doc))
    application.add_handler(MessageHandler(filters.PHOTO, handle_router_photo))
    application.add_handler(MessageHandler(filters.VOICE | filters.AUDIO, handle_router_audio))
    application.add_handler(MessageHandler(filters.StatusUpdate.WEB_APP_DATA, web_app_data_handler))

    async def handle_router_other(update: Update, context: ContextTypes.DEFAULT_TYPE):
        if not await unified_router_check(update, context): return
        if await process_admin_state_input(update, context): return
    application.add_handler(MessageHandler(filters.ALL & ~filters.COMMAND, handle_router_other))

    application.add_error_handler(error_handler)
    return application

def main():
    application = setup_application()
    if application:
        logger.info("✅ Bot is starting in POLLING mode...")
        application.run_polling(allowed_updates=Update.ALL_TYPES, drop_pending_updates=True)

if __name__ == '__main__':
    main()
