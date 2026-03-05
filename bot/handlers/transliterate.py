"""
Kirill ↔ Lotin Transliteration Handler

Correct flow:
  Step 1: User presses "Krill-Lotin"
          Bot: "Send text or a document."

  Step 2: User sends text OR document
          Bot stores it in user_data, then shows inline buttons:
            [Kirill → Lotin]  [Lotin → Kirill]

  Step 3: User picks a direction (inline callback)
          Bot processes and returns ONLY the result (one-click copyable).
          For documents: sends converted DOCX file.
"""
import os
import time
import logging
from telegram import Update, InputFile, InlineKeyboardMarkup, InlineKeyboardButton
from telegram.ext import ContextTypes
from telegram.constants import ChatAction
from bot.keyboards.reply_keyboards import get_main_menu, get_back_button
from bot.services.transliterate_service import transliterate
from bot.services.user_service import get_user_lang, increment_file_count

logger = logging.getLogger(__name__)

# ── Inline keyboard shown after user sends content ───────────────────────────

def _direction_keyboard() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([
        [
            InlineKeyboardButton("🔡 Kirill → Lotin", callback_data="trl_krill_to_lotin"),
            InlineKeyboardButton("🔠 Lotin → Kirill", callback_data="trl_lotin_to_krill"),
        ]
    ])


# ── Step 1: User pressed "Krill-Lotin" button ────────────────────────────────

async def transliterate_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Step 1 — ask user to send text or document."""
    context.user_data.pop("trl_text", None)
    context.user_data.pop("trl_file_id", None)
    context.user_data.pop("trl_file_name", None)

    await update.message.reply_text(
        "🔤 *Kirill ↔ Lotin*\n\n"
        "Matn yozing yoki DOCX hujjat yuboring.",
        parse_mode="Markdown",
        reply_markup=get_back_button(),
    )
    context.user_data["waiting_for"] = "translit_content"


# ── Step 2: Receive text ──────────────────────────────────────────────────────

async def process_transliteration(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Called by handle_router_text / handle_router_doc.
    Stores the content and shows direction inline buttons.
    """
    msg = update.message

    # ── Text input ──────────────────────────────────────────────────────────
    if msg.text:
        text = msg.text.strip()
        if not text:
            return

        context.user_data["trl_text"] = text
        context.user_data.pop("trl_file_id", None)
        context.user_data.pop("trl_file_name", None)

        # Store the prompt message id so we can delete it later
        sent = await msg.reply_text(
            "↔️ Yo'nalishni tanlang:",
            reply_markup=_direction_keyboard(),
        )
        context.user_data["trl_prompt_id"] = sent.message_id
        return

    # ── Document input ───────────────────────────────────────────────────────
    if msg.document:
        file_name = msg.document.file_name or "document.docx"
        if not file_name.lower().endswith(".docx"):
            await msg.reply_text(
                "❌ Faqat *.DOCX* fayl qabul qilinadi.",
                parse_mode="Markdown",
                reply_markup=get_back_button(),
            )
            return

        context.user_data["trl_file_id"] = msg.document.file_id
        context.user_data["trl_file_name"] = file_name
        context.user_data.pop("trl_text", None)

        sent = await msg.reply_text(
            f"📄 *{file_name}* qabul qilindi.\n\n↔️ Yo'nalishni tanlang:",
            parse_mode="Markdown",
            reply_markup=_direction_keyboard(),
        )
        context.user_data["trl_prompt_id"] = sent.message_id
        return


# ── Step 3: Direction chosen via inline callback ──────────────────────────────

async def translit_direction_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Inline button callback — process and return result."""
    query = update.callback_query
    await query.answer()

    direction = query.data.replace("trl_", "")   # "krill_to_lotin" or "lotin_to_krill"
    uid = query.from_user.id
    lang = get_user_lang(uid)

    trl_text   = context.user_data.pop("trl_text", None)
    file_id    = context.user_data.pop("trl_file_id", None)
    file_name  = context.user_data.pop("trl_file_name", None)
    context.user_data.pop("waiting_for", None)

    # Delete the direction-picker message so chat stays clean
    try:
        await query.message.delete()
    except Exception:
        pass

    # ── Text conversion ───────────────────────────────────────────────────────
    if trl_text:
        try:
            result = transliterate(trl_text, direction)
        except Exception as e:
            await context.bot.send_message(
                chat_id=uid,
                text=f"❌ Xatolik: {e}",
                reply_markup=get_back_button(lang),
            )
            return

        # Send ONLY the result — no prefix words — for one-click copy
        await context.bot.send_message(
            chat_id=uid,
            text=result,
            reply_markup=get_back_button(lang),
        )
        increment_file_count(uid, "Transliterate Text")
        return

    # ── Document conversion ───────────────────────────────────────────────────
    if file_id and file_name:
        status_msg = await context.bot.send_message(
            chat_id=uid,
            text=f"⏳ *{file_name}* aylantirılmoqda...",
            parse_mode="Markdown",
        )
        await context.bot.send_chat_action(chat_id=uid, action=ChatAction.UPLOAD_DOCUMENT)

        temp_path   = None
        output_path = None
        try:
            from docx import Document

            file_obj  = await context.bot.get_file(file_id)
            temp_path = f"temp_translit_{uid}_{int(time.time())}.docx"
            await file_obj.download_to_drive(temp_path)

            doc = Document(temp_path)

            # Paragraphs
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.text:
                        run.text = transliterate(run.text, direction)

            # Tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if run.text:
                                    run.text = transliterate(run.text, direction)

            dir_label   = "lotin" if direction == "krill_to_lotin" else "krill"
            output_path = f"Translit_{dir_label}_{file_name}"
            doc.save(output_path)

            with open(output_path, "rb") as f:
                await context.bot.send_document(
                    chat_id=uid,
                    document=InputFile(f, filename=output_path),
                    caption=f"✅ Tayyor!",
                    reply_markup=get_back_button(lang),
                )
            increment_file_count(uid, "Transliterate Doc")
            await status_msg.delete()

        except Exception as e:
            logger.error(f"Translit DOCX error: {e}", exc_info=True)
            await status_msg.edit_text(f"❌ Xatolik: {e}")
        finally:
            for p in [temp_path, output_path]:
                if p and os.path.exists(p):
                    try:
                        os.remove(p)
                    except Exception:
                        pass
        return

    # Nothing stored — prompt user again
    await context.bot.send_message(
        chat_id=uid,
        text="❓ Avval matn yoki hujjat yuboring.",
        reply_markup=get_back_button(lang),
    )


# ── Legacy stubs kept so main.py imports don't break ─────────────────────────
# (main.py still imports these names; they are now no-ops / redirect)

async def krill_to_lotin_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Legacy — no longer used in the new flow."""
    await transliterate_handler(update, context)


async def lotin_to_krill_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Legacy — no longer used in the new flow."""
    await transliterate_handler(update, context)
