"""
OCR to Word AI Handler (HTML Table Support)
"""
import os
import time
import logging
import asyncio
from telegram import Update
from telegram.ext import ContextTypes
from telegram.constants import ParseMode, ChatAction
from docx import Document
from bs4 import BeautifulSoup
from bot.keyboards.reply_keyboards import get_back_button, get_main_menu, get_ocr_to_word_keyboard
from bot.utils.helpers import is_back_button
from bot.services.plan_limits import CAT_OCR
from bot.services.user_service import get_user_lang, record_service_completion
from bot.services.usage_tracker import ensure_can_use_or_notify
from bot.services.ocr_service import extract_text_from_image
from bot.utils.progress import send_progress, update_progress
from bot.utils.delivery import send_docx_with_confirmation

from docx.shared import Cm
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

from bs4.element import NavigableString, Tag


def _style_dict(style_text: str) -> dict[str, str]:
    out: dict[str, str] = {}
    for chunk in (style_text or "").split(";"):
        if ":" not in chunk:
            continue
        k, v = chunk.split(":", 1)
        key = (k or "").strip().lower()
        val = (v or "").strip()
        if key:
            out[key] = val
    return out


def _parse_percent(raw: str) -> float | None:
    s = (raw or "").strip().lower()
    if not s or "%" not in s:
        return None
    try:
        return float(s.replace("%", "").strip())
    except Exception:
        return None


def _parse_px(raw: str) -> float | None:
    s = (raw or "").strip().lower()
    if not s:
        return None
    if s.startswith("calc("):
        return None
    if s.endswith("px"):
        s = s[:-2].strip()
    try:
        return float(s)
    except Exception:
        return None


def _add_layout_html_to_docx(doc, layout_root) -> bool:
    """
    data-ocr-layout HTML ni DOCX ga approximate joylashuv bilan o‘tkazadi.
    Absolute CSS to‘liq qo‘llanmasa ham, line/indent/spacing saqlanadi.
    """
    nodes = []
    for el in layout_root.find_all("div", recursive=False):
        text = (el.get_text(" ", strip=False) or "").replace("\xa0", " ")
        if not text.strip():
            continue
        st = _style_dict(el.get("style", ""))
        left = _parse_percent(st.get("left", ""))
        top = _parse_percent(st.get("top", ""))
        width = _parse_percent(st.get("width", ""))
        font_px = _parse_px(st.get("font-size", ""))
        if left is None or top is None:
            continue
        nodes.append(
            {
                "text": text.strip(),
                "left": left,
                "top": top,
                "width": width if width is not None else 100.0,
                "font_px": font_px,
            }
        )

    if not nodes:
        return False

    nodes.sort(key=lambda n: (n["top"], n["left"]))
    heights = [max(6.0, float(n["font_px"] or 11.0) * 1.15) for n in nodes]
    median_h = sorted(heights)[len(heights) // 2] if heights else 12.0
    top_gap_threshold = max(0.75, median_h * 0.06)

    lines = []
    current = []
    cur_top = None
    for n in nodes:
        if not current:
            current = [n]
            cur_top = n["top"]
            continue
        if abs(float(n["top"]) - float(cur_top)) <= top_gap_threshold:
            current.append(n)
            cur_top = (float(cur_top) + float(n["top"])) / 2.0
        else:
            lines.append(current)
            current = [n]
            cur_top = n["top"]
    if current:
        lines.append(current)

    section = doc.sections[0]
    content_width_cm = float(section.page_width - section.left_margin - section.right_margin) / 360000.0
    prev_line_top = None
    for line in lines:
        line = sorted(line, key=lambda n: n["left"])
        p = doc.add_paragraph()
        min_left = max(0.0, min(float(n["left"]) for n in line))
        p.paragraph_format.left_indent = Cm(content_width_cm * (min_left / 100.0))

        if prev_line_top is not None:
            dy = max(0.0, float(line[0]["top"]) - float(prev_line_top))
            if dy > 1.8:
                p.paragraph_format.space_before = Pt(min(28.0, dy * 0.95))
        prev_line_top = float(line[0]["top"])

        prev_right = min_left
        for idx, n in enumerate(line):
            gap = max(0.0, float(n["left"]) - float(prev_right))
            if idx > 0 and gap > 3.0:
                tab_count = max(1, int(gap / 7.5))
                for _ in range(tab_count):
                    p.add_run("\t")
            run = p.add_run(str(n["text"]))
            fsz = n.get("font_px")
            if fsz:
                run.font.size = Pt(max(8.0, min(28.0, float(fsz))))
            prev_right = float(n["left"]) + float(n.get("width", 0.0))
    return True


def _schedule_ocr_auto_process(
    bot,
    chat_id: int,
    user_id: int,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """After the last photo, wait ~3s; if no new photos, start batch OCR (debounced)."""
    old = context.user_data.get("_ocr_debounce_task")
    if old and not old.done():
        old.cancel()

    async def _job():
        try:
            await asyncio.sleep(1.2)
            if context.user_data.get("waiting_for") != "ocr_image":
                return
            imgs = list(context.user_data.get("ocr_images") or [])
            if not imgs:
                return
            if not await ensure_can_use_or_notify(
                bot,
                chat_id,
                user_id,
                category=CAT_OCR,
                lang=get_user_lang(user_id),
            ):
                return
            context.user_data["ocr_images"] = []
            _run_ocr_batch_background(bot, chat_id, user_id, imgs, context.user_data)
            await bot.send_message(
                chat_id=chat_id,
                text=(
                    f"⏳ {len(imgs)} ta rasm avtomatik qayta ishlanmoqda.\n"
                    "Bir nechta rasm yuborganingizda, oxirgi rasmdan keyin ~1 soniya kutamiz."
                ),
            )
        except asyncio.CancelledError:
            return
        except Exception as e:
            logger.error("OCR debounce error: %s", e, exc_info=True)

    context.user_data["_ocr_debounce_task"] = asyncio.create_task(_job())


def _add_run_with_style(paragraph_obj, element, bold=False, italic=False, underline=False):
    """Recursively parses HTML elements and adds styled runs to a docx paragraph."""
    is_bold = bold or element.name in ['b', 'strong', 'h1', 'h2', 'h3', 'th']
    is_italic = italic or element.name in ['i', 'em']
    is_underline = underline or element.name in ['u']
    
    for child in element.children:
        if isinstance(child, NavigableString):
            raw = str(child).replace('\r\n', '\n').replace('\r', '\n')
            if not raw:
                continue
            parts = raw.split('\n')
            for idx, part in enumerate(parts):
                text = part
                if not text.strip() and text:
                    text = ' '
                if text:
                    run = paragraph_obj.add_run(text)
                    run.bold = is_bold
                    run.italic = is_italic
                    run.underline = is_underline
                if idx < len(parts) - 1:
                    paragraph_obj.add_run().add_break()
        elif isinstance(child, Tag):
            # If we hit block elements inside text contexts, just add a line break
            if child.name in ['br', 'p', 'div'] and child.name != 'br':
                paragraph_obj.add_run().add_break()
                _add_run_with_style(paragraph_obj, child, is_bold, is_italic, is_underline)
            elif child.name == 'br':
                paragraph_obj.add_run().add_break()
            else:
                _add_run_with_style(paragraph_obj, child, is_bold, is_italic, is_underline)

def get_alignment(element):
    """Extract alignment from align attribute or inline style."""
    align_str = element.get('align', '')
    style_str = element.get('style', '')
    if not align_str and style_str:
        style_lower = style_str.lower()
        if 'text-align: center' in style_lower or 'text-align:center' in style_lower: align_str = 'center'
        elif 'text-align: right' in style_lower or 'text-align:right' in style_lower: align_str = 'right'
        elif 'text-align: justify' in style_lower or 'text-align:justify' in style_lower: align_str = 'justify'
    return align_str

def add_html_to_docx(doc, html_content):
    """Parses HTML and maps it to Word layout (tables, widths, alignment, inline fonts, lists)"""
    # python-docx Document doim kamida bitta section beradi — jadval kengligi uchun shart
    section = doc.sections[0]
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)

    soup = BeautifulSoup(html_content, 'html.parser')
    root = soup.body if soup.body else soup

    layout_root = root.find(attrs={"data-ocr-layout": "1"}) or root.find("div", class_="ocr-visual")
    if layout_root is not None and _add_layout_html_to_docx(doc, layout_root):
        return
    
    def apply_align(p, align_str):
        if not align_str: return
        align_str = align_str.lower()
        if 'center' in align_str: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif 'right' in align_str: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif 'justify' in align_str: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for element in root.children:
        if isinstance(element, NavigableString):
            raw = str(element).replace('\r\n', '\n').replace('\r', '\n').strip('\n')
            if raw.strip():
                for line in raw.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())
            continue
            
        if element.name == 'table':
            rows = element.find_all('tr', recursive=False)
            if not rows and element.tbody:
                rows = element.tbody.find_all('tr', recursive=False)
            if not rows: continue
            
            # Count max cols exactly
            max_cols = 0
            for row in rows:
                cols = row.find_all(['td', 'th'], recursive=False)
                if len(cols) > max_cols: max_cols = len(cols)
            
            if max_cols > 0:
                table = doc.add_table(rows=len(rows), cols=max_cols)
                table.style = 'Table Grid'
                table.autofit = False
                table.allow_autofit = False
                
                total_width = section.page_width - section.left_margin - section.right_margin
                
                # Try to apply widths from first row
                if len(rows) > 0:
                    first_row_cols = rows[0].find_all(['td', 'th'], recursive=False)
                    for j, col in enumerate(first_row_cols):
                        width_attr = col.get('width', '').replace('%', '')
                        if width_attr and width_attr.isdigit() and j < max_cols:
                            percent = int(width_attr)
                            width_val = total_width * (percent / 100)
                            for r_idx in range(len(rows)):
                                try: table.cell(r_idx, j).width = width_val
                                except: pass

                # Fill data
                for i, row in enumerate(rows):
                    cols = row.find_all(['td', 'th'], recursive=False)
                    for j, col in enumerate(cols):
                        if j < max_cols:
                            cell = table.cell(i, j)
                            # Clear default text run
                            p = cell.paragraphs[0]
                            p.text = ""
                            
                            align = get_alignment(col)
                            apply_align(p, align)
                            _add_run_with_style(p, col)
        
        elif element.name in ['p', 'h1', 'h2', 'h3', 'h4', 'div', 'center', 'article', 'section', 'main', 'header', 'footer']:
            style = 'Normal'
            if element.name in ['h1', 'h2', 'h3']:
                style = f"Heading {element.name[-1]}"
            
            p = doc.add_paragraph(style=style)
            align = get_alignment(element)
            if element.name == 'center': align = 'center'
            apply_align(p, align)
            _add_run_with_style(p, element)
            
        elif element.name in ['ul', 'ol']:
            style = 'List Bullet' if element.name == 'ul' else 'List Number'
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(style=style)
                _add_run_with_style(p, li)
                
        elif element.name == 'br':
            doc.add_paragraph()
            
        elif element.name and element.name not in ['html', 'body', 'head', 'style', 'script', 'title', 'meta']:
            # For unrecognized wrappers, process their children directly
            for child in element.children:
                if isinstance(child, NavigableString):
                    text = str(child).strip()
                    if text:
                        doc.add_paragraph(text)
                elif isinstance(child, Tag):
                    style = 'Normal'
                    p = doc.add_paragraph()
                    align = get_alignment(child)
                    apply_align(p, align)
                    _add_run_with_style(p, child)


async def perform_ocr_and_send(context, image_path, chat_id, user_id):
    """
    Reusable function: Takes image path, performs OCR, creates Word doc, and sends it.
    Runs fully async; safe to call from a background task.
    """
    t0 = time.perf_counter()
    logger.info("OCR task started for user_id=%s chat_id=%s", user_id, chat_id)
    lang = get_user_lang(user_id)
    if not await ensure_can_use_or_notify(
        context.bot, chat_id, user_id, category=CAT_OCR, lang=lang
    ):
        return
    progress_msg = await send_progress(context, chat_id, "Jarayon boshlandi...")
    doc_path = None

    try:
        await update_progress(context, progress_msg, 20, "AI matnni o'qimoqda...")
        # Extract Text (HTML format)
        extracted_text = await extract_text_from_image(image_path)
        logger.info("OCR extract done in %.1fs user_id=%s", time.perf_counter() - t0, user_id)
        
        if not extracted_text:
            await progress_msg.edit_text("❌ **Xatolik:** Matn ajratilmadi.")
            return
            
        await update_progress(context, progress_msg, 70, "Word hujjat shakllantirilmoqda...")
        # Create Word Document asynchronously so we don't block the loop
        doc_path = f"Ocr_Natija_{user_id}_{int(time.time())}_@DastyorAiBot.docx"
        def create_and_save_doc(html_text, path):
            doc = Document()
            try:
                add_html_to_docx(doc, html_text)
            except Exception as parse_err:
                logger.error(f"HTML Parse error: {parse_err}")
                doc.add_paragraph(str(html_text))
            doc.save(path)

        await asyncio.to_thread(create_and_save_doc, extracted_text, doc_path)
        
        await update_progress(context, progress_msg, 90, "Fayl yuborilmoqda...")
        
        # Send Document
        with open(doc_path, 'rb') as f:
            ok = await send_docx_with_confirmation(
                context.bot,
                chat_id,
                f,
                filename=doc_path,
                caption="✅ **Marhamat!**\n\nSizning hujjatingiz tayyor.",
                parse_mode=ParseMode.MARKDOWN,
                reply_markup=get_main_menu(user_id, get_user_lang(user_id)),
            )
            if not ok:
                return

        record_service_completion(user_id, CAT_OCR, "OCR Image")
        await progress_msg.delete()
        # CLEAR STATE AFTER SUCCESS (when run from background task, user_data is shared)
        if getattr(context, "user_data", None) and context.user_data.get("waiting_for") == "ocr_image":
            context.user_data.pop("waiting_for", None)
        logger.info("OCR task completed in %.1fs user_id=%s", time.perf_counter() - t0, user_id)
    except Exception as e:
        logger.error("OCR Error user_id=%s: %s", user_id, e, exc_info=True)
        try:
            await progress_msg.edit_text(f"❌ **Xatolik yuz berdi:** {str(e)}")
        except Exception:
            pass
        
    finally:
        # Cleanup
        try:
            if doc_path and os.path.exists(doc_path):
                os.remove(doc_path)
        except Exception: pass


async def ocr_to_word_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Start OCR process: collect images then process on 'Tayyor'."""
    context.user_data["waiting_for"] = "ocr_image"
    context.user_data["ocr_images"] = []
    uid = update.effective_user.id if update.effective_user else None
    lang = get_user_lang(uid) if uid else "uz_lat"

    msg = (
        "📜 **Hujjat rasmi → Word AI** ✨\n\n"
        "Rasmlarni yuboring (1–20 ta).\n"
        "Pastdagi **«✅ Tayyor — Word yaratish»** tugmasini bosing yoki *tayyor* deb yozing.\n"
        "Bitta yoki bir nechta rasm yuborganingizdan keyin ~3 soniya kutib, avtomatik ham boshlanadi."
    )
    await update.message.reply_text(
        msg,
        reply_markup=get_ocr_to_word_keyboard(lang),
        parse_mode=ParseMode.MARKDOWN,
    )


def _run_ocr_background(
    bot, chat_id: int, user_id: int, temp_image_path: str, user_data: dict
) -> None:
    """
    Run OCR in a fire-and-forget background task. Does NOT block the event loop.
    Cleans up temp file and updates user_data on completion.
    """
    async def _task():
        try:
            # Build a minimal context-like object for progress/send (no full Update)
            class _Ctx:
                def __init__(self, b, ud):
                    self.bot = b
                    self.user_data = ud
            ctx = _Ctx(bot, user_data)
            await perform_ocr_and_send(ctx, temp_image_path, chat_id, user_id)
        except Exception as e:
            logger.error(f"OCR background task failed: {e}", exc_info=True)
            try:
                await bot.send_message(
                    chat_id=chat_id,
                    text=f"❌ **OCR xatolik:** {str(e)}",
                    parse_mode=ParseMode.MARKDOWN,
                )
            except Exception:
                pass
        finally:
            try:
                if temp_image_path and os.path.exists(temp_image_path):
                    os.remove(temp_image_path)
            except Exception:
                pass

    asyncio.create_task(_task())


async def _perform_ocr_batch_and_send(context, bot, chat_id: int, user_id: int, file_ids: list) -> None:
    """
    Download all files, run OCR on each with progress (e.g. "Processing 3/10"),
    merge HTML into one Word doc, send. Runs in background; cleans up temp files.
    """
    t0 = time.perf_counter()
    n = len(file_ids)
    logger.info("OCR batch started user_id=%s chat_id=%s count=%s", user_id, chat_id, n)
    try:
        from bot.utils.system_tracker import track_event_fire_and_forget

        track_event_fire_and_forget(
            telegram_id=user_id,
            username=None,
            event_type="START",
            action_name="bot:ocr_batch",
            status="ok",
            metadata={"images": n},
        )
    except Exception:
        pass

    progress_msg = None
    temp_paths = []
    doc_path = None
    try:
        try:
            from bot.utils.system_tracker import track_span

            async with track_span(
                telegram_id=user_id,
                username=None,
                action_name="bot:ocr_batch",
                metadata={"images": n},
            ):
                pass
        except Exception:
            # span created in background; ignore
            pass

        progress_msg = await send_progress(context, chat_id, f"0/{n} — Yuklanmoqda...")
        temp_dir = "temp"
        os.makedirs(temp_dir, exist_ok=True)
        for i, fid in enumerate(file_ids):
            try:
                f = await bot.get_file(fid)
                ext = os.path.splitext(f.file_path or "")[1] or ".jpg"
                if not ext.startswith("."):
                    ext = "." + ext
                path = os.path.join(
                    temp_dir,
                    f"ocr_batch_{user_id}_{int(time.time())}_{i}{ext}",
                )
                await f.download_to_drive(path)
                temp_paths.append(path)
            except Exception as e:
                logger.warning("Batch download failed for file %s: %s", i, e)
        if not temp_paths:
            await progress_msg.edit_text("❌ Hech qanday rasm yuklanmadi.")
            return

        # 1 ta rasm bo'lsa batch yo'lga tushirmaymiz — tezroq single oqim.
        if len(temp_paths) == 1:
            img_path = temp_paths[0]
            await update_progress(context, progress_msg, 35, "AI matnni o'qimoqda...")
            extracted_text = await extract_text_from_image(img_path)
            if not extracted_text:
                await progress_msg.edit_text("❌ Matn ajratilmadi.")
                return
            await update_progress(context, progress_msg, 80, "Word yaratilmoqda...")
            doc_path = f"Ocr_Natija_{user_id}_{int(time.time())}_@DastyorAiBot.docx"
            def _create_single_doc():
                doc = Document()
                add_html_to_docx(doc, extracted_text)
                doc.save(doc_path)
                return doc_path

            build_timeout = max(15, int(os.getenv("OCR_DOCX_BUILD_TIMEOUT_SECONDS", "45")))
            try:
                await asyncio.wait_for(asyncio.to_thread(_create_single_doc), timeout=build_timeout)
            except asyncio.TimeoutError:
                logger.warning("OCR single-from-batch DOCX timeout (%ss), fallback plain-text user=%s", build_timeout, user_id)
                def _fallback_single():
                    doc = Document()
                    plain = BeautifulSoup(str(extracted_text or ""), "html.parser").get_text("\n", strip=True)
                    if plain:
                        for line in plain.splitlines():
                            line = line.strip()
                            if line:
                                doc.add_paragraph(line)
                    doc.save(doc_path)
                await asyncio.to_thread(_fallback_single)

            await update_progress(context, progress_msg, 95, "Yuborilmoqda...")
            with open(doc_path, "rb") as f:
                ok_send = await send_docx_with_confirmation(
                    bot, chat_id, f,
                    filename=doc_path,
                    caption="✅ **Word fayl tayyor.**",
                    parse_mode=ParseMode.MARKDOWN,
                    reply_markup=get_main_menu(user_id, get_user_lang(user_id)),
                )
            if ok_send:
                record_service_completion(user_id, CAT_OCR, "OCR Single")
            await progress_msg.delete()
            if getattr(context, "user_data", None):
                context.user_data.pop("waiting_for", None)
                context.user_data.pop("ocr_images", None)
            return

        html_parts = []
        for i, img_path in enumerate(temp_paths):
            pct = 20 + int(70 * (i + 1) / len(temp_paths))
            await update_progress(
                context, progress_msg, pct,
                f"O'qilmoqda {i + 1}/{len(temp_paths)}...",
            )
            text = await extract_text_from_image(img_path)
            if text:
                html_parts.append(f"<div class=\"page-break\">{text}</div>")
            else:
                html_parts.append("<p>[Matn ajratilmadi]</p>")

        await update_progress(context, progress_msg, 90, "Word yaratilmoqda...")
        merged_html = "<body>" + "\n".join(html_parts) + "</body>"
        doc_path = f"Ocr_Natija_{user_id}_{int(time.time())}_@DastyorAiBot.docx"
        def _create_doc():
            doc = Document()
            try:
                add_html_to_docx(doc, merged_html)
            except Exception as parse_err:
                logger.error("Batch HTML parse error: %s", parse_err)
                doc.add_paragraph(merged_html.replace("<br>", "\n").replace("</p>", "\n"))
            doc.save(doc_path)
            return doc_path

        await asyncio.to_thread(_create_doc)

        await update_progress(context, progress_msg, 95, "Yuborilmoqda...")
        with open(doc_path, "rb") as f:
            ok_send = await send_docx_with_confirmation(
                bot, chat_id, f,
                filename=doc_path,
                caption="✅ **Barcha rasmlar bitta Word faylga birlashtirildi.**",
                parse_mode=ParseMode.MARKDOWN,
                reply_markup=get_main_menu(user_id, get_user_lang(user_id)),
            )
        if ok_send:
            record_service_completion(user_id, CAT_OCR, "OCR Batch")
        await progress_msg.delete()
        if getattr(context, "user_data", None):
            context.user_data.pop("waiting_for", None)
            context.user_data.pop("ocr_images", None)
        logger.info("OCR batch completed in %.1fs user_id=%s count=%s", time.perf_counter() - t0, user_id, n)
        try:
            from bot.utils.system_tracker import track_event_fire_and_forget

            track_event_fire_and_forget(
                telegram_id=user_id,
                username=None,
                event_type="END",
                action_name="bot:ocr_batch",
                status="success",
                execution_time_ms=int((time.perf_counter() - t0) * 1000),
                metadata={"images": n},
            )
        except Exception:
            pass
    except Exception as e:
        logger.error("OCR batch error user_id=%s: %s", user_id, e, exc_info=True)
        try:
            from bot.utils.system_tracker import track_event_fire_and_forget

            track_event_fire_and_forget(
                telegram_id=user_id,
                username=None,
                event_type="ERROR",
                action_name="bot:ocr_batch",
                status="failed",
                error_message=str(e)[:2000],
                execution_time_ms=int((time.perf_counter() - t0) * 1000),
                metadata={"images": n},
            )
        except Exception:
            pass
        try:
            if progress_msg:
                await progress_msg.edit_text(f"❌ **Xatolik:** {str(e)}")
        except Exception:
            pass
        if getattr(context, "user_data", None):
            context.user_data.pop("waiting_for", None)
            context.user_data.pop("ocr_images", None)
    finally:
        for p in temp_paths:
            try:
                if os.path.exists(p):
                    os.remove(p)
            except Exception:
                pass
        if doc_path:
            try:
                if os.path.exists(doc_path):
                    os.remove(doc_path)
            except Exception:
                pass


def _run_ocr_batch_background(bot, chat_id: int, user_id: int, file_ids: list, user_data: dict) -> None:
    """Start batch OCR in background; does not block the event loop."""
    class _Ctx:
        def __init__(self, b, ud):
            self.bot = b
            self.user_data = ud
    ctx = _Ctx(bot, user_data)

    async def _task():
        await _perform_ocr_batch_and_send(ctx, bot, chat_id, user_id, file_ids)

    asyncio.create_task(_task())


async def process_ocr_tayyor(update: Update, context: ContextTypes.DEFAULT_TYPE) -> bool:
    """
    Called when user says 'Tayyor' in OCR mode. Starts batch OCR in background.
    Returns True if batch was started, False otherwise.
    """
    images = context.user_data.get("ocr_images") or []
    if not images:
        return False
    t = context.user_data.get("_ocr_debounce_task")
    if t and not t.done():
        t.cancel()
    chat_id = update.effective_chat.id
    user_id = update.effective_user.id if update.effective_user else 0
    if not await ensure_can_use_or_notify(
        context.bot,
        chat_id,
        user_id,
        category=CAT_OCR,
        lang=get_user_lang(user_id),
    ):
        return True
    context.user_data["ocr_images"] = []  # clear so we don't process twice
    _run_ocr_batch_background(context.bot, chat_id, user_id, images, context.user_data)
    await update.message.reply_text(
        f"⏳ {len(images)} ta rasm qayta ishlanmoqda. Natija tez orada yuboriladi.",
        parse_mode=ParseMode.MARKDOWN,
    )
    return True


async def handle_ocr_image(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle image upload (Direct menu usage). Downloads file then runs OCR in background."""
    message = update.message

    # Check if back button
    uid = update.effective_user.id if update.effective_user else None
    lang = get_user_lang(uid) if uid else "uz_lat"

    if message.text and is_back_button(message.text):
        context.user_data.pop("waiting_for", None)
        context.user_data.pop("ocr_images", None)
        t = context.user_data.get("_ocr_debounce_task")
        if t and not t.done():
            t.cancel()
        await update.message.reply_text(
            "🏠 **Asosiy menyuga qaytildi**",
            reply_markup=get_main_menu(uid, lang),
            parse_mode=ParseMode.MARKDOWN
        )
        return

    if not message.photo and not message.document:
        await update.message.reply_text(
            "⚠️ Iltimos, rasm yuboring (JPG yoki PNG formatda).",
            reply_markup=get_ocr_to_word_keyboard(lang),
        )
        return

    # Collect file_id (no download yet — batch will download on Tayyor)
    if message.document:
        mime = (message.document.mime_type or "").lower()
        fname = (message.document.file_name or "").lower()
        ok = mime.startswith("image/") or mime == "application/pdf" or fname.endswith(
            (".jpg", ".jpeg", ".png", ".webp", ".gif", ".pdf", ".heic", ".heif")
        )
        if mime and not ok:
            await update.message.reply_text(
                "⚠️ OCR uchun **rasm** (fotosurat sifatida) yoki **PDF** yuboring.\n"
                "Word/Excel fayllar bu yerda ishlamaydi.",
                reply_markup=get_ocr_to_word_keyboard(lang),
                parse_mode=ParseMode.MARKDOWN,
            )
            return
        file_id = message.document.file_id
    else:
        file_id = message.photo[-1].file_id

    images = context.user_data.setdefault("ocr_images", [])
    if len(images) >= 20:
        await update.message.reply_text(
            "❌ Maksimum 20 ta rasm. *Tayyor* deb yozing yoki tugmani bosing.",
            reply_markup=get_ocr_to_word_keyboard(lang),
            parse_mode=ParseMode.MARKDOWN,
        )
        return
    images.append(file_id)
    context.user_data["ocr_images"] = images

    await update.message.reply_text(
        f"✅ {len(images)} ta rasm qabul qilindi.\n\n"
        "Yana rasm yuboring, **Tayyor** tugmasini bosing yoki biroz kuting (avtomatik boshlanadi).",
        reply_markup=get_ocr_to_word_keyboard(lang),
        parse_mode=ParseMode.MARKDOWN,
    )
    _schedule_ocr_auto_process(
        context.bot,
        update.effective_chat.id,
        uid or 0,
        context,
    )
