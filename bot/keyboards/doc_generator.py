"""
CV DOCX generator.
"""

import os
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


def _as_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()


def _parse_list(value: Any) -> list[dict[str, Any]]:
    if isinstance(value, list):
        return value
    return []


def generate_cv_docx(data: dict[str, Any], output_dir: str = "temp") -> str:
    os.makedirs(output_dir, exist_ok=True)
    safe_name = (_as_text(data.get("name")) or "CV").replace(" ", "_").replace("/", "_")
    filepath = os.path.join(output_dir, f"cv_{safe_name}_@DastyorAiBot.docx")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(_as_text(data.get("name")) or "CV")
    r.bold = True
    r.font.size = Pt(18)

    spec = _as_text(data.get("spec"))
    if spec:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sub.add_run(spec)
        sr.italic = True
        sr.font.size = Pt(12)

    contacts = []
    for key in ("phone", "email", "loc"):
        val = _as_text(data.get(key))
        if val:
            contacts.append(val)
    if contacts:
        p = doc.add_paragraph(" | ".join(contacts))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)

    about = _as_text(data.get("about"))
    if about:
        h = doc.add_paragraph("HAQIDA")
        h.runs[0].bold = True
        doc.add_paragraph(about)

    skills = _as_text(data.get("skills"))
    if skills:
        h = doc.add_paragraph("KO'NIKMALAR")
        h.runs[0].bold = True
        for part in [s.strip() for s in skills.replace(",", "\n").splitlines() if s.strip()]:
            doc.add_paragraph(f"- {part}")

    works = _parse_list(data.get("works") or data.get("work_experience"))
    if works:
        h = doc.add_paragraph("ISH TAJRIBASI")
        h.runs[0].bold = True
        for item in works:
            title = _as_text(item.get("title"))
            date_s = _as_text(item.get("date") or item.get("year") or item.get("from"))
            place = _as_text(item.get("company") or item.get("place") or item.get("co"))
            desc = _as_text(item.get("desc") or item.get("description") or item.get("d"))
            head_line = " · ".join([p for p in (title, date_s) if p])
            if head_line:
                doc.add_paragraph(head_line)
            if place:
                pp = doc.add_paragraph(place)
                for r in pp.runs:
                    r.italic = True
            if desc:
                doc.add_paragraph(desc)

    education = _parse_list(data.get("education_list") or data.get("education"))
    if education:
        h = doc.add_paragraph("TA'LIM")
        h.runs[0].bold = True
        for item in education:
            title_text = _as_text(item.get("title") or item.get("name"))
            date_text = _as_text(item.get("date") or item.get("year"))
            place_text = _as_text(item.get("place") or item.get("company") or item.get("institution"))
            head_line = " · ".join([p for p in (title_text, date_text) if p])
            if head_line:
                doc.add_paragraph(head_line)
            if place_text:
                ep = doc.add_paragraph(place_text)
                for r in ep.runs:
                    r.italic = True

    langs = _parse_list(data.get("languages_list") or data.get("language_levels"))
    if langs:
        h = doc.add_paragraph("TILLAR")
        h.runs[0].bold = True
        for row in langs:
            lang_name = _as_text(row.get("lang")) or "—"
            try:
                li = int(row.get("listen") or 0)
            except (TypeError, ValueError):
                li = 0
            try:
                rd = int(row.get("read") or 0)
            except (TypeError, ValueError):
                rd = 0
            try:
                sp = int(row.get("speak") or 0)
            except (TypeError, ValueError):
                sp = 0
            try:
                wr = int(row.get("write") or 0)
            except (TypeError, ValueError):
                wr = 0
            doc.add_paragraph(
                f"{lang_name}: tinglash {li}/6, o'qish {rd}/6, gapirish {sp}/6, yozish {wr}/6"
            )

    achievements = _parse_list(data.get("achievements_list"))
    if achievements:
        h = doc.add_paragraph("YUTUQLAR")
        h.runs[0].bold = True
        for row in achievements:
            t = _as_text(row.get("type"))
            tit = _as_text(row.get("title"))
            yr = _as_text(row.get("year"))
            line = f"{t}: {tit}" if t else tit
            if yr:
                line = f"{line} ({yr})" if line else yr
            if line:
                doc.add_paragraph(line)

    doc.save(filepath)
    return filepath
