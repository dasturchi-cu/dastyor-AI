import os
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────────────────────
# LABELS
# ─────────────────────────────────────────────────────────────────────────────
_OBY_LBL = {
    "uz_lat": {
        "title":      "MA'LUMOTNOMA",
        "birthdate":  "Tug'ilgan yili:",
        "birthplace": "Tug'ilgan joyi:",
        "nation":     "Millati:",
        "party":      "Partiyaviyligi:",
        "education":  "Ma'lumoti:",
        "graduated":  "Tamomlagan:",
        "specialty":  "Ma'lumoti bo'yicha mutaxassisligi:",
        "degree":     "Ilmiy darajasi:",
        "scititle":   "Ilmiy unvoni:",
        "langs":      "Qaysi chet tillarini biladi:",
        "military":   "Harbiy unvoni:",
        "awards":     "Davlat mukofotlari bilan taqdirlanganmi (qanaqa):",
        "deputy":     "Xalq deputatlari respublika, viloyat, shahar va tuman Kengashi deputatimi yoki boshqa saylanadigan organlarning a'zosimi (to'liq ko'rsatilishi lozim):",
        "work_title": "MEHNAT FAOLIYATI",
        "rel_suffix": "ning yaqin qarindoshlari haqida",
        "rel_sub":    "MA'LUMOTNOMA",
        "rel_c1":     "Qarindoshligi",
        "rel_c2":     "Familiyasi ismi va\notasining ismi",
        "rel_c3":     "Tug'ilgan yili\nva joyi",
        "rel_c4":     "Ish joyi va\nlavozimi",
        "rel_c5":     "Turar joyi",
        "no_work":    "Ish joylari qo'shilmagan",
        "no_rel":     "Qarindoshlar qo'shilmagan",
        "photo":      "3x4",
    },
    "uz_cyr": {
        "title":      "\u041c\u0410\u042a\u041b\u0423\u041c\u041e\u0422\u041d\u041e\u041c\u0410",
        "birthdate":  "\u0422\u0443\u0493\u0438\u043b\u0433\u0430\u043d \u0439\u0438\u043b\u0438:",
        "birthplace": "\u0422\u0443\u0493\u0438\u043b\u0433\u0430\u043d \u0436\u043e\u0439\u0438:",
        "nation":     "\u041c\u0438\u043b\u043b\u0430\u0442\u0438:",
        "party":      "\u041f\u0430\u0440\u0442\u0438\u044f\u0432\u0438\u0439\u043b\u0438\u0433\u0438:",
        "education":  "\u041c\u0430\u044a\u043b\u0443\u043c\u043e\u0442\u0438:",
        "graduated":  "\u0422\u0430\u043c\u043e\u043c\u043b\u0430\u0433\u0430\u043d:",
        "specialty":  "\u041c\u0430\u044a\u043b\u0443\u043c\u043e\u0442\u0438 \u0431\u045e\u0439\u0438\u0447\u0430 \u043c\u0443\u0442\u0430\u0445\u0430\u0441\u0441\u0438\u0441\u043b\u0438\u0433\u0438:",
        "degree":     "\u0418\u043b\u043c\u0438\u0439 \u0434\u0430\u0440\u0430\u0436\u0430\u0441\u0438:",
        "scititle":   "\u0418\u043b\u043c\u0438\u0439 \u0443\u043d\u0432\u043e\u043d\u0438:",
        "langs":      "\u049a\u0430\u0439\u0441\u0438 \u0447\u0435\u0442 \u0442\u0438\u043b\u043b\u0430\u0440\u0438\u043d\u0438 \u0431\u0438\u043b\u0430\u0434\u0438:",
        "military":   "\u04b2\u0430\u0440\u0431\u0438\u0439 \u0443\u043d\u0432\u043e\u043d\u0438:",
        "awards":     "\u0414\u0430\u0432\u043b\u0430\u0442 \u043c\u0443\u043a\u043e\u0444\u043e\u0442\u043b\u0430\u0440\u0438 \u0431\u0438\u043b\u0430\u043d \u0442\u0430\u049b\u0434\u0438\u0440\u043b\u0430\u043d\u0433\u0430\u043d\u043c\u0438 (\u049b\u0430\u043d\u0430\u049b\u0430):",
        "deputy":     "\u0425\u0430\u043b\u049b \u0434\u0435\u043f\u0443\u0442\u0430\u0442\u043b\u0430\u0440\u0438 \u0440\u0435\u0441\u043f\u0443\u0431\u043b\u0438\u043a\u0430, \u0432\u0438\u043b\u043e\u044f\u0442, \u0448\u0430\u04b3\u0430\u0440 \u0432\u0430 \u0442\u0443\u043c\u0430\u043d \u041a\u0435\u043d\u0433\u0430\u0448\u0438 \u0434\u0435\u043f\u0443\u0442\u0430\u0442\u0438\u043c\u0438 \u0451\u043a\u0438 \u0431\u043e\u0448\u049b\u0430 \u0441\u0430\u0439\u043b\u0430\u043d\u0430\u0434\u0438\u0433\u0430\u043d \u043e\u0440\u0433\u0430\u043d\u043b\u0430\u0440\u043d\u0438\u043d\u0433 \u0430\u044a\u0437\u043e\u0441\u0438\u043c\u0438 (\u0442\u045e\u043b\u0438\u049b \u043a\u045e\u0440\u0441\u0430\u0442\u0438\u043b\u0438\u0448\u0438 \u043b\u043e\u0437\u0438\u043c):",
        "work_title": "\u041c\u0415\u04b2\u041d\u0410\u0422 \u0424\u0410\u041e\u041b\u0418\u042f\u0422\u0418",
        "rel_suffix": "\u043d\u0438\u043d\u0433 \u044f\u049b\u0438\u043d \u049b\u0430\u0440\u0438\u043d\u0434\u043e\u0448\u043b\u0430\u0440\u0438 \u04b3\u0430\u049b\u0438\u0434\u0430",
        "rel_sub":    "\u041c\u0410\u042a\u041b\u0423\u041c\u041e\u0422\u041d\u041e\u041c\u0410",
        "rel_c1":     "\u049a\u0430\u0440\u0438\u043d\u0434\u043e\u0448\u043b\u0438\u0433\u0438",
        "rel_c2":     "\u0424\u0430\u043c\u0438\u043b\u0438\u044f\u0441\u0438 \u0438\u0441\u043c\u0438 \u0432\u0430\n\u043e\u0442\u0430\u0441\u0438\u043d\u0438\u043d\u0433 \u0438\u0441\u043c\u0438",
        "rel_c3":     "\u0422\u0443\u0493\u0438\u043b\u0433\u0430\u043d \u0439\u0438\u043b\u0438\n\u0432\u0430 \u0436\u043e\u0439\u0438",
        "rel_c4":     "\u0418\u0448 \u0436\u043e\u0439\u0438 \u0432\u0430\n\u043b\u0430\u0432\u043e\u0437\u0438\u043c\u0438",
        "rel_c5":     "\u0422\u0443\u0440\u0430\u0440 \u0436\u043e\u0439\u0438",
        "no_work":    "\u0418\u0448 \u0436\u043e\u0439\u043b\u0430\u0440\u0438 \u049b\u045e\u0448\u0438\u043b\u043c\u0430\u0433\u0430\u043d",
        "no_rel":     "\u049a\u0430\u0440\u0438\u043d\u0434\u043e\u0448\u043b\u0430\u0440 \u049b\u045e\u0448\u0438\u043b\u043c\u0430\u0433\u0430\u043d",
        "photo":      "3x4",
    },
}

FONT = "Times New Roman"


# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _set_margins(doc):
    for s in doc.sections:
        s.page_width    = Mm(210)
        s.page_height   = Mm(297)
        s.top_margin    = Cm(1.5)
        s.bottom_margin = Cm(1.5)
        s.left_margin   = Cm(2.0)
        s.right_margin  = Cm(1.5)


def _remove_table_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tb = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "none")
        b.set(qn("w:sz"),    "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tb.append(b)
    tblPr.append(tb)


def _add_run(para, text, bold=False, size=11):
    r = para.add_run(text)
    r.bold       = bold
    r.font.name  = FONT
    r.font.size  = Pt(size)
    return r


def _add_item(para, label, value):
    """Bold label + newline + plain value in one paragraph."""
    _add_run(para, label + "\n", bold=True, size=10)
    _add_run(para, value or "yo'q", bold=False, size=11)
    para.paragraph_format.space_after = Pt(4)


def _parse_json(val):
    if isinstance(val, list):
        return val
    try:
        return json.loads(val or "[]")
    except Exception:
        return []


# ─────────────────────────────────────────────────────────────────────────────
# MAIN GENERATOR
# ─────────────────────────────────────────────────────────────────────────────

def generate_obyektivka_docx(user_data: dict,
                              photo_path: str,
                              output_filepath: str) -> str:

    lang = user_data.get("lang", "uz_lat") or "uz_lat"
    lb   = _OBY_LBL.get(lang, _OBY_LBL["uz_lat"])
    g    = user_data.get

    doc = Document()
    _set_margins(doc)

    ns = doc.styles["Normal"]
    ns.font.name = FONT
    ns.font.size = Pt(11)
    ns.paragraph_format.space_before = Pt(0)
    ns.paragraph_format.space_after  = Pt(0)

    fullname = (g("fullname", "") or "").strip()

    # ── Title ────────────────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p_title, lb["title"], bold=True, size=14)

    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p_name, fullname, bold=True, size=14)
    p_name.paragraph_format.space_after = Pt(10)

    # ── Main info table: 3 cols (left | right | photo) ──────────────────────
    C2 = Cm(3.5)
    table = doc.add_table(rows=0, cols=3)
    table.autofit       = False
    table.allow_autofit = False

    # Set column widths
    pw_emu = (doc.sections[0].page_width
              - doc.sections[0].left_margin
              - doc.sections[0].right_margin)
    col0_w = int((pw_emu - C2) / 2)
    col1_w = int((pw_emu - C2) / 2)
    col2_w = int(C2)

    for col, w in zip(table.columns, [col0_w, col1_w, col2_w]):
        for cell in col.cells:
            tc  = cell._tc
            tcPr = tc.get_or_add_tcPr()
            old = tcPr.find(qn("w:tcW"))
            if old is not None:
                tcPr.remove(old)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(w))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)

    def add_row(l0, v0, l1, v1, span_right=False):
        row = table.add_row()
        _add_item(row.cells[0].paragraphs[0], l0, v0)
        if not span_right:
            _add_item(row.cells[1].paragraphs[0], l1, v1)
        return row

    # Row 0: birthdate | birthplace | (photo — span rows 0-2)
    r0 = add_row(lb["birthdate"],  g("birthdate","")  or "",
                 lb["birthplace"], g("birthplace","") or "")

    # Row 1: nation | party
    r1 = add_row(lb["nation"],    g("nation","")    or "",
                 lb["party"],     g("party","")     or "")

    # Row 2: education | graduated
    r2 = add_row(lb["education"], g("education","") or "",
                 lb["graduated"], g("graduated","") or "")

    # Merge photo cell rows 0-2 col 2
    photo_cell = table.cell(0, 2)
    photo_cell.merge(table.cell(2, 2))

    has_photo = photo_path and os.path.exists(photo_path)
    p_ph = photo_cell.paragraphs[0]
    p_ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    photo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if has_photo:
        r = p_ph.add_run()
        r.add_picture(photo_path, width=Cm(3), height=Cm(4))
    else:
        _add_run(p_ph, lb["photo"], bold=False, size=9)

    # Row 3: specialty — spans all 3 cols
    r3 = table.add_row()
    _add_item(r3.cells[0].paragraphs[0],
              lb["specialty"], g("specialty","") or "")
    r3.cells[0].merge(r3.cells[1]).merge(r3.cells[2])

    # Row 4: degree | scititle
    r4 = add_row(lb["degree"],   g("degree","")          or "",
                 lb["scititle"], g("scientific_title","") or "")

    # Row 5: langs | military
    r5 = add_row(lb["langs"],    g("languages","")    or "",
                 lb["military"], g("military_rank","") or "")

    # Row 6: awards — spans all
    r6 = table.add_row()
    _add_item(r6.cells[0].paragraphs[0],
              lb["awards"], g("awards","") or "")
    r6.cells[0].merge(r6.cells[1]).merge(r6.cells[2])

    # Row 7: deputy — spans all
    r7 = table.add_row()
    _add_item(r7.cells[0].paragraphs[0],
              lb["deputy"], g("deputy","") or "")
    r7.cells[0].merge(r7.cells[1]).merge(r7.cells[2])

    _remove_table_borders(table)
    doc.add_paragraph()

    # ── Work experience ──────────────────────────────────────────────────────
    wt = doc.add_paragraph(lb["work_title"])
    wt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    wt.runs[0].bold      = True
    wt.runs[0].font.size = Pt(12)

    works = _parse_json(g("work_experience", []))

    if works:
        work_table = doc.add_table(rows=0, cols=1)
        work_table.autofit       = False
        work_table.allow_autofit = False
        # Set full width
        for cell in work_table.columns[0].cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW  = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(int(pw_emu)))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)

        for work in works:
            year = (work.get("year") or work.get("years") or "").strip()
            pos  = (work.get("position") or work.get("pos") or
                    work.get("description") or "").strip()
            if not (year or pos):
                continue
            row  = work_table.add_row()
            cell = row.cells[0]
            p    = cell.paragraphs[0]
            if year:
                _add_run(p, year, bold=True, size=11)
                _add_run(p, " yy. - ", bold=False, size=11)
            _add_run(p, pos, bold=False, size=11)
            p.paragraph_format.space_after = Pt(2)

        # Add borders only to work table (box style)
        wtbl = work_table._tbl
        tblPr = wtbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            wtbl.insert(0, tblPr)
        tblBorders = OxmlElement("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"),   "single")
            b.set(qn("w:sz"),    "4")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "000000")
            tblBorders.append(b)
        tblPr.append(tblBorders)
    else:
        p_nd = doc.add_paragraph(lb["no_work"])
        p_nd.paragraph_format.left_indent = Cm(0.5)

    # ── Page break ───────────────────────────────────────────────────────────
    doc.add_page_break()

    # ── Relatives ────────────────────────────────────────────────────────────
    p_rt = doc.add_paragraph(f"{fullname} {lb['rel_suffix']}")
    p_rt.alignment       = WD_ALIGN_PARAGRAPH.CENTER
    p_rt.runs[0].bold    = True
    p_rt.runs[0].font.size = Pt(12)

    p_rs = doc.add_paragraph(lb["rel_sub"])
    p_rs.alignment       = WD_ALIGN_PARAGRAPH.CENTER
    p_rs.runs[0].bold    = True
    p_rs.runs[0].font.size = Pt(12)
    p_rs.paragraph_format.space_after = Pt(8)

    relatives = _parse_json(g("relatives", []))

    headers = [lb["rel_c1"], lb["rel_c2"], lb["rel_c3"],
               lb["rel_c4"], lb["rel_c5"]]
    rel_table = doc.add_table(rows=1, cols=5)
    rel_table.style = "Table Grid"

    # Header row
    hdr_cells = rel_table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs:
            p.runs[0].bold      = True
            p.runs[0].font.name = FONT
            p.runs[0].font.size = Pt(10)

    # Data rows
    if relatives:
        for rel in relatives:
            row = rel_table.add_row()
            vals = [
                rel.get("type","")  or rel.get("degree","") or rel.get("kin",""),
                rel.get("name","")  or rel.get("fullname","") or rel.get("fish",""),
                rel.get("birth","") or rel.get("birth_year_place",""),
                rel.get("job","")   or rel.get("work_place","") or rel.get("work",""),
                rel.get("addr","")  or rel.get("address",""),
            ]
            for i, (cell, val) in enumerate(zip(row.cells, vals)):
                cell.text = val or ""
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if p.runs:
                    p.runs[0].font.name = FONT
                    p.runs[0].font.size = Pt(10)
                    if i == 0:
                        p.runs[0].bold = True
    else:
        row  = rel_table.add_row()
        mc   = row.cells[0].merge(row.cells[4])
        mc.text = lb["no_rel"]
        p = mc.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs:
            p.runs[0].font.name = FONT
            p.runs[0].font.size = Pt(10)

    doc.save(output_filepath)
    return output_filepath


# ─────────────────────────────────────────────────────────────────────────────
# TEST
# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    sample = {
        "lang":             "uz_lat",
        "fullname":         "Aliyev Ali Aliyevich",
        "birthdate":        "22.02.2000",
        "birthplace":       "Farg'ona viloyati, Farg'ona tumani",
        "nation":           "O'zbek",
        "party":            "yo'q",
        "education":        "Oliy",
        "graduated":        "Andijon Davlat tibbiyot instituti",
        "specialty":        "Umumiy amaliyot shifokori",
        "degree":           "yo'q",
        "scientific_title": "yo'q",
        "languages":        "Ingliz tili",
        "military_rank":    "yo'q",
        "awards":           "yo'q",
        "deputy":           "yo'q",
        "work_experience": json.dumps([
            {"year": "2018-2024", "position": "Andijon davlat tibbiyot instituti davolash ishi fakulteti talabasi"},
            {"year": "2024-2025", "position": "Respublika Shoshilinch Tez Yordam Ilmiy Markazi 1-stansiya vrachi"},
            {"year": "2025 - h.v.", "position": "Andijon davlat tibbiyot instituti magistratura talabasi"},
        ]),
        "relatives": json.dumps([
            {"type": "Otasi",   "name": "Valiyev Vali Valiyevich",   "birth": "1971-yil, Farg'ona",  "job": "Nafaqada",       "addr": "Farg'ona, Ulug'nor MFY"},
            {"type": "Onasi",   "name": "Botirova Nargiza Sodiqovna","birth": "1977-yil, Farg'ona",  "job": "Uy bekasi",      "addr": "Farg'ona, Ulug'nor MFY"},
        ]),
    }

    out = generate_obyektivka_docx(sample, "", "/mnt/user-data/outputs/obyektivka_docx_test.docx")
    print(f"Tayyor: {out}")
