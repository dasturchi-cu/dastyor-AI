"""
Document Generator Service — DASTYOR AI
Pure black & white design. No colors.
"""
import os, logging
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

FONT      = "Times New Roman"
FONT_CV   = "Calibri"
BLACK     = RGBColor(0x00,0x00,0x00)
DARK      = RGBColor(0x1A,0x1A,0x1A)
GREY_TXT  = RGBColor(0x66,0x66,0x66)
GREY_DARK = RGBColor(0x44,0x44,0x44)
LABEL_DARK = RGBColor(0x22,0x22,0x22)
MUTED     = RGBColor(0xAA,0xAA,0xAA)
ITALIC_C  = RGBColor(0x77,0x77,0x77)
NAVY      = RGBColor(0x1A,0x3A,0x6B)
NAVY_HEX  = "1A3A6B"
NAVY_PALE = "E8EEF7"
MUTED_CLR = RGBColor(0x5A,0x6A,0x8A)
DARK_TEXT = RGBColor(0x0D,0x0D,0x1A)
WHITE_HEX = "FFFFFF"
BLK       = "000000"
LGT       = "CCCCCC"
FONT_SIZE = 10.5

_OBY_LABELS = {
    "uz_lat": {
        "title":"MA'LUMOTNOMA","name_label":"F.I.SH.",
        "bdate":"TUG'ILGAN YILI, KUNI VA OYI","bplace":"TUG'ILGAN JOYI",
        "nation":"MILLATI","party":"PARTIYAVIYLIGI",
        "edu":"MA'LUMOTI","grad":"TAMOMLAGAN",
        "spec":"MUTAXASSISLIGI","degree":"ILMIY DARAJASI",
        "stitle":"ILMIY UNVONI","langs":"CHET TILLARNI BILISHI",
        "military":"HARBIY/MAXSUS UNVONI","awards":"DAVLAT MUKOFOTLARI",
        "deputy":"DEPUTATLIK STATUSI","phone":"TELEFON",
        "addr":"YASHASH MANZILI",
        "work_title":"MEHNAT  FAOLIYATI",
        "yillar":"YILLARI","lavozim":"ISH JOYI VA EGALLAGAN LAVOZIMLARI",
        "rel_title":"YAQIN QARINDOSHLARI HAQIDA MA'LUMOT",
        "rel_kin":"QARINDOSHLIGI","rel_fish":"F.I.SH.",
        "rel_birth":"TUG'ILGAN YILI VA JOYI",
        "rel_work":"ISH JOYI VA LAVOZIMI","rel_addr":"YASHASH MANZILI",
        "photo":"3x4","no_data":"Ma'lumot yo'q",
    },
    "uz_cyr": {
        "title":"МАЪЛУМОТНОМА","name_label":"Ф.И.Ш.",
        "bdate":"ТУҒИЛГАН ЙИЛИ, КУНИ ВА ОЙИ","bplace":"ТУҒИЛГАН ЖОЙИ",
        "nation":"МИЛЛАТИ","party":"ПАРТИЯВИЙЛИГИ",
        "edu":"МАЪЛУМОТИ","grad":"ТАМОМЛАГАН",
        "spec":"МУТАХАССИСЛИГИ","degree":"ИЛМИЙ ДАРАЖАСИ",
        "stitle":"ИЛМИЙ УНВОНИ","langs":"ЧЕТ ТИЛЛАРНИ БИЛИШИ",
        "military":"ҲАРБИЙ/МАХСУС УНВОНИ","awards":"ДАВЛАТ МУКОФОТЛАРИ",
        "deputy":"ДЕПУТАТЛИК СТАТУСИ","phone":"ТЕЛЕФОН",
        "addr":"ЯШАШ МАНЗИЛИ",
        "work_title":"МЕҲНАТ  ФАОЛИЯТИ",
        "yillar":"ЙИЛЛАРИ","lavozim":"ИШ ЖОЙИ ВА ЭГАЛЛАГАН ЛАВОЗИМЛАРИ",
        "rel_title":"ЯҚИН ҚАРИНДОШЛАРИ ҲАҚИДА МАЪЛУМОТ",
        "rel_kin":"ҚАРИНДОШЛИГИ","rel_fish":"Ф.И.Ш.",
        "rel_birth":"ТУҒИЛГАН ЙИЛИ ВА ЖОЙИ",
        "rel_work":"ИШ ЖОЙИ ВА ЛАВОЗИМИ","rel_addr":"ЯШАШ МАНЗИЛИ",
        "photo":"3x4","no_data":"Маълумот йўқ",
    },
    "en": {
        "title":"CURRICULUM VITAE","name_label":"Full Name",
        "bdate":"DATE OF BIRTH","bplace":"PLACE OF BIRTH",
        "nation":"NATIONALITY","party":"PARTY MEMBERSHIP",
        "edu":"EDUCATION","grad":"GRADUATED FROM",
        "spec":"SPECIALTY","degree":"ACADEMIC DEGREE",
        "stitle":"ACADEMIC TITLE","langs":"FOREIGN LANGUAGES",
        "military":"MILITARY RANK","awards":"STATE AWARDS",
        "deputy":"DEPUTY STATUS","phone":"PHONE",
        "addr":"HOME ADDRESS",
        "work_title":"WORK  EXPERIENCE",
        "yillar":"PERIOD","lavozim":"POSITION & WORKPLACE",
        "rel_title":"CLOSE RELATIVES INFORMATION",
        "rel_kin":"RELATIONSHIP","rel_fish":"FULL NAME",
        "rel_birth":"YEAR & PLACE OF BIRTH",
        "rel_work":"WORKPLACE & POSITION","rel_addr":"HOME ADDRESS",
        "photo":"3x4","no_data":"No data",
    },
    "ru": {
        "title":"ОБЪЕКТИВКА","name_label":"Ф.И.О.",
        "bdate":"ГОД, ЧИСЛО И МЕСЯЦ РОЖДЕНИЯ","bplace":"МЕСТО РОЖДЕНИЯ",
        "nation":"НАЦИОНАЛЬНОСТЬ","party":"ПАРТИЙНОСТЬ",
        "edu":"ОБРАЗОВАНИЕ","grad":"ОКОНЧИЛ(А)",
        "spec":"СПЕЦИАЛЬНОСТЬ","degree":"УЧЁНАЯ СТЕПЕНЬ",
        "stitle":"УЧЁНОЕ ЗВАНИЕ","langs":"ИНОСТРАННЫЕ ЯЗЫКИ",
        "military":"ВОИНСКОЕ ЗВАНИЕ","awards":"ГОСУДАРСТВЕННЫЕ НАГРАДЫ",
        "deputy":"СТАТУС ДЕПУТАТА","phone":"ТЕЛЕФОН",
        "addr":"ДОМАШНИЙ АДРЕС",
        "work_title":"ТРУДОВАЯ  ДЕЯТЕЛЬНОСТЬ",
        "yillar":"ГОДЫ","lavozim":"МЕСТО РАБОТЫ И ДОЛЖНОСТЬ",
        "rel_title":"СВЕДЕНИЯ О БЛИЗКИХ РОДСТВЕННИКАХ",
        "rel_kin":"СТЕПЕНЬ РОДСТВА","rel_fish":"Ф.И.О.",
        "rel_birth":"ГОД И МЕСТО РОЖДЕНИЯ",
        "rel_work":"МЕСТО РАБОТЫ И ДОЛЖНОСТЬ","rel_addr":"МЕСТО ЖИТЕЛЬСТВА",
        "photo":"3x4","no_data":"Нет данных",
    },
}

_CV_SECTIONS = {
    "uz_lat":{"about":"Haqida","exp":"Ish Tajribasi","edu":"Ta'lim","skills":"Ko'nikmalar"},
    "uz_cyr":{"about":"Ҳақида","exp":"Иш Тажрибаси","edu":"Таълим","skills":"Кўникмалар"},
    "en":    {"about":"About","exp":"Experience","edu":"Education","skills":"Skills"},
    "ru":    {"about":"О себе","exp":"Опыт работы","edu":"Образование","skills":"Навыки"},
}

# ── XML helpers (1:1 preview: padding 18mm 16mm 20mm 20mm) ─────────────────────
def _margins(doc,top=1.8,bottom=2.0,left=2.0,right=1.6):
    for s in doc.sections:
        s.page_width=Mm(210);s.page_height=Mm(297)
        s.top_margin=Cm(top);s.bottom_margin=Cm(bottom)
        s.left_margin=Cm(left);s.right_margin=Cm(right)

def _pw(doc):
    s=doc.sections[0]
    return s.page_width-s.left_margin-s.right_margin

def _spc(p,before=0,after=0,line=1.15):
    pPr=p._p.get_or_add_pPr()
    sp=OxmlElement("w:spacing")
    sp.set(qn("w:before"),str(int(before*20)))
    sp.set(qn("w:after"),str(int(after*20)))
    sp.set(qn("w:line"),str(int(line*240)))
    sp.set(qn("w:lineRule"),"auto")
    pPr.append(sp)

def _run(p,text,bold=False,italic=False,size=11.0,color=None,font=FONT,lspc=None):
    r=p.add_run(text)
    r.bold=bold;r.italic=italic
    r.font.name=font;r.font.size=Pt(size)
    if color:r.font.color.rgb=color
    if lspc is not None:
        rPr=r._r.get_or_add_rPr()
        sc=OxmlElement("w:spacing")
        sc.set(qn("w:val"),str(int(lspc*20)))
        rPr.append(sc)
    return r

def _no_tbl_bdr(tbl):
    t=tbl._tbl
    tblPr=t.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr=OxmlElement("w:tblPr");t.insert(0,tblPr)
    tb=OxmlElement("w:tblBorders")
    for s in("top","left","bottom","right","insideH","insideV"):
        b=OxmlElement(f"w:{s}")
        b.set(qn("w:val"),"none");b.set(qn("w:sz"),"0")
        b.set(qn("w:space"),"0");b.set(qn("w:color"),"auto")
        tb.append(b)
    tblPr.append(tb)

def _col_w(tbl,pw,pcts):
    for ci,pct in enumerate(pcts):
        w=int(pw*pct/100)
        for cell in tbl.columns[ci].cells:
            tcPr=cell._tc.get_or_add_tcPr()
            old=tcPr.find(qn("w:tcW"))
            if old is not None:tcPr.remove(old)
            tcW=OxmlElement("w:tcW")
            tcW.set(qn("w:w"),str(w));tcW.set(qn("w:type"),"dxa")
            tcPr.append(tcW)

def _shading(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr()
    shd=OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear");shd.set(qn("w:color"),"auto")
    shd.set(qn("w:fill"),fill.lstrip("#"))
    tcPr.append(shd)

def _pad(cell,top=80,bottom=80,left=80,right=80):
    tcPr=cell._tc.get_or_add_tcPr()
    tcMar=OxmlElement("w:tcMar")
    for s,v in[("top",top),("bottom",bottom),("left",left),("right",right)]:
        m=OxmlElement(f"w:{s}");m.set(qn("w:w"),str(v));m.set(qn("w:type"),"dxa")
        tcMar.append(m)
    tcPr.append(tcMar)

def _vcenter(cell):
    tcPr=cell._tc.get_or_add_tcPr()
    va=OxmlElement("w:vAlign");va.set(qn("w:val"),"center")
    tcPr.append(va)

def _bdr_none(cell):
    tcPr=cell._tc.get_or_add_tcPr()
    tb=OxmlElement("w:tcBorders")
    for s in("top","left","bottom","right"):
        b=OxmlElement(f"w:{s}");b.set(qn("w:val"),"none");tb.append(b)
    tcPr.append(tb)

def _bdr_bottom_only(cell,color="CCCCCC",sz="4"):
    tcPr=cell._tc.get_or_add_tcPr()
    tb=OxmlElement("w:tcBorders")
    for s in ("top","left","right"):
        b=OxmlElement(f"w:{s}");b.set(qn("w:val"),"none");tb.append(b)
    bot=OxmlElement("w:bottom")
    bot.set(qn("w:val"),"single");bot.set(qn("w:sz"),sz)
    bot.set(qn("w:space"),"0");bot.set(qn("w:color"),color.lstrip("#") if hasattr(color,"lstrip") else color)
    tb.append(bot);tcPr.append(tb)

def _bdr_bottom(cell,color=LGT,sz="4"):
    tcPr=cell._tc.get_or_add_tcPr()
    tb=OxmlElement("w:tcBorders")
    for s in("top","left","right"):
        b=OxmlElement(f"w:{s}");b.set(qn("w:val"),"none");tb.append(b)
    bot=OxmlElement("w:bottom")
    bot.set(qn("w:val"),"single");bot.set(qn("w:sz"),sz)
    bot.set(qn("w:space"),"0");bot.set(qn("w:color"),color.lstrip("#"))
    tb.append(bot);tcPr.append(tb)

def _bdr_all(cell,color=BLK,sz="4"):
    tcPr=cell._tc.get_or_add_tcPr()
    tb=OxmlElement("w:tcBorders")
    for s in("top","left","bottom","right"):
        b=OxmlElement(f"w:{s}")
        b.set(qn("w:val"),"single");b.set(qn("w:sz"),sz)
        b.set(qn("w:space"),"0");b.set(qn("w:color"),color.lstrip("#"))
        tb.append(b)
    tcPr.append(tb)

def _bdr_dashed(cell,color="AAAAAA",sz="6"):
    tcPr=cell._tc.get_or_add_tcPr()
    tb=OxmlElement("w:tcBorders")
    for s in("top","left","bottom","right"):
        b=OxmlElement(f"w:{s}")
        b.set(qn("w:val"),"dashed");b.set(qn("w:sz"),sz)
        b.set(qn("w:space"),"0");b.set(qn("w:color"),color.lstrip("#"))
        tb.append(b)
    tcPr.append(tb)

def _para_bdr_bot(p,color=BLK,sz="8"):
    pPr=p._p.get_or_add_pPr()
    pBdr=OxmlElement("w:pBdr")
    bot=OxmlElement("w:bottom")
    bot.set(qn("w:val"),"single");bot.set(qn("w:sz"),sz)
    bot.set(qn("w:space"),"6");bot.set(qn("w:color"),color.lstrip("#"))
    pBdr.append(bot);pPr.append(pBdr)

# ── Components (1:1 preview: sect-title 14pt, letter-spacing 3px) ───────────────
def _sect_title(doc,text):
    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    _spc(p,before=0,after=12,line=1.0)
    _para_bdr_bot(p,BLK,"8")
    _run(p,text,bold=True,size=14.0,color=BLACK,lspc=3.0)

def _info_row_2(doc,pw,lbl1,val1,lbl2,val2):
    t=doc.add_table(rows=1,cols=2)
    _no_tbl_bdr(t);t.autofit=False
    _col_w(t,pw,[50,50])
    for j,(lbl,val) in enumerate([(lbl1,val1),(lbl2,val2)]):
        c=t.cell(0,j)
        _shading(c,WHITE_HEX)
        _bdr_none(c)
        # Match HTML-ish paddings: ~10px vertical, left column has extra right gap (~24px)
        _pad(c,top=90,bottom=70,left=0,right=360 if j==0 else 0)
        p_lbl=c.paragraphs[0]
        _spc(p_lbl,before=0,after=0,line=1.0)
        _run(p_lbl,lbl,size=9.0,color=GREY_TXT,lspc=1.0)
        p_val=c.add_paragraph()
        # Values with a bit more line height to resemble HTML line-height: 1.5
        _spc(p_val,before=2,after=0,line=1.5)
        _run(p_val,val or "",size=12.0,color=BLACK)

def _info_row_1(doc,pw,lbl,val):
    t=doc.add_table(rows=1,cols=1)
    _no_tbl_bdr(t);t.autofit=False
    _col_w(t,pw,[100])
    c=t.cell(0,0)
    _shading(c,WHITE_HEX)
    _bdr_none(c)
    _pad(c,top=90,bottom=70,left=0,right=0)
    p_lbl=c.paragraphs[0]
    _spc(p_lbl,before=0,after=0,line=1.0)
    _run(p_lbl,lbl,size=9.0,color=GREY_TXT,lspc=1.0)
    p_val=c.add_paragraph()
    _spc(p_val,before=2,after=0,line=1.5)
    _run(p_val,val or "",size=12.0,color=BLACK)

def _bdr_top_bottom(cell,col="000000",sz="8"):
    tcPr=cell._tc.get_or_add_tcPr()
    tb=OxmlElement("w:tcBorders")
    for s in ("top","bottom"):
        b=OxmlElement(f"w:{s}")
        b.set(qn("w:val"),"single");b.set(qn("w:sz"),sz)
        b.set(qn("w:space"),"0");b.set(qn("w:color"),col.lstrip("#") if isinstance(col,str) else "000000")
        tb.append(b)
    for s in ("left","right"):
        b=OxmlElement(f"w:{s}");b.set(qn("w:val"),"none");tb.append(b)
    tcPr.append(tb)

def _empty_state(doc,text):
    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _spc(p,before=12,after=12,line=1.0)
    _run(p,text,italic=True,size=11.0,color=RGBColor(0x88,0x88,0x88))

def _set_tabs_2col(p, left_col_w_cm: float):
    """Create a consistent 2-column layout using a single tab stop."""
    try:
        p.paragraph_format.tab_stops.clear_all()
    except Exception:
        pass
    p.paragraph_format.tab_stops.add_tab_stop(Cm(left_col_w_cm), alignment=WD_TAB_ALIGNMENT.LEFT)

def _set_tab_right(p, pos_cm: float):
    """Right-aligned tab stop (for photo on the right)."""
    try:
        p.paragraph_format.tab_stops.clear_all()
    except Exception:
        pass
    p.paragraph_format.tab_stops.add_tab_stop(Cm(pos_cm), alignment=WD_TAB_ALIGNMENT.RIGHT)

def _cm_to_pt(cm: float) -> float:
    return cm * 28.3464567

def _add_vml_dashed_box(run, width_cm: float, height_cm: float, text: str = "3×4"):
    """
    Insert a dashed rectangle placeholder WITHOUT tables.
    Uses VML so Word renders a real border (not gridlines).
    """
    # NOTE: Use cm units directly in VML style.
    # Some Word viewers mis-handle pt sizing in w:pict, but cm is reliable.
    # Use v:rect (more stable than v:shape in many Word viewers)
    xml = f"""
    <w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:v="urn:schemas-microsoft-com:vml"
            xmlns:o="urn:schemas-microsoft-com:office:office">
      <v:rect id="oby_photo_ph"
              style="width:{width_cm:.2f}cm;height:{height_cm:.2f}cm;v-text-anchor:middle"
              strokecolor="#CCCCCC" fillcolor="#FFFFFF">
        <v:stroke dashstyle="dash" endcap="flat"/>
        <v:textbox inset="0pt,0pt,0pt,0pt" style="mso-fit-shape-to-text:t">
          <w:txbxContent>
            <w:p>
              <w:pPr>
                <w:jc w:val="center"/>
              </w:pPr>
              <w:r>
                <w:rPr>
                  <w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}"/>
                  <w:sz w:val="18"/>
                  <w:color w:val="999999"/>
                </w:rPr>
                <w:t>{text}</w:t>
              </w:r>
            </w:p>
          </w:txbxContent>
        </v:textbox>
      </v:rect>
    </w:pict>
    """.strip()

    run._r.append(parse_xml(xml))

def _info_pair_tabs(doc, left_col_w_cm: float, lbl1: str, val1: str, lbl2: str, val2: str):
    """
    Tableless 2-col row: label line then value line (preview-like).
    Avoids Word table gridlines completely.
    """
    p_lbl = doc.add_paragraph()
    _set_tabs_2col(p_lbl, left_col_w_cm)
    _spc(p_lbl, before=0, after=0, line=1.0)
    _run(p_lbl, lbl1, bold=True, size=9.5, color=DARK, lspc=0.8)
    p_lbl.add_run("\t")
    _run(p_lbl, lbl2, bold=True, size=9.5, color=DARK, lspc=0.8)

    p_val = doc.add_paragraph()
    _set_tabs_2col(p_val, left_col_w_cm)
    _spc(p_val, before=2, after=8, line=1.5)
    _run(p_val, val1 or "", size=12.0, color=BLACK)
    p_val.add_run("\t")
    _run(p_val, val2 or "", size=12.0, color=BLACK)

def _info_single_tabs(doc, lbl: str, val: str):
    p_lbl = doc.add_paragraph()
    _spc(p_lbl, before=0, after=0, line=1.0)
    _run(p_lbl, lbl, bold=True, size=9.5, color=DARK, lspc=0.8)

    p_val = doc.add_paragraph()
    _spc(p_val, before=2, after=8, line=1.5)
    _run(p_val, val or "", size=12.0, color=BLACK)

def _pick_current_job(work_experience: list) -> tuple[str, str] | None:
    """
    Return (position, year) for the 'current' (last / present) job if found.
    Heuristics:
    - Prefer entries whose year contains 'h.v' / 'hozir' / 'present' or ends with '-' (open-ended)
    - Otherwise, fall back to the last non-empty entry.
    """
    if not work_experience:
        return None

    def _norm(s: str) -> str:
        return (s or "").strip().lower()

    candidates: list[tuple[int, str, str]] = []
    for i, w in enumerate(work_experience):
        year = (w.get("year") or w.get("years") or w.get("f") or "").strip()
        pos  = (w.get("position") or w.get("pos") or w.get("d") or "").strip()
        if not (year or pos):
            continue
        candidates.append((i, pos, year))

    if not candidates:
        return None

    for _, pos, year in candidates:
        y = _norm(year)
        if "h.v" in y or "hozir" in y or "present" in y or y.endswith("-"):
            return (pos, year)

    # fallback: last provided entry
    _, pos, year = candidates[-1]
    return (pos, year)

def _work_table(doc,pw,lb,works):
    if not works:
        return _empty_state(doc,lb["no_data"])
    t=doc.add_table(rows=1+len(works),cols=2)
    _no_tbl_bdr(t);t.autofit=False
    _col_w(t,pw,[26,74])
    for j,txt in enumerate([lb["yillar"],lb["lavozim"]]):
        c=t.cell(0,j)
        _shading(c,WHITE_HEX)
        _bdr_top_bottom(c,"000000","8")
        _pad(c,top=100,bottom=100,left=100 if j==0 else 120,right=80)
        p=c.paragraphs[0]
        _spc(p,before=0,after=0,line=1.0)
        _run(p,txt,bold=True,size=11.0,color=BLACK,lspc=1.0)
    if works:
        for i,w in enumerate(works):
            year=w.get("year","") or w.get("years","") or w.get("f","") or ""
            pos=w.get("position","") or w.get("pos","") or w.get("d","") or ""
            for j,(c,val) in enumerate([(t.cell(1+i,0),year),(t.cell(1+i,1),pos)]):
                _bdr_bottom_only(c,"CCCCCC","4")
                _pad(c,top=96,bottom=96,left=100 if j==0 else 120,right=80)
                p=c.paragraphs[0]
                _spc(p,before=0,after=0,line=1.15)
                r=_run(p,val,size=11.5,color=BLACK)
                if j==0: r.bold=True

def _rel_table(doc,pw,lb,rels):
    if not rels:
        return _empty_state(doc,lb["no_data"])
    t=doc.add_table(rows=1+len(rels),cols=5)
    _no_tbl_bdr(t);t.autofit=False
    _col_w(t,pw,[12,22,18,26,22])
    hdrs=[lb["rel_kin"],lb["rel_fish"],lb["rel_birth"],lb["rel_work"],lb["rel_addr"]]
    for j,txt in enumerate(hdrs):
        c=t.cell(0,j)
        _shading(c,"EFEFEF")
        _bdr_all(c,"000000","4")
        _pad(c,top=70,bottom=70,left=60,right=60)
        p=c.paragraphs[0]
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        _spc(p,before=0,after=0,line=1.1)
        _run(p,txt,bold=True,size=10.0,color=BLACK,lspc=0.5)
    if rels:
        for i,rel in enumerate(rels):
            vals=[
                rel.get("degree","") or rel.get("kin",""),
                rel.get("fullname","") or rel.get("fish",""),
                rel.get("birth_year_place","") or rel.get("birth",""),
                rel.get("work_place","") or rel.get("work",""),
                rel.get("address","") or rel.get("addr",""),
            ]
            for j,val in enumerate(vals):
                c=t.cell(1+i,j);_bdr_all(c,"000000","4")
                _pad(c,top=80,bottom=80,left=60,right=60)
                p=c.paragraphs[0]
                p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                _spc(p,before=0,after=0,line=1.1)
                _run(p,val or "",size=10.0,color=BLACK)

# ── MAIN FUNCTIONS ────────────────────────────────────────────────────────────
def generate_obyektivka_docx(data:dict, photo_path:str|None=None,
                              output_filepath:str|None=None,
                              output_dir:str="temp") -> str:
    lang=data.get("lang","uz_lat") or "uz_lat"
    lb=_OBY_LABELS.get(lang,_OBY_LABELS["uz_lat"])
    g=data.get
    if output_filepath is None:
        os.makedirs(output_dir,exist_ok=True)
        safe=(g("fullname","") or "unknown").replace(" ","_").replace("/","_")
        output_filepath=os.path.join(output_dir,f"obyektivka_{safe}_@DastyorAiBot.docx")
    doc=Document()
    # USER GRID: A4, Top 2.5cm, Bottom 2cm, Left 2.5cm, Right 2.5cm
    _margins(doc, top=2.5, bottom=2.0, left=2.5, right=2.5)
    ns=doc.styles["Normal"]
    ns.font.name=FONT;ns.font.size=Pt(11)
    ns.paragraph_format.space_before=Pt(0)
    ns.paragraph_format.space_after=Pt(0)
    pw=_pw(doc)
    # Content width: 21cm - 2.5cm - 2.5cm = 16cm
    # USER GRID: column gap target ~7–8cm ⇒ set right col start at 8cm
    left_col_w_cm = 8.0
    right_edge_cm = 16.0

    # 1. Sarlavha (preview: 20pt, margin-bottom 36px)
    p_t=doc.add_paragraph()
    p_t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    # USER GRID: Title top distance 3cm.
    # With top margin 2.5cm, add ~0.5cm (~14pt) before spacing.
    _spc(p_t,before=14,after=71,line=1.0)  # after=2.5cm ≈ 71pt to F.I.SH block
    _run(p_t,lb["title"],bold=True,size=18.0,color=BLACK,lspc=4.0)

    # 2. Ism | Foto (TABLELESS to avoid gridlines)
    # Left: name label + name, Right: photo (or "3x4") aligned to the right margin.
    p_hdr = doc.add_paragraph()
    _set_tab_right(p_hdr, right_edge_cm)
    _spc(p_hdr, before=0, after=17, line=1.0)  # USER GRID: 0.6cm ≈ 17pt between label and name
    _run(p_hdr, lb["name_label"], size=10.0, color=DARK, lspc=1.0)
    p_hdr.add_run("\t")
    has_photo = bool(photo_path and os.path.exists(photo_path))
    if has_photo:
        r = p_hdr.add_run()
        # USER GRID: 3.5cm x 4.5cm
        r.add_picture(photo_path, width=Cm(3.5), height=Cm(4.5))
    else:
        ph_run = p_hdr.add_run()
        _add_vml_dashed_box(ph_run, 3.5, 4.5, lb.get("photo", "3×4"))

    p_name = doc.add_paragraph()
    _set_tab_right(p_name, right_edge_cm)
    # USER GRID: Name 14pt bold, then 3cm gap before the info grid.
    # If current job exists, show it under the name (within this 3cm area).
    _spc(p_name, before=0, after=0, line=1.2)
    _run(p_name, (g("fullname","") or "").upper(), bold=True, size=14.0, color=BLACK, lspc=0.5)
    p_name.add_run("\t")
    p_name.add_run("")  # keep tab stop effective

    # Prefer explicit current_job from frontend (more reliable), else infer from work_experience
    role = (g("current_job", "") or "").strip()
    role_year = (g("current_job_year", "") or "").strip()
    if not role and not role_year:
        cj = _pick_current_job(g("work_experience", []) or [])
        if cj:
            role, role_year = (cj[0] or "").strip(), (cj[1] or "").strip()

    if role or role_year:
        role_line = role
        if role_year and role_year not in role_line:
            role_line = f"{role_line} ({role_year})" if role_line else role_year

        p_role = doc.add_paragraph()
        _set_tab_right(p_role, right_edge_cm)
        # Keep overall spacing similar, but slightly tighter than a full blank gap
        _spc(p_role, before=0, after=85, line=1.15)  # keep 3cm to the info grid
        _run(p_role, role_line, italic=True, size=11.0, color=DARK, lspc=0.2)
        p_role.add_run("\t")
        p_role.add_run("")
    else:
        # No role line → keep the 3cm gap here
        sp_gap = doc.add_paragraph()
        _spc(sp_gap, before=0, after=85, line=1.0)

    # 3. Info
    pairs=[
        (lb["bdate"],g("birthdate","") or g("bdate",""),lb["bplace"],g("birthplace","") or g("bplace","")),
        (lb["nation"],g("nation",""),lb["party"],g("party","")),
        (lb["edu"],g("education","") or g("edu",""),lb["grad"],g("graduated","") or g("grad","")),
        (lb["spec"],g("specialty","") or g("spec",""),lb["degree"],g("degree","")),
        (lb["stitle"],g("scientific_title","") or g("stitle",""),lb["langs"],g("languages","") or g("langs","")),
        (lb["military"],g("military_rank","") or g("military",""),lb["awards"],g("awards","")),
        (lb["deputy"],g("deputy",""),lb["phone"],g("phone","")),
    ]
    for lbl1,val1,lbl2,val2 in pairs:
        _info_pair_tabs(doc,left_col_w_cm,lbl1,val1,lbl2,val2)
    _info_single_tabs(doc,lb["addr"],g("address","") or g("addr",""))

    # 4. Mehnat faoliyati (preview section-header margin-top 40px)
    sp2=doc.add_paragraph();_spc(sp2,before=0,after=28)
    _sect_title(doc,lb["work_title"])
    _work_table(doc,pw,lb,g("work_experience",[]) or [])

    # 5. Qarindoshlar (preview margin-top 40px)
    sp3=doc.add_paragraph();_spc(sp3,before=0,after=28)
    fname=(g("fullname","") or "").split()[0].upper() if g("fullname","") else ""
    rel_h=lb["rel_title"] if lang=="ru" else (f"{fname} {lb['rel_title']}" if fname else lb["rel_title"])
    _sect_title(doc,rel_h)
    _rel_table(doc,pw,lb,g("relatives",[]) or [])

    doc.save(output_filepath)
    return output_filepath


def generate_cv_docx(data:dict, output_dir:str="temp") -> str:
    os.makedirs(output_dir,exist_ok=True)
    lang=data.get("lang","uz_lat")
    if lang not in _CV_SECTIONS:lang="uz_lat"
    sec=_CV_SECTIONS[lang]
    safe=(data.get("name","") or "unknown").replace(" ","_").replace("/","_")
    template=(data.get("template","minimal") or "minimal").lower()
    filepath=os.path.join(output_dir,f"cv_{safe}_{template}_@DastyorAiBot.docx")
    doc=Document()
    _margins(doc,top=1.5,bottom=1.5,left=2.0,right=1.5)
    pw=_pw(doc)

    def _cv_head(text,level=1,center=False):
        p=doc.add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        _spc(p,before=0,after=4 if level>0 else 6,line=1.15)
        sizes={0:20.0,1:12.5,2:11.0}
        _run(p,text.upper() if level<=1 else text,bold=True,font=FONT_CV,
             size=sizes.get(level,11.0),
             color=NAVY if level==1 else DARK_TEXT,
             lspc=2.5 if level==1 else 0)
        if level==1:_para_bdr_bot(p,NAVY_HEX,"6")

    def _cv_para(text,size=10.5,italic=False,bold=False,center=False,indent=0.0):
        p=doc.add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        _spc(p,before=0,after=3,line=1.2)
        if indent:p.paragraph_format.left_indent=Cm(indent)
        _run(p,text,italic=italic,bold=bold,size=size,font=FONT_CV,color=DARK_TEXT)

    nt=doc.add_table(rows=1,cols=1)
    _no_tbl_bdr(nt);_col_w(nt,pw,[100])
    nc=nt.cell(0,0)
    _shading(nc,NAVY_PALE)
    tcPr=nc._tc.get_or_add_tcPr()
    tb=OxmlElement("w:tcBorders")
    for s,col,sz in[("top",NAVY_HEX,"20"),("bottom",NAVY_HEX,"20"),("left",NAVY_HEX,"28")]:
        b=OxmlElement(f"w:{s}")
        b.set(qn("w:val"),"single");b.set(qn("w:sz"),sz)
        b.set(qn("w:space"),"0");b.set(qn("w:color"),col)
        tb.append(b)
    br=OxmlElement("w:right");br.set(qn("w:val"),"none");tb.append(br)
    tcPr.append(tb)
    _pad(nc,top=160,bottom=160,left=200,right=120)
    np_=nc.paragraphs[0]
    _spc(np_,before=0,after=4,line=1.1)
    _run(np_,(data.get("name","") or "").upper(),bold=True,size=18.0,font=FONT_CV,color=NAVY)
    rp=nc.add_paragraph()
    _spc(rp,before=2,after=0,line=1.0)
    _run(rp,data.get("spec","") or data.get("role",""),italic=True,size=11.0,font=FONT_CV,color=MUTED_CLR)

    sp=doc.add_paragraph();_spc(sp,before=0,after=6)
    parts=[]
    if data.get("phone"):parts.append(f"📞 {data['phone']}")
    if data.get("email"):parts.append(f"✉ {data['email']}")
    if data.get("loc"):parts.append(f"📍 {data['loc']}")
    if parts:
        cp=doc.add_paragraph()
        cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        _spc(cp,before=0,after=8,line=1.0)
        _run(cp,"  •  ".join(parts),size=9.5,font=FONT_CV,color=MUTED_CLR)

    if data.get("about"):
        _cv_head(sec["about"]);_cv_para(data["about"])
        sp2=doc.add_paragraph();_spc(sp2,before=0,after=6)

    works=data.get("works",[]) or data.get("work_experience",[]) or []
    if works:
        _cv_head(sec["exp"])
        for w in works:
            fy=w.get("f","") or w.get("from","") or w.get("year","")
            ty=w.get("t","") or w.get("to","")
            period=f"{fy} – {ty}" if ty else fy
            if not period and w.get("date"):period=w["date"]
            title_=w.get("title","") or w.get("pos","")
            co=w.get("company","") or w.get("co","")
            desc=w.get("d","") or w.get("position","") or w.get("desc","") or ""
            rp2=doc.add_paragraph();_spc(rp2,before=6,after=1,line=1.1)
            _run(rp2,title_ or period,bold=True,size=10.5,font=FONT_CV,color=NAVY)
            if co:_run(rp2,f"  —  {co}",size=10.0,font=FONT_CV,color=MUTED_CLR)
            if period and title_:
                sp_=doc.add_paragraph();_spc(sp_,before=0,after=1,line=1.0)
                _run(sp_,period,italic=True,size=9.0,font=FONT_CV,color=MUTED_CLR)
            if desc:_cv_para(desc,size=10.0,indent=0.3)
        sp3=doc.add_paragraph();_spc(sp3,before=0,after=6)

    edus=data.get("education_list",[]) or []
    if edus or data.get("edu") or data.get("grad"):
        _cv_head(sec["edu"])
        if edus:
            for edu in edus:
                ep=doc.add_paragraph();_spc(ep,before=5,after=1,line=1.1)
                _run(ep,edu.get("title","") or edu.get("name",""),bold=True,size=10.5,font=FONT_CV,color=NAVY)
                det=[x for x in[edu.get("date") or edu.get("year"),edu.get("company") or edu.get("field")]if x]
                if det:_run(ep,f"  —  {', '.join(det)}",size=9.5,font=FONT_CV,color=MUTED_CLR)
        else:
            if data.get("edu"):_cv_para(f"• {data['edu']}")
            if data.get("grad"):_cv_para(f"• {data['grad']}")
        sp4=doc.add_paragraph();_spc(sp4,before=0,after=6)

    skills_raw=data.get("skills","") or ""
    skills=[s.strip() for s in skills_raw.replace(",","\n").splitlines() if s.strip()]
    if skills:
        _cv_head(sec["skills"])
        skp=doc.add_paragraph();_spc(skp,before=3,after=3,line=1.3)
        _run(skp,"  •  ".join(skills),size=10.0,font=FONT_CV,color=DARK_TEXT)

    doc.save(filepath)
    return filepath


def convert_to_pdf_safe(docx_path:str, output_dir:str="temp") -> str|None:
    pdf_path=docx_path.replace(".docx",".pdf")
    os.makedirs(output_dir,exist_ok=True)
    try:
        from docx2pdf import convert
        convert(docx_path,pdf_path)
        if os.path.exists(pdf_path):return pdf_path
    except Exception as e:logger.debug(f"docx2pdf:{e}")
    try:
        import subprocess
        subprocess.run(["libreoffice","--headless","--convert-to","pdf",docx_path,"--outdir",output_dir],
                       check=True,timeout=60,capture_output=True)
        if os.path.exists(pdf_path):return pdf_path
    except Exception as e:logger.debug(f"LibreOffice:{e}")
    return None


if __name__=="__main__":
    sample={
        "lang":"uz_lat","fullname":"Karimov Jasur Abdullayevich",
        "birthdate":"15.03.1985","birthplace":"Toshkent shahri",
        "nation":"O'zbek","party":"Yo'q","education":"Oliy","graduated":"TDTU",
        "specialty":"Muhandis","degree":"PhD","scientific_title":"Dotsent",
        "languages":"Ingliz, Rus","military":"Leytenant",
        "awards":"Mehnat shuhrati","deputy":"Yo'q",
        "phone":"+998 90 123 45 67","address":"Toshkent sh., 22-uy",
        "work_experience":[
            {"year":"2005–2010","position":"TDTU — Muhandis"},
            {"year":"2010–2015","position":"Texnopark — Bosh muhandis"},
            {"year":"2015–hozir","position":"O'zbekiston FA — Lab. mudiri"},
        ],
        "relatives":[
            {"degree":"Xotini","fullname":"Karimova Dilnoza",
             "birth_year_place":"1987, Toshkent","work_place":"1-maktab","address":"Toshkent"},
            {"degree":"O'g'li","fullname":"Karimov Amir",
             "birth_year_place":"2010, Toshkent","work_place":"O'quvchi","address":"Toshkent"},
        ],
    }
    empty={"lang":"uz_lat","fullname":"Asad","birthplace":"Sad","nation":"Sda",
           "work_experience":[],"relatives":[]}
    out1=generate_obyektivka_docx(sample,"","/home/claude/test_new_full.docx")
    out2=generate_obyektivka_docx(empty,"","/home/claude/test_new_empty.docx")
    print(f"✓ To'liq: {out1}")
    print(f"✓ Bo'sh:  {out2}")