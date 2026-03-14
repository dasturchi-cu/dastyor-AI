"""
Official Obyektivka (Ma'lumotnoma) DOCX generator.
Reference layout: user's sample document style.
"""

import json
import logging
import os
import re
import base64
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

logger = logging.getLogger(__name__)


def _to_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()


def _parse_list(value: Any) -> list[dict[str, Any]]:
    if isinstance(value, list):
        return value
    if isinstance(value, str) and value.strip():
        try:
            parsed = json.loads(value)
            return parsed if isinstance(parsed, list) else []
        except Exception:
            return []
    return []


def _add_label_value(paragraph, label: str, value: str, *, value_on_new_line: bool = True) -> None:
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    if value_on_new_line:
        paragraph.add_run().add_break()
    paragraph.add_run(value or "yo'q")
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15


def _set_cell_no_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for side in ("top", "left", "bottom", "right"):
        border = tc_borders.find(qn(f"w:{side}"))
        if border is None:
            border = OxmlElement(f"w:{side}")
            tc_borders.append(border)
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")


def _set_table_no_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    # Remove style influence so table is layout-only.
    tbl_style = tbl_pr.find(qn("w:tblStyle"))
    if tbl_style is not None:
        tbl_pr.remove(tbl_style)

    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = tbl_borders.find(qn(f"w:{side}"))
        if border is None:
            border = OxmlElement(f"w:{side}")
            tbl_borders.append(border)
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")

    # Disable table look flags that may introduce visible styling.
    tbl_look = tbl_pr.find(qn("w:tblLook"))
    if tbl_look is None:
        tbl_look = OxmlElement("w:tblLook")
        tbl_pr.append(tbl_look)
    tbl_look.set(qn("w:firstRow"), "0")
    tbl_look.set(qn("w:lastRow"), "0")
    tbl_look.set(qn("w:firstColumn"), "0")
    tbl_look.set(qn("w:lastColumn"), "0")
    tbl_look.set(qn("w:noHBand"), "1")
    tbl_look.set(qn("w:noVBand"), "1")

    for row in table.rows:
        for cell in row.cells:
            _set_cell_no_borders(cell)


def _set_table_no_borders_strict(table) -> None:
    _set_table_no_borders(table)
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            for side in ("top", "left", "bottom", "right"):
                border = tc_borders.find(qn(f"w:{side}"))
                if border is None:
                    border = OxmlElement(f"w:{side}")
                    tc_borders.append(border)
                border.set(qn("w:val"), "none")
                border.set(qn("w:sz"), "0")
                border.set(qn("w:space"), "0")


def _set_cell_margins(cell, *, top_twips: int = 0, right_twips: int = 36, bottom_twips: int = 0, left_twips: int = 36) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (
        ("top", top_twips),
        ("right", right_twips),
        ("bottom", bottom_twips),
        ("left", left_twips),
    ):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_borders(table, size_pt: float = 1.0) -> None:
    size_eighth_pt = str(int(size_pt * 8))
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = tbl_borders.find(qn(f"w:{side}"))
        if border is None:
            border = OxmlElement(f"w:{side}")
            tbl_borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size_eighth_pt)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def _add_two_col_text_line(
    doc: Document,
    left_label: str,
    left_value: str,
    right_label: str | None = None,
    right_value: str | None = None,
    tab_pos: int | None = None,
) -> None:
    p = doc.add_paragraph()
    if tab_pos is not None:
        p.paragraph_format.tab_stops.add_tab_stop(tab_pos)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(6)

    left_run = p.add_run(f"{left_label}: ")
    left_run.bold = True
    p.add_run(left_value or "yo'q")

    if right_label:
        p.add_run("\t")
        right_run = p.add_run(f"{right_label}: ")
        right_run.bold = True
        p.add_run(right_value or "yo'q")


def generate_obyektivka_docx(
    user_data: dict[str, Any] | None = None,
    photo_path: str | None = None,
    output_filepath: str | None = None,
    **kwargs: Any,
) -> str:
    data = user_data or kwargs.get("data") or {}
    temp_photo_from_data = None
    if (not photo_path or not os.path.exists(photo_path)):
        photo_data = _to_text(data.get("photo_data"))
        try:
            if photo_data.startswith("data:image/") and "," in photo_data:
                header, b64 = photo_data.split(",", 1)
                mime = header.split(";")[0].split(":")[1].lower()
                ext = {
                    "image/png": "png",
                    "image/jpeg": "jpg",
                    "image/jpg": "jpg",
                    "image/webp": "webp",
                }.get(mime, "jpg")
                os.makedirs("temp", exist_ok=True)
                temp_photo_from_data = os.path.join("temp", f"oby_local_photo_{os.getpid()}.{ext}")
                with open(temp_photo_from_data, "wb") as f:
                    f.write(base64.b64decode(b64))
                photo_path = temp_photo_from_data
        except Exception as exc:
            logger.warning("Failed to decode photo_data in generator: %s", exc)
            photo_path = None

    if not isinstance(data, dict):
        raise ValueError("user_data must be a dictionary")

    if not output_filepath:
        os.makedirs("temp", exist_ok=True)
        safe_name = (_to_text(data.get("fullname")) or "Obyektivka").replace(" ", "_").replace("/", "_")
        output_filepath = os.path.join("temp", f"obyektivka_{safe_name}_@DastyorAiBot.docx")
    else:
        os.makedirs(os.path.dirname(output_filepath) or ".", exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.3
    style.paragraph_format.space_after = Pt(6)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(1.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("MA'LUMOTNOMA")
    # Oddiy holatda (sarlavha bold emas)
    title_run.font.size = Pt(16)
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.line_spacing = 1.3

    full_name = _to_text(data.get("fullname")) or "FAMILIYA ISM SHARIF"
    work_items = _parse_list(data.get("work_experience"))
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name.add_run(full_name)
    # Oddiy holatda
    name_run.font.size = Pt(14)
    name.paragraph_format.space_after = Pt(4)
    name.paragraph_format.line_spacing = 1.3

    current_job = _to_text(data.get("current_job"))
    current_job_year = _to_text(data.get("current_job_year"))
    if not current_job:
        # Fallback: if frontend confirm was skipped, infer current job
        # from work rows that end with h.v./hozirgacha variants.
        for idx, item in enumerate(work_items):
            year_raw = _to_text(item.get("year") or item.get("from"))
            year_norm = re.sub(r"[\s.\-_/]", "", year_raw.lower())
            is_current = any(
                key in year_norm
                for key in ("hv", "hvgacha", "hozirgacha", "ҳв", "ҳвгача", "ҳозиргача")
            )
            position_raw = _to_text(item.get("position") or item.get("description") or item.get("job"))
            if is_current and position_raw:
                current_job = position_raw
                if not current_job_year:
                    from_raw = _to_text(item.get("from"))
                    if from_raw:
                        current_job_year = from_raw
                    else:
                        match = re.search(r"(19|20)\d{2}", year_raw)
                        if match:
                            current_job_year = match.group(0)
                work_items.pop(idx)
                break
    if current_job:
        current_line = doc.add_paragraph()
        current_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if current_job_year:
            year_text = current_job_year.rstrip(".")
            current_line.add_run(f"{year_text} yil:")
            current_line.add_run().add_break()
            current_line.add_run(current_job)
        else:
            current_line.add_run(current_job)
        current_line.paragraph_format.line_spacing = 1.15
        current_line.paragraph_format.space_after = Pt(14)
    else:
        name.paragraph_format.space_after = Pt(20)

    # Main info block: invisible 3-column table for precise alignment.
    section = doc.sections[0]
    total_width = section.page_width - section.left_margin - section.right_margin
    photo_col_width = Cm(4)
    info_col_width = int((total_width - photo_col_width) / 2)

    info_tbl = doc.add_table(rows=8, cols=3)
    info_tbl.autofit = False
    info_tbl.columns[0].width = info_col_width
    info_tbl.columns[1].width = info_col_width
    info_tbl.columns[2].width = photo_col_width
    _set_table_no_borders_strict(info_tbl)
    info_rows = [
        (("Tug'ilgan yili", _to_text(data.get("birthdate"))), ("Tug'ilgan joyi", _to_text(data.get("birthplace"))), False),
        (("Millati", _to_text(data.get("nation"))), ("Partiyaviyligi", _to_text(data.get("party"))), False),
        (("Ma'lumoti", _to_text(data.get("education"))), ("Tamomlagan", _to_text(data.get("graduated"))), False),
        (("Ma'lumoti bo'yicha mutaxassisligi", _to_text(data.get("specialty"))), None, False),
        (("Ilmiy darajasi", _to_text(data.get("degree"))), ("Ilmiy unvoni", _to_text(data.get("scientific_title"))), False),
        (("Qaysi chet tillarini biladi", _to_text(data.get("languages"))), ("Harbiy unvoni", _to_text(data.get("military_rank"))), False),
        (("Davlat mukofotlari bilan taqdirlanganmi (qanaqa)", _to_text(data.get("awards"))), None, True),
        (("Xalq deputatlari respublika, viloyat, shahar va tuman Kengashi deputatimi yoki boshqa saylanadigan organlarning a'zosimi (to'liq ko'rsatish lozim)", _to_text(data.get("deputy"))), None, True),
    ]

    for row_idx, (left_data, right_data, merge_row) in enumerate(info_rows):
        left_cell = info_tbl.cell(row_idx, 0)
        right_cell = info_tbl.cell(row_idx, 1)
        left_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _set_cell_margins(left_cell)
        _set_cell_margins(right_cell)

        _add_label_value(left_cell.paragraphs[0], left_data[0], left_data[1], value_on_new_line=True)

        if right_data is not None:
            _add_label_value(right_cell.paragraphs[0], right_data[0], right_data[1], value_on_new_line=True)
        elif merge_row:
            merged = left_cell.merge(right_cell)
            merged.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            _set_cell_margins(merged)
            _set_cell_no_borders(merged)
        else:
            right_cell.paragraphs[0].paragraph_format.space_after = Pt(6)
            right_cell.paragraphs[0].paragraph_format.line_spacing = 1.3

    photo_cell = info_tbl.cell(0, 2).merge(info_tbl.cell(3, 2))
    photo_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _set_cell_margins(photo_cell, right_twips=0, left_twips=0)
    _set_cell_no_borders(photo_cell)
    photo_p = photo_cell.paragraphs[0]
    photo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    photo_p.paragraph_format.space_before = Pt(0)
    photo_p.paragraph_format.space_after = Pt(0)
    if photo_path and os.path.exists(photo_path):
        try:
            run = photo_p.add_run()
            run.add_picture(photo_path, width=Cm(3), height=Cm(4))
        except Exception as exc:
            logger.warning("Failed to insert photo: %s", exc)
            ph = photo_p.add_run("[3x4 foto]")
            ph.italic = True
    else:
        ph = photo_p.add_run("[3x4 foto]")
        ph.italic = True

    work_title = doc.add_paragraph()
    work_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    wr = work_title.add_run("MEHNAT FAOLIYATI")
    # Oddiy holatda
    wr.font.size = Pt(14)
    work_title.paragraph_format.space_before = Pt(20)
    work_title.paragraph_format.space_after = Pt(6)
    work_title.paragraph_format.line_spacing = 1.3

    if work_items:
        for item in work_items:
            year = _to_text(item.get("year")) or _to_text(item.get("from"))
            end = _to_text(item.get("to"))
            if year and end and "yy" not in year.lower():
                year = f"{year}-{end} yy."
            year = year.rstrip(".") if year else year  # 2009. -> 2009, keyin : va pastga matn
            position = _to_text(item.get("position") or item.get("description") or item.get("job"))
            if not (year or position):
                continue
            p = doc.add_paragraph()
            if year and position:
                p.add_run(f"{year}:")
                p.add_run().add_break()
                p.add_run(position)
            else:
                p.add_run(year or position)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.3
    else:
        doc.add_paragraph("-")

    # Page 2 relatives (table visible as in sample).
    relatives = _parse_list(data.get("relatives"))
    doc.add_page_break()
    t1 = doc.add_paragraph()
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = t1.add_run(f"{full_name}ning yaqin qarindoshlari haqida")
    r1.bold = True
    r1.font.size = Pt(12)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("MA'LUMOTNOMA")
    r2.bold = True
    r2.font.size = Pt(12)

    rel_tbl = doc.add_table(rows=1, cols=5)
    rel_tbl.autofit = False
    rel_tbl.columns[0].width = int(total_width * 0.18)
    rel_tbl.columns[1].width = int(total_width * 0.26)
    rel_tbl.columns[2].width = int(total_width * 0.19)
    rel_tbl.columns[3].width = int(total_width * 0.19)
    rel_tbl.columns[4].width = int(total_width * 0.18)
    _set_table_borders(rel_tbl, size_pt=1.0)

    headers = [
        "Qarindoshligi",
        "Familiyasi ismi va\notasining ismi",
        "Tug'ilgan yili\nva joyi",
        "Ish joyi va\nlavozimi",
        "Turar joyi",
    ]
    for i, h in enumerate(headers):
        p = rel_tbl.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(11)

    if relatives:
        for rel in relatives:
            row = rel_tbl.add_row().cells
            vals = [
                _to_text(rel.get("degree") or rel.get("type")),
                _to_text(rel.get("fullname") or rel.get("name")),
                _to_text(rel.get("birth_year_place") or rel.get("birth")),
                _to_text(rel.get("work_place") or rel.get("job")),
                _to_text(rel.get("address") or rel.get("addr")),
            ]
            for i, val in enumerate(vals):
                p = row[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rr = p.add_run(val or "yo'q")
                rr.font.size = Pt(11)
                if i == 0:
                    rr.bold = True
    else:
        row = rel_tbl.add_row().cells
        merged = row[0].merge(row[4])
        p = merged.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("Yaqin qarindoshlar haqida ma'lumot kiritilmagan.")

    doc.save(output_filepath)
    if temp_photo_from_data and os.path.exists(temp_photo_from_data):
        try:
            os.remove(temp_photo_from_data)
        except Exception:
            pass
    return output_filepath
