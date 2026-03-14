import os
import logging
from fastapi import FastAPI, Request, Response, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from telegram import Update
from contextlib import asynccontextmanager
from fastapi.staticfiles import StaticFiles
from main import setup_application

# Setup Logging
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# Initialize Application
application = setup_application()

@asynccontextmanager
async def lifespan(app: FastAPI):
    # Initialize application
    await application.initialize()
    await application.start()
    
    # Set webhook automatically on startup (configurable via .env)
    webhook_url = os.getenv("WEBHOOK_URL", "https://dastyor-ai.onrender.com/webhook").strip()
    logger.info(f"Setting webhook to: {webhook_url}")
    await application.bot.set_webhook(url=webhook_url, drop_pending_updates=True)
    
    logger.info("🚀 Webhook Application Started")
    yield
    # Shutdown
    await application.stop()
    await application.shutdown()
    logger.info("🛑 Webhook Application Stopped")

from fastapi.responses import RedirectResponse

app = FastAPI(lifespan=lifespan)

# ── CORS: allow webapp pages to call /api/* from any origin ──
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Serve Web App Files
app.mount("/webapp", StaticFiles(directory="webapp"), name="webapp")

@app.get("/")
async def root():
    return RedirectResponse(url="/webapp/index.html")

from pydantic import BaseModel
from fastapi import File, UploadFile, Form, Query
from fastapi.responses import StreamingResponse, HTMLResponse
from typing import List, Optional
from telegram import InputFile
import asyncio, time, io
import base64

from bot.services.render_service import (
    generate_cv_pdf, safe_filename,
)
from bot.utils.delivery import send_docx_with_confirmation

SITE_BASE_URL = os.getenv("SITE_BASE_URL", "https://dastyor-ai.onrender.com")


def _safe_name(name: str, fallback: str) -> str:
    base = os.path.basename(name or "").strip()
    if not base:
        return fallback
    return "".join(ch for ch in base if ch.isalnum() or ch in ("-", "_", ".", " ")).strip() or fallback


# ═══════════════════════════════════════════════════════════════════════════
# HELPER: Resolve telegram_id from ANY source (token > URL param > form)
# ═══════════════════════════════════════════════════════════════════════════
def _resolve_uid(
    telegram_id_param: Optional[str] = None,
    session_token: Optional[str] = None,
) -> Optional[str]:
    """
    Priority order:
      1. Session token (most secure — validated against our DB)
      2. Raw telegram_id param (from URL / form, legacy / WebApp fallback)
    Returns the telegram_id string or None.
    """
    if session_token:
        from bot.services.session_service import resolve_telegram_id
        uid = resolve_telegram_id(session_token)
        if uid:
            return uid
    if telegram_id_param and telegram_id_param.strip().isdigit():
        return telegram_id_param.strip()
    return None


# ═══════════════════════════════════════════════════════════════════════════
# /api/auth — Create session from Telegram identity
# Called by index.html after Telegram WebApp is ready.
# ═══════════════════════════════════════════════════════════════════════════
class AuthRequest(BaseModel):
    telegram_id: int
    first_name:  str = ""
    username:    str = ""
    photo_url:   str = ""

@app.post("/api/auth")
async def api_auth(req: AuthRequest):
    """
    Exchange Telegram identity → server session token.
    The token is stored in sessionStorage and sent back
    with every subsequent API request.
    """
    try:
        from bot.services.session_service import create_session
        token = create_session(
            telegram_id = req.telegram_id,
            first_name  = req.first_name,
            username    = req.username,
            photo_url   = req.photo_url,
        )
        # Also upsert the user profile in the CRM
        from bot.services.user_service import track_user_activity, save_chat_id
        class _FakeUser:
            id = req.telegram_id
            first_name = req.first_name
            username   = req.username
        track_user_activity(_FakeUser(), command="web_auth")
        # Persist chat_id so file-delivery endpoints can find the correct chat
        save_chat_id(req.telegram_id, req.telegram_id)  # For private chats uid == chat_id

        return {"ok": True, "token": token, "telegram_id": req.telegram_id}
    except Exception as e:
        logger.error(f"/api/auth error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e)[:200])


# ═══════════════════════════════════════════════════════════════════════════
# /api/me — Return user profile for given token or telegram_id
# ═══════════════════════════════════════════════════════════════════════════
@app.get("/api/me")
async def api_me(
    token:       Optional[str] = Query(None),
    telegram_id: Optional[str] = Query(None),
):
    uid = _resolve_uid(telegram_id, token)
    if not uid:
        raise HTTPException(status_code=401, detail="Foydalanuvchi aniqlanmadi")

    from bot.services.user_service       import get_user_profile
    from bot.services.settings_service   import is_premium
    from bot.services.session_service    import get_session_by_telegram_id

    profile = get_user_profile(uid) or {}
    session = get_session_by_telegram_id(uid) or {}

    return {
        "ok"          : True,
        "telegram_id" : uid,
        "first_name"  : session.get("first_name", profile.get("first_name", "")),
        "username"    : session.get("username",   profile.get("username", "")),
        "photo_url"   : session.get("photo_url",  ""),
        "is_premium"  : is_premium(int(uid)),
        "files_processed": profile.get("files_processed", 0),
        "joined_at"   : profile.get("joined_at", ""),
        "last_active" : profile.get("last_active", ""),
    }


# ═══════════════════════════════════════════════════════════════════════════
# /api/translit — Server-side Cyrillic ↔ Latin (consistent with bot)
# ═══════════════════════════════════════════════════════════════════════════
class TranslitRequest(BaseModel):
    text:      str
    direction: str   # "krill_to_lotin" | "lotin_to_krill"

@app.post("/api/translit")
async def api_translit(req: TranslitRequest):
    if not req.text or not req.text.strip():
        raise HTTPException(status_code=400, detail="Matn bo'sh bo'lishi mumkin emas")
    if len(req.text) > 50_000:
        raise HTTPException(status_code=400, detail="Matn 50 000 belgidan oshmasligi kerak")

    valid = {"krill_to_lotin", "lotin_to_krill"}
    if req.direction not in valid:
        raise HTTPException(status_code=400, detail=f"Noto'g'ri yo'nalish: {req.direction}")

    try:
        from bot.services.transliterate_service import transliterate
        result = transliterate(req.text, req.direction)  # type: ignore[arg-type]
        return {"ok": True, "result": result, "direction": req.direction}
    except Exception as e:
        logger.error(f"/api/translit error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e)[:200])


# ═══════════════════════════════════════════════════════════════════════════
# /api/bot-link — Generate deep-link for a bot command
# Lets the website open the bot at exactly the right state.
# ═══════════════════════════════════════════════════════════════════════════
BOT_USERNAME = os.getenv("BOT_USERNAME", "DastyorAiBot")
WEBAPP_BASE  = os.getenv("WEBAPP_BASE",  "https://dastyor-ai.onrender.com/webapp")
WEBAPP_VERSION = os.getenv("WEBAPP_VERSION", "20260311")

@app.get("/api/bot-link")
async def api_bot_link(
    action:      str = Query(..., description="cv | obyektivka | ocr | pdf | translit | translate"),
    telegram_id: Optional[str] = Query(None),
):
    """
    Returns a t.me deep-link that opens the bot AND, for web-first
    actions, also returns the direct webapp URL so the website can
    redirect without going through Telegram.
    """
    page_map = {
        "cv"         : "cv.html",
        "obyektivka" : "obyektivka.html",
        "ocr"        : "ocr.html",
        "pdf"        : "img2pdf.html",
        "translit"   : "translit.html",
        "translate"  : "translate.html",
        "premium"    : "premium.html",
    }
    page = page_map.get(action)
    if not page:
        raise HTTPException(status_code=400, detail=f"Noma'lum action: {action}")

    # Deep-link: t.me/BotUsername?start=action
    bot_link  = f"https://t.me/{BOT_USERNAME}?start={action}"
    # Direct webapp URL (already has telegram_id → user stays logged in)
    tid_param = f"?telegram_id={telegram_id}" if telegram_id else ""
    # cache-bust for Telegram mobile WebView
    v_param = f"{'&' if tid_param else '?'}v={WEBAPP_VERSION}"
    webapp_url = f"{WEBAPP_BASE}/{page}{tid_param}{v_param}"

    return {"ok": True, "bot_link": bot_link, "webapp_url": webapp_url, "action": action}


# ═══════════════════════════════════════════════════════════════════════════
# /api/stats — Per-user usage statistics (for the website dashboard)
# ═══════════════════════════════════════════════════════════════════════════
@app.get("/api/stats")
async def api_stats(
    token:       Optional[str] = Query(None),
    telegram_id: Optional[str] = Query(None),
):
    uid = _resolve_uid(telegram_id, token)
    if not uid:
        raise HTTPException(status_code=401, detail="Foydalanuvchi aniqlanmadi")

    from bot.services.user_service     import get_user_profile
    from bot.services.settings_service import is_premium

    profile = get_user_profile(uid) or {}
    premium = is_premium(int(uid))

    return {
        "ok"              : True,
        "telegram_id"     : uid,
        "is_premium"      : premium,
        "files_processed" : profile.get("files_processed", 0),
        "sessions"        : profile.get("sessions", 1),
        "last_service"    : profile.get("last_service", ""),
        "last_active"     : profile.get("last_active", ""),
        "joined_at"       : profile.get("joined_at", ""),
    }


# ═══════════════════════════════════════════════════════════════════════════
# /api/notify — Push a text notification to user's Telegram chat
# Lets any website action trigger a Telegram message.
# ═══════════════════════════════════════════════════════════════════════════
class NotifyRequest(BaseModel):
    telegram_id : int
    message     : str
    token       : Optional[str] = None

class SupportRequest(BaseModel):
    telegram_id : int
    username    : Optional[str] = ""
    message     : str
    token       : Optional[str] = None

@app.post("/api/notify")
async def api_notify(req: NotifyRequest):
    """Send a plain-text notification to the user's Telegram chat."""
    # Validate token OR just trust telegram_id (within same server — no public exposure risk)
    uid = _resolve_uid(str(req.telegram_id), req.token)
    if not uid:
        raise HTTPException(status_code=401, detail="Foydalanuvchi aniqlanmadi")

    if not req.message or not req.message.strip():
        raise HTTPException(status_code=400, detail="Xabar bo'sh")
    if len(req.message) > 4000:
        raise HTTPException(status_code=400, detail="Xabar 4000 belgi oshmasligi kerak")

    try:
        await application.bot.send_message(
            chat_id    = int(uid),
            text       = req.message,
            parse_mode = "HTML",
        )
        return {"ok": True}
    except Exception as e:
        logger.warning(f"/api/notify failed for {uid}: {e}")
        raise HTTPException(status_code=502, detail=f"Telegram xatosi: {str(e)[:200]}")

@app.post("/api/support")
async def api_support(req: SupportRequest):
    """
    Forward support/contact message from webapp to admin chat(s).
    """
    uid = _resolve_uid(str(req.telegram_id), req.token)
    if not uid:
        raise HTTPException(status_code=401, detail="Foydalanuvchi aniqlanmadi")

    msg = (req.message or "").strip()
    if not msg:
        raise HTTPException(status_code=400, detail="Xabar bo'sh")
    if len(msg) > 4000:
        raise HTTPException(status_code=400, detail="Xabar 4000 belgidan oshmasligi kerak")

    admin_ids = []
    raw_admin = (os.getenv("ADMIN_USER_ID") or "").strip()
    if raw_admin:
        admin_ids.extend(
            int(x.strip()) for x in raw_admin.split(",")
            if x.strip().lstrip("-").isdigit()
        )
    support_group = os.getenv("SUPPORT_GROUP_ID", "").strip()
    if support_group and support_group.lstrip("-").isdigit():
        admin_ids.append(int(support_group))

    # Safe fallback to existing feedback group if no env values were configured
    if not admin_ids:
        admin_ids.append(-1003457224552)

    username = (req.username or "").strip()
    username_text = f"@{username}" if username else "yo'q"
    text = (
        "📩 <b>WebApp support so'rovi</b>\n\n"
        f"🆔 User ID: <code>{uid}</code>\n"
        f"👤 Username: {username_text}\n\n"
        f"💬 Xabar:\n{msg}"
    )

    sent = 0
    for chat_id in dict.fromkeys(admin_ids):
        try:
            await application.bot.send_message(
                chat_id=chat_id,
                text=text,
                parse_mode="HTML",
            )
            sent += 1
        except Exception as e:
            logger.warning(f"/api/support forward failed to {chat_id}: {e}")

    if sent == 0:
        raise HTTPException(status_code=502, detail="Support xabarini yuborib bo'lmadi")

    return {"ok": True, "forwarded_to": sent}


# ═══════════════════════════════════════════════════════════════════════════
# /api/translate (existing, kept in place)
# ═══════════════════════════════════════════════════════════════════════════
class TranslateRequest(BaseModel):
    text: str
    direction: str

@app.post("/api/translate")
async def api_translate(req: TranslateRequest):
    if not req.text or not req.text.strip():
        raise HTTPException(status_code=400, detail="Matn bo'sh bo'lishi mumkin emas")
    if len(req.text) > 5000:
        raise HTTPException(status_code=400, detail="Matn 5000 belgidan oshmasligi kerak")

    valid_dirs = {"uz_en", "en_uz", "ru_uz", "uz_ru", "ru_en"}
    if req.direction not in valid_dirs:
        raise HTTPException(status_code=400, detail=f"Noto'g'ri yo'nalish: {req.direction}")

    try:
        from bot.services.ai_service import translate_text
        result = await translate_text(req.text, req.direction)
        if not result or result.startswith("Tarjimada xato") or result.startswith("AI model"):
            raise HTTPException(status_code=502, detail=result or "Tarjima bo'sh qaytdi")
        return {"ok": True, "translated_text": result, "direction": req.direction}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Translate API error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Tarjima serveri xatosi: {str(e)[:200]}")


# ─────────────────────────────────────────────────────────────────────
# /api/ocr_direct — Website-first OCR endpoint
#
# Architecture:
#   1. Image uploaded from browser (no telegram_id required)
#   2. Gemini OCR extracts text as HTML
#   3. DOCX created in-memory via python-docx
#   4. DOCX bytes streamed back to browser → auto-download
#   5. If telegram_id provided → background task sends copy to Telegram
#
# Why old /api/upload_ocr was broken:
#   - Required telegram_id (Form required field) → error without Telegram
#   - Imported non-existent process_ocr_logic from smart_logic → ImportError
#   - int('') raises ValueError at line 1
#   - DOCX was deleted after bot send — never returned to browser
# ─────────────────────────────────────────────────────────────────────
@app.post("/api/ocr_direct")
async def api_ocr_direct(
    file: UploadFile = File(...),
    telegram_id: Optional[str] = Form(None),   # optional — works without Telegram
):
    ts = int(time.time())
    safe_upload_name = _safe_name(file.filename or "", "img.jpg")
    img_path  = f"temp/ocr_{ts}_{safe_upload_name}"
    os.makedirs("temp", exist_ok=True)

    # ── 1. Save uploaded image to disk ──────────────────────────────
    try:
        raw = await file.read()
        if len(raw) == 0:
            raise HTTPException(status_code=400, detail="Fayl bo'sh")
        if len(raw) > 15 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="Fayl 15 MB dan oshmasligi kerak")
        with open(img_path, "wb") as f:
            f.write(raw)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Fayl saqlashda xato: {e}")

    # ── 2. OCR: extract text from image using Gemini ─────────────────
    try:
        from bot.services.ocr_service import extract_text_from_image
        html_text = await extract_text_from_image(img_path)
    except Exception as e:
        logger.error(f"OCR extract error: {e}", exc_info=True)
        _cleanup(img_path)
        raise HTTPException(status_code=502, detail=f"OCR xatosi: {str(e)[:200]}")

    if not html_text or not html_text.strip():
        _cleanup(img_path)
        raise HTTPException(status_code=422, detail="Rasmdan matn ajratib bo'lmadi. Aniqroq rasm yuboring.")

    # ── 3. Build DOCX in a thread (CPU-bound) ────────────────────────
    try:
        from docx import Document
        from bot.handlers.ocr_to_word import add_html_to_docx

        def build_docx() -> bytes:
            doc = Document()
            doc.add_heading("OCR Natijasi", 0)
            try:
                add_html_to_docx(doc, html_text)
            except Exception as parse_err:
                logger.warning(f"HTML parse fallback: {parse_err}")
                doc.add_paragraph(html_text)  # raw text fallback
            # Save to bytes buffer (no disk needed)
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf.read()

        loop = asyncio.get_running_loop()
        docx_bytes = await loop.run_in_executor(None, build_docx)
    except Exception as e:
        logger.error(f"DOCX build error: {e}", exc_info=True)
        _cleanup(img_path)
        raise HTTPException(status_code=500, detail=f"Word hujjat yaratishda xato: {str(e)[:200]}")

    # ── 4. Optional Telegram send (background, non-blocking) ─────────
    if telegram_id and telegram_id.strip().isdigit():
        chat_id = int(telegram_id)

        async def send_to_telegram():
            try:
                buf = io.BytesIO(docx_bytes)
                buf.name = f"OCR_Natija_{ts}.docx"
                ok = await send_docx_with_confirmation(
                    application.bot,
                    chat_id,
                    buf,
                    filename=buf.name,
                    caption="✅ Rasm Word ga aylantirildi!\n📎 Fayl tayyor.",
                )
                if ok:
                    logger.info(f"OCR DOCX sent to Telegram chat {chat_id}")
            except Exception as tg_err:
                logger.warning(f"Telegram send failed (non-fatal): {tg_err}")

        asyncio.create_task(send_to_telegram())

    # ── 5. Cleanup temp image ─────────────────────────────────────────
    _cleanup(img_path)

    # ── 6. Stream DOCX bytes to browser → triggers auto-download ─────
    filename = f"DASTYOR_OCR_{ts}_@DastyorAiBot.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _cleanup(*paths):
    """Safely remove temp files."""
    for p in paths:
        try:
            if p and os.path.exists(p):
                os.remove(p)
        except Exception:
            pass


# ── Legacy endpoint kept for backward compatibility ───────────────────
@app.post("/api/upload_ocr")
async def api_upload_ocr(telegram_id: str = Form(...), files: List[UploadFile] = File(...)):
    """
    Legacy endpoint — kept for backward compat.
    New website code should use /api/ocr_direct instead.
    """
    try:
        chat_id = int(telegram_id)
        legacy_ts = int(time.time())
        os.makedirs("temp", exist_ok=True)
        img_paths = []
        for f in files:
            path = f"temp/legacy_ocr_{legacy_ts}_{f.filename}"
            content = await f.read()
            with open(path, "wb") as fh:
                fh.write(content)
            img_paths.append(path)

        async def run_legacy_ocr():
            from bot.services.ocr_service import extract_text_from_image
            from bot.handlers.ocr_to_word import add_html_to_docx
            from docx import Document
            try:
                msg = await application.bot.send_message(
                    chat_id=chat_id,
                    text="⏳ Rasm qabul qilindi. OCR qilinmoqda..."
                )
                for img_path in img_paths:
                    html_text = await extract_text_from_image(img_path)
                    if not html_text:
                        continue
                    ts2 = int(time.time())
                    docx_path = f"temp/legacy_{chat_id}_{ts2}.docx"
                    doc = Document()
                    doc.add_heading("OCR Natijasi", 0)
                    add_html_to_docx(doc, html_text)
                    doc.save(docx_path)
                    with open(docx_path, "rb") as df:
                        await send_docx_with_confirmation(
                            application.bot,
                            chat_id,
                            df,
                            filename=f"OCR_{ts2}.docx",
                            caption="✅ Word fayl tayyor!",
                        )
                    _cleanup(docx_path)
                await application.bot.delete_message(chat_id=chat_id, message_id=msg.message_id)
            except Exception as ex:
                logger.error(f"Legacy OCR error: {ex}", exc_info=True)
            finally:
                for p in img_paths:
                    _cleanup(p)

        asyncio.create_task(run_legacy_ocr())
        return {"status": "ok"}
    except Exception as e:
        logger.error(f"Upload OCR API error: {e}", exc_info=True)
        return {"error": str(e)}



@app.post("/api/pdf_direct")
async def api_pdf_direct(
    files: List[UploadFile] = File(...),
    telegram_id: Optional[str] = Form(None)
):
    try:
        import time
        logger.info(f"[PDF API] Boshlandi. telegram_id={telegram_id}, files_count={len(files)}")
        os.makedirs("temp", exist_ok=True)
        
        # Save uploaded files with unique names
        img_paths = []
        ts = int(time.time())
        try:
            for i, file in enumerate(files):
                if not file.filename:
                    raise HTTPException(status_code=400, detail="Noto'g'ri fayl nomi")
                ext = os.path.splitext(file.filename)[1] or ".jpg"
                path = f"temp/pdf_req_{ts}_{i}{ext}"
                content = await file.read()
                if not content:
                    raise HTTPException(status_code=400, detail=f"Bo'sh fayl: {file.filename}")
                if len(content) > 15 * 1024 * 1024:
                    raise HTTPException(status_code=400, detail=f"Fayl juda katta: {file.filename}")
                with open(path, "wb") as f:
                    f.write(content)
                img_paths.append(path)
            logger.info(f"[PDF API] Barcha {len(img_paths)} rasm serverga yuklandi va vaqtinchalik papkaga saqlandi: {img_paths}")
        except Exception as e:
            logger.error(f"[PDF API] Rasm yuklashda yoki saqlashda xatolik: {e}", exc_info=True)
            raise HTTPException(status_code=500, detail=f"Rasm yuklash xatosi: {e}")
        
        if not img_paths:
            logger.warning("[PDF API] Hech qanday rasm topilmadi. img_paths bo'sh.")
            raise HTTPException(status_code=400, detail="Fayl yuklanmadi")
        
        pdf_path = f"temp/merged_{ts}.pdf"
        
        # CPU-bound PDF creation
        try:
            from bot.services.pdf_service import images_to_pdf
            logger.info(f"[PDF API] PDF generatsiya funksiyasi ishga tushmoqda... output: {pdf_path}")
            loop = asyncio.get_running_loop()
            await loop.run_in_executor(
                None, images_to_pdf, img_paths, pdf_path
            )
            logger.info(f"[PDF API] PDF yaratish yakunlandi: {pdf_path}")
        except Exception as build_err:
            logger.error(f"[PDF API] PDF yaralishda (images_to_pdf) xato: {build_err}", exc_info=True)
            _cleanup(*img_paths)
            raise HTTPException(status_code=500, detail=f"PDF yaratishda xato: {build_err}")
            
        if not os.path.exists(pdf_path):
            logger.error(f"[PDF API] PDF fayl yaratilmadi, file yo'q: {pdf_path}")
            _cleanup(*img_paths)
            raise HTTPException(status_code=500, detail="PDF fayl yaratilmadi")
            
        # Read the generated PDF into memory so we can safely delete the file
        with open(pdf_path, 'rb') as f:
            pdf_bytes = f.read()
            
        # Cleanup ALL temp files now that we have bytes
        _cleanup(pdf_path, *img_paths)
        logger.info(f"[PDF API] Temp fayllar tozalandi. PDF xotiraga yuklandi, size: {len(pdf_bytes)} bytes")
        
        # Optional: Send to Telegram in background
        if telegram_id and telegram_id.strip().isdigit():
            chat_id = int(telegram_id)
            async def send_pdf_to_telegram():
                try:
                    logger.info(f"[PDF API Background] Telegram ga jo'natish (chat_id={chat_id}) boshlandi.")
                    buf = io.BytesIO(pdf_bytes)
                    buf.name = f"DASTYOR_AI_Rasmlar_{ts}_@DastyorAiBot.pdf"
                    await application.bot.send_document(
                        chat_id=chat_id,
                        document=InputFile(buf, filename=buf.name),
                        caption=f"✅ **PDF tayyor!**\n📄 {len(img_paths)} ta rasm birlashtirildi.",
                        parse_mode="Markdown"
                    )
                    logger.info(f"[PDF API Background] PDF Telegram {chat_id} -ga muvaffaqiyatli jo'natildi.")
                except Exception as tg_err:
                    logger.error(f"[PDF API Background] Telegram'ga jo'natish xatosi: {tg_err}", exc_info=True)
                    
            asyncio.create_task(send_pdf_to_telegram())
        else:
            logger.info("[PDF API] telegram_id yo'q yoki invalid, background jo'natish bajarilmaydi.")
            
        # Stream back to browser
        filename = f"DASTYOR_AI_Rasmlar_{ts}_@DastyorAiBot.pdf"
        logger.info(f"[PDF API] Frontend ga {filename} nomli streaming javob qaytarilmoqda.")
        
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
        
    except HTTPException:
        logger.warning("[PDF API] HTTPException ko'tarildi")
        raise
    except Exception as e:
        logger.error(f"[PDF API] Umumiy / Kutilmagan Xatolik: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/health")
async def health():
    return {"status": "healthy"}


# ═══════════════════════════════════════════════════════════════════════════
# /api/upload_to_telegram — Direct file sender to Telegram from Frontend JS
# ═══════════════════════════════════════════════════════════════════════════
@app.post("/api/upload_to_telegram")
async def api_upload_to_telegram(
    file: UploadFile = File(...),
    telegram_id: str = Form(None),
    token: str = Form(None),
    caption: str = Form("")
):
    uid = _resolve_uid(telegram_id, token)
    if not uid:
        raise HTTPException(status_code=401, detail="Unauthorized")
        
    try:
        content = await file.read()
        logger.info(f"Sending frontend-generated file {file.filename} to UID {uid}")
        
        from telegram import InputFile
        import io
        buf = io.BytesIO(content)
        buf.name = file.filename
        
        if (file.filename or "").lower().endswith(".docx"):
            ok = await send_docx_with_confirmation(
                application.bot,
                int(uid),
                buf,
                filename=file.filename,
                caption=caption or "✅ Faylingiz tayyorlandi!",
            )
            if not ok:
                return {"ok": False, "error": "Word fayl yuborilmadi"}
        else:
            await application.bot.send_document(
                chat_id=int(uid),
                document=InputFile(buf, filename=file.filename),
                caption=caption or "✅ Faylingiz tayyorlandi!"
            )
        return {"ok": True}
    except Exception as e:
        logger.error(f"Error sending file via /api/upload_to_telegram: {e}", exc_info=True)
        return {"ok": False, "error": str(e)}


# ═══════════════════════════════════════════════════════════════════════════
# /api/generate_cv — CV DOCX generator (Website → Server → Browser + Telegram)
# ─────────────────────────────────────────────────────────────────────────────
# Architecture:
#   1. Website POSTs JSON payload (form fields + template choice)
#   2. Server calls bot.services.doc_generator.generate_cv_docx()
#   3. DOCX bytes streamed to browser → auto-download
#   4. If telegram_id provided → background sends same file to Telegram chat
# ═══════════════════════════════════════════════════════════════════════════
class CVRequest(BaseModel):
    telegram_id : Optional[int] = None
    token       : Optional[str] = None
    # CV form fields
    name        : str = ""
    spec        : str = ""
    phone       : str = ""
    email       : str = ""
    loc         : str = ""
    addr        : str = ""
    birth       : str = ""
    place       : str = ""
    nation      : str = ""
    langs       : str = ""
    about       : str = ""
    template    : str = "minimal"
    works       : list = []
    education_list: list = []
    skills      : str = ""

@app.post("/api/generate_cv")
async def api_generate_cv(req: CVRequest):
    """
    Generate a CV DOCX on the server from website form data.
    • Returns the file as a streaming download for the browser.
    • Simultaneously sends the file to the user’s Telegram chat (background).
    """
    ts = int(time.time())
    os.makedirs("temp", exist_ok=True)

    # Resolve Telegram identity
    uid_str = _resolve_uid(
        str(req.telegram_id) if req.telegram_id else None,
        req.token
    )

    # Build payload for doc_generator (mirrors webapp_data.py schema)
    payload = req.dict(exclude={"telegram_id", "token"})

    # CPU-bound generation in thread
    try:
        from bot.services.doc_generator import generate_cv_docx
        loop = asyncio.get_running_loop()
        docx_path = await loop.run_in_executor(
            None, generate_cv_docx, payload
        )
    except Exception as e:
        logger.error(f"/api/generate_cv build error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"CV yaratishda xato: {str(e)[:200]}")

    if not docx_path or not os.path.exists(docx_path):
        raise HTTPException(status_code=500, detail="CV fayl yaratilmadi")

    with open(docx_path, "rb") as fh:
        docx_bytes = fh.read()
    _cleanup(docx_path)

    # ── Background Telegram delivery ────────────────────────────────────
    if uid_str:
        from bot.services.user_service import get_chat_id, increment_file_count
        chat_id = get_chat_id(int(uid_str)) or int(uid_str)
        increment_file_count(int(uid_str), "CV Generator")

        async def _send_cv():
            try:
                buf = io.BytesIO(docx_bytes)
                safe = (req.name or "CV").replace(" ", "_")[:30]
                buf.name = f"DASTYOR_CV_{safe}_{ts}_@DastyorAiBot.docx"
                ok = await send_docx_with_confirmation(
                    application.bot,
                    chat_id,
                    buf,
                    filename=buf.name,
                    caption=(
                        f"✅ <b>CV tayyor!</b>\n"
                        f"📄 <b>{req.name or 'CV'}</b>\n"
                        f"🆆 Shablon: <i>{req.template}</i>\n"
                        f"📎 Veb-saytdan ham yuklab olishingiz mumkin."
                    ),
                    parse_mode="HTML",
                )
                if ok:
                    logger.info(f"CV DOCX sent to Telegram chat {chat_id}")
            except Exception as tg_err:
                logger.warning(f"CV Telegram send failed (non-fatal): {tg_err}")

        asyncio.create_task(_send_cv())

    # ── Stream to browser ───────────────────────────────────────────────
    safe_name = (req.name or "CV").replace(" ", "_")[:30]
    filename = f"DASTYOR_CV_{safe_name}_{ts}_@DastyorAiBot.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ═══════════════════════════════════════════════════════════════════════════
# /api/get_oby_data — Fetch temporary parsed data for Obyektivka
# ═══════════════════════════════════════════════════════════════════════════
@app.get("/api/get_oby_data")
async def api_get_oby_data(
    token: str = Query(None),
    telegram_id: str = Query(None),
):
    uid = _resolve_uid(telegram_id, token)
    if not uid:
        raise HTTPException(status_code=401, detail="Foydalanuvchi aniqlanmadi")
    
    # 1. Try persistent user_profiles storage (survives server restarts)
    from bot.services.user_service import get_pending_oby_data
    data = get_pending_oby_data(uid)
    if data:
        return {"ok": True, "data": data}

    # 2. Fallback: temp file (legacy / same-session)
    path = f"temp/oby_data_{uid}.json"
    if os.path.exists(path):
        import json
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            return {"ok": True, "data": data}
        except Exception as e:
            logger.error(f"Error reading oby_data: {e}")
    return {"ok": False}

# ═══════════════════════════════════════════════════════════════════════════
# /api/generate_obyektivka — Obyektivka DOCX generator (WORD ONLY)
# ═══════════════════════════════════════════════════════════════════════════
class ObyektivkaRequest(BaseModel):
    telegram_id : Optional[int] = None
    token       : Optional[str] = None
    # format field kept for backward‑compat but ignored (WORD only now)
    format      : str = "word"
    lang        : str = "uz_lat"    # uz_lat | uz_cyr | en | ru
    # personal info
    fullname    : str = ""
    birthdate   : str = ""
    birthplace  : str = ""
    nation      : str = ""
    party       : str = ""
    education   : str = ""
    graduated   : str = ""
    specialty   : str = ""
    degree      : str = ""
    scientific_title: str = ""
    languages   : str = ""
    military_rank: str = ""
    awards      : str = ""
    deputy      : str = ""
    address     : str = ""
    phone       : str = ""
    work_experience: list = []
    relatives   : list = []
    # Optional photo: data URL (base64) from webapp (e.g. "data:image/jpeg;base64,...")
    photo_data  : Optional[str] = None
    # Optional: explicitly provided current job (to show under name)
    current_job : Optional[str] = None
    current_job_year: Optional[str] = None

@app.post("/api/generate_obyektivka")
async def api_generate_obyektivka(req: ObyektivkaRequest):
    """
    Generate an Obyektivka (Ma’lumotnoma) DOCX from website form data.
    • Returns the file as a streaming download for the browser.
    • Sends the same file to the user’s Telegram chat in the background.
    NOTE: PDF format is no longer supported; all outputs are DOCX.
    """
    ts = int(time.time())
    os.makedirs("temp", exist_ok=True)

    uid_str = _resolve_uid(
        str(req.telegram_id) if req.telegram_id else None,
        req.token
    )

    doc_data = {
        "lang"            : req.lang,
        # Keep Word output visually consistent with HTML preview placeholders
        "fullname"        : req.fullname or "FAMILIYA ISM SHARIF",
        "birthdate"       : req.birthdate,
        "birthplace"      : req.birthplace,
        "nation"          : req.nation,
        "party"           : req.party,
        "education"       : req.education,
        "graduated"       : req.graduated,
        "specialty"       : req.specialty,
        "degree"          : req.degree,
        "scientific_title": req.scientific_title,
        "languages"       : req.languages,
        "military_rank"   : req.military_rank,
        "awards"          : req.awards,
        "deputy"          : req.deputy,
        "address"         : req.address,
        "phone"           : req.phone,
        "work_experience" : req.work_experience,
        "current_job"     : req.current_job,
        "current_job_year": req.current_job_year,
        "relatives"       : req.relatives,
    }

    # ── Optional photo handling (data URL → temp file path) ───────────────
    photo_path = None
    try:
        if req.photo_data and isinstance(req.photo_data, str) and req.photo_data.startswith("data:image/"):
            header, b64 = req.photo_data.split(",", 1)
            # header: data:image/png;base64
            mime = header.split(";")[0].split(":")[1].lower()
            ext = {
                "image/png": "png",
                "image/jpeg": "jpg",
                "image/jpg": "jpg",
                "image/webp": "webp",
            }.get(mime, "png")
            raw = base64.b64decode(b64)
            photo_path = os.path.join("temp", f"oby_photo_{ts}.{ext}")
            with open(photo_path, "wb") as f:
                f.write(raw)
    except Exception as e:
        logger.warning(f"/api/generate_obyektivka photo decode failed: {e}")
        photo_path = None

    try:
        from bot.services.doc_generator import generate_obyektivka_docx
        loop = asyncio.get_running_loop()
        docx_path = await loop.run_in_executor(
            None, generate_obyektivka_docx, doc_data, photo_path
        )
    except Exception as e:
        logger.error(f"/api/generate_obyektivka build error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Obyektivka yaratishda xato: {str(e)[:200]}")

    if not docx_path or not os.path.exists(docx_path):
        raise HTTPException(status_code=500, detail="Obyektivka fayl yaratilmadi")

    with open(docx_path, "rb") as fh:
        file_bytes = fh.read()
    _cleanup(docx_path)
    if photo_path:
        _cleanup(photo_path)

    ext = "docx"
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    # ── Background Telegram delivery ────────────────────────────────────
    if uid_str:
        from bot.services.user_service import get_chat_id, increment_file_count
        chat_id = get_chat_id(int(uid_str)) or int(uid_str)
        increment_file_count(int(uid_str), "Obyektivka Generator")

        async def _send_oby():
            try:
                buf = io.BytesIO(file_bytes)
                safe = (req.fullname or "Obyektivka").replace(" ", "_")[:30]
                buf.name = f"DASTYOR_Obyektivka_{safe}_{ts}_@DastyorAiBot.{ext}"
                ok = await send_docx_with_confirmation(
                    application.bot,
                    chat_id,
                    buf,
                    filename=buf.name,
                    caption=(
                        f"✅ <b>Obyektivka tayyor!</b>\n"
                        f"👤 <b>{req.fullname or 'Ma’lumotnoma'}</b>\n"
                        f"📎 Format: <i>{ext.upper()}</i>\n"
                        f"📥 Veb-saytdan ham yuklab olishingiz mumkin."
                    ),
                    parse_mode="HTML",
                )
                if ok:
                    logger.info(f"Obyektivka sent to Telegram chat {chat_id}")
            except Exception as tg_err:
                logger.error(f"Obyektivka Telegram yuborishda xato (send_document): {tg_err}", exc_info=True)

        asyncio.create_task(_send_oby())

    # ── Stream to browser ───────────────────────────────────────────────
    safe_name = (req.fullname or "Obyektivka").replace(" ", "_")[:30]
    filename = f"DASTYOR_Obyektivka_{safe_name}_{ts}_@DastyorAiBot.{ext}"
    return StreamingResponse(
        io.BytesIO(file_bytes),
        media_type=mime,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )




# ═══════════════════════════════════════════════════════════════════════════
# /api/export_cv  — Server-side CV export (PDF or Word)
# ═══════════════════════════════════════════════════════════════════════════
class ExportCVRequest(BaseModel):
    telegram_id : Optional[int]  = None
    token       : Optional[str]  = None
    format      : str            = "pdf"   # "pdf" | "word"
    lang        : str            = "uz_lat"
    # CV fields (mirrors CVRequest above)
    name        : str  = ""
    spec        : str  = ""
    phone       : str  = ""
    email       : str  = ""
    loc         : str  = ""
    addr        : str  = ""
    birth       : str  = ""
    place       : str  = ""
    nation      : str  = ""
    langs       : str  = ""
    about       : str  = ""
    template    : str  = "minimal"
    works       : list = []
    education_list: list = []
    skills      : str  = ""
    img         : str  = ""   # absolute URL of profile photo (optional)

@app.post("/api/export_cv")
async def api_export_cv(req: ExportCVRequest):
    """
    Server-side CV export.
    Renders the SAME Jinja2 template used for the browser preview, so
    the exported file is pixel-perfect regardless of client environment.
    """
    ts = int(time.time())
    uid_str = _resolve_uid(str(req.telegram_id) if req.telegram_id else None, req.token)
    fmt = req.format.lower()
    data = req.dict(exclude={"telegram_id", "token", "format"})
    safe = safe_filename(req.name or "CV")
    bot_suffix = "_@DastyorAiBot"

    if fmt == "word":
        # ── Word export — real .docx (python-docx, mobile-compatible) ──
        filename = f"DASTYOR_CV_{safe}_{ts}{bot_suffix}.docx"
        from bot.services.doc_generator import generate_cv_docx
        loop = asyncio.get_running_loop()
        docx_path = await loop.run_in_executor(None, generate_cv_docx, data)
        if not docx_path or not os.path.exists(docx_path):
            raise HTTPException(status_code=500, detail="Word fayl yaratishda xato")
        with open(docx_path, "rb") as fh:
            file_bytes = fh.read()
        try:
            os.remove(docx_path)
        except Exception:
            pass
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        # ── PDF export (Playwright → exact browser render) ─────────────
        filename = f"DASTYOR_CV_{safe}_{ts}{bot_suffix}.pdf"
        pdf_bytes = await generate_cv_pdf(data, base_url=SITE_BASE_URL)
        if not pdf_bytes:
            # Playwright unavailable → fall back to python-docx PDF
            from bot.services.doc_generator import generate_cv_docx, convert_to_pdf_safe
            loop = asyncio.get_running_loop()
            docx_path = await loop.run_in_executor(None, generate_cv_docx, data)
            pdf_path = convert_to_pdf_safe(docx_path) if docx_path else None
            if pdf_path and os.path.exists(pdf_path):
                with open(pdf_path, "rb") as fh:
                    pdf_bytes = fh.read()
                for p in [pdf_path, docx_path]:
                    try: os.remove(p)
                    except: pass
            else:
                logger.error("All PDF generation methods failed inside api_webhook.py")
                raise HTTPException(status_code=500, detail="PDF yaratishda xato")
        file_bytes = pdf_bytes
        media_type = "application/pdf"

    # ── Background Telegram delivery ────────────────────────────────────
    if uid_str:
        from bot.services.user_service import increment_file_count
        increment_file_count(int(uid_str), f"CV Export {fmt.upper()}")

        async def _send():
            try:
                buf = io.BytesIO(file_bytes)
                buf.name = filename
                chat_id = int(uid_str)
                if filename.lower().endswith(".docx"):
                    await send_docx_with_confirmation(
                        application.bot,
                        chat_id,
                        buf,
                        filename=filename,
                        caption=(
                            f"✅ <b>CV tayyor!</b>\n"
                            f"👤 <b>{req.name}</b>  · 🎨 {req.template}\n"
                            f"📎 <code>{filename}</code>"
                        ),
                        parse_mode="HTML",
                    )
                else:
                    await application.bot.send_document(
                        chat_id=chat_id,
                        document=InputFile(buf, filename=filename),
                        caption=(
                            f"✅ <b>CV tayyor!</b>\n"
                            f"👤 <b>{req.name}</b>  · 🎨 {req.template}\n"
                            f"📎 <code>{filename}</code>"
                        ),
                        parse_mode="HTML",
                    )
            except Exception as tg_err:
                logger.warning(f"CV export Telegram send failed: {tg_err}")

        asyncio.create_task(_send())

    return StreamingResponse(
        io.BytesIO(file_bytes),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ═══════════════════════════════════════════════════════════════════════════
# /api/export_obyektivka  — Server-side Obyektivka export (WORD ONLY)
# ═══════════════════════════════════════════════════════════════════════════
class ExportObyektivkaRequest(BaseModel):
    telegram_id    : Optional[int] = None
    token          : Optional[str] = None
    # format field kept for compatibility but ignored; always DOCX
    format         : str           = "word"
    lang           : str           = "uz_lat"
    fullname       : str = ""
    birthdate      : str = ""
    birthplace     : str = ""
    nation         : str = ""
    party          : str = ""
    education      : str = ""
    graduated      : str = ""
    specialty      : str = ""
    degree         : str = ""
    scientific_title: str = ""
    languages      : str = ""
    military_rank  : str = ""
    awards         : str = ""
    deputy         : str = ""
    address        : str = ""
    phone          : str = ""
    work_experience: list = []
    relatives      : list = []
    # optional: data URL or public URL of photo
    photo_data     : Optional[str] = None
    # optional current job line under full name
    current_job    : Optional[str] = None
    current_job_year: Optional[str] = None

@app.post("/api/export_obyektivka")
async def api_export_obyektivka(req: ExportObyektivkaRequest):
    """
    Server-side Obyektivka export — WORD (DOCX) only.
    PDF export/conversion has been removed.
    """
    ts  = int(time.time())
    os.makedirs("temp", exist_ok=True)
    uid_str = _resolve_uid(str(req.telegram_id) if req.telegram_id else None, req.token)
    data = req.dict(exclude={"telegram_id", "token", "format"})
    safe = safe_filename(req.fullname or "Obyektivka")
    bot_suffix = "_@DastyorAiBot"

    # Optional photo_data (data URL) -> temp image for DOCX insertion.
    photo_path = None
    try:
        if req.photo_data and isinstance(req.photo_data, str) and req.photo_data.startswith("data:image/"):
            header, b64 = req.photo_data.split(",", 1)
            mime = header.split(";")[0].split(":")[1].lower()
            ext = {
                "image/png": "png",
                "image/jpeg": "jpg",
                "image/jpg": "jpg",
                "image/webp": "webp",
            }.get(mime, "png")
            raw = base64.b64decode(b64)
            photo_path = os.path.join("temp", f"oby_export_photo_{ts}.{ext}")
            with open(photo_path, "wb") as f:
                f.write(raw)
    except Exception as e:
        logger.warning(f"/api/export_obyektivka photo decode failed: {e}")
        photo_path = None

    # Har doim DOCX yaratamiz (rassmiy minimal layout bilan)
    from bot.services.doc_generator import generate_obyektivka_docx

    loop = asyncio.get_running_loop()
    docx_path = await loop.run_in_executor(None, generate_obyektivka_docx, data, photo_path)
    if not docx_path or not os.path.exists(docx_path):
        raise HTTPException(status_code=500, detail="Word fayl yaratishda xato")

    filename   = f"DASTYOR_Obyektivka_{safe}_{ts}{bot_suffix}.docx"
    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    with open(docx_path, "rb") as fh:
        file_bytes = fh.read()
    try:
        os.remove(docx_path)
    except Exception:
        pass
    if photo_path:
        try:
            os.remove(photo_path)
        except Exception:
            pass

    if uid_str:
        from bot.services.user_service import increment_file_count
        increment_file_count(int(uid_str), "Obyektivka Export WORD")

        async def _send_oby():
            try:
                buf = io.BytesIO(file_bytes)
                buf.name = filename
                await send_docx_with_confirmation(
                    application.bot,
                    int(uid_str),
                    buf,
                    filename=filename,
                    caption=(
                        f"✅ <b>Obyektivka tayyor!</b>\n"
                        f"👤 <b>{req.fullname}</b>\n"
                        f"📎 <code>{filename}</code>"
                    ),
                    parse_mode="HTML",
                )
            except Exception as e:
                logger.warning(f"Obyektivka export Telegram send failed: {e}")

        asyncio.create_task(_send_oby())

    return StreamingResponse(
        io.BytesIO(file_bytes),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ═══════════════════════════════════════════════════════════════════════════
# /api/preview_obyektivka — HTML preview for webapp (A4 canvas)
# Returns rendered obyektivka_template.html so the frontend can show
# a real-time preview identical to the exported Word/PDF design.
# ═══════════════════════════════════════════════════════════════════════════
class PreviewObyektivkaRequest(BaseModel):
    lang           : str = "uz_lat"
    fullname       : str = ""
    birthdate      : str = ""
    birthplace     : str = ""
    nation         : str = ""
    party          : str = ""
    education      : str = ""
    graduated      : str = ""
    specialty      : str = ""
    degree         : str = ""
    scientific_title: str = ""
    languages      : str = ""
    military_rank  : str = ""
    awards         : str = ""
    deputy         : str = ""
    address        : str = ""
    phone          : str = ""
    work_experience: list = []
    relatives      : list = []


@app.post("/api/preview_obyektivka", response_class=HTMLResponse)
async def api_preview_obyektivka(req: PreviewObyektivkaRequest):
    """
    Render obyektivka_template.html with given data and return raw HTML.
    Used only by the webapp for live A4 preview; does NOT send anything
    to Telegram and does not persist data.
    """
    from bot.services.render_service import render_obyektivka_html

    html = render_obyektivka_html(req.dict())
    return HTMLResponse(content=html)


@app.post("/webhook")

async def webhook(request: Request):
    """
    Handle incoming Telegram updates via POST request.
    """
    try:
        data = await request.json()
        update = Update.de_json(data, application.bot)
        await application.process_update(update)
        return Response(status_code=200)
    except Exception as e:
        logger.error(f"Error processing update: {e}")
        return Response(status_code=500)

if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
